from flask import Flask, render_template, request, redirect, url_for, session, g, flash, send_file, jsonify
from datetime import datetime, date, timedelta
from functools import wraps
from werkzeug.utils import secure_filename
from markupsafe import Markup, escape
from io import BytesIO
import os
import json
import re

from database import (get_db, init_db, hash_password, get_setting, set_setting,
                       get_test_price, find_or_create_doctor, find_or_create_referral_center,
                       get_report_template, save_report_template,
                       get_examining_doctors, set_examining_doctors, add_examining_doctor,
                       get_examining_doctors_full, add_examining_doctor_full,
                       update_examining_doctor, delete_examining_doctor, move_examining_doctor,
                       get_letterhead_doctors,
                       get_examining_tests, get_examining_rates_map,
                       set_examining_doctor_rate, compute_examining_doctor_fee,
                       find_reference_range,
                       save_saved_report, get_saved_report, search_saved_reports)
from translations import t
from barcode_gen import generate_code39, generate_code128, generate_qr
import astm_host
import license_manager
import auto_updater
import secrets

app = Flask(__name__)
# مفتاح جلسة عشوائي مختلف بكل مرة يُشغَّل فيها السيرفر فعليًا (وليس نفس
# نص ثابت دائمًا) — بهذا الشكل أي كوكي دخول قديم صار غير صالح تلقائيًا بعد
# أي إعادة تشغيل حقيقية للبرنامج (إعادة تشغيل الجهاز، إيقاف ثم تشغيل من
# جديد...)، فيُطلب اسم المستخدم وكلمة المرور من جديد بشكل مضمون، بدل أن
# يبقى المتصفح مسجّل دخول تلقائيًا لأشهر لمجرد أن الكوكي القديم لسا شغّال.
app.secret_key = secrets.token_hex(32)
app.config["PERMANENT_SESSION_LIFETIME"] = 60 * 60 * 24 * 30  # 30 days when "remember me" is checked

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "static", "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)
ALLOWED_LOGO_EXT = {"png", "jpg", "jpeg", "gif", "svg", "webp"}
# صور خلفية شاشة الترحيب (dashboard) — صور فوتوغرافية بس، بدون svg/gif
# (svg ما إله فايدة كخلفية ممتدة، وgif المتحرك يشتت الانتباه بشاشة ترحيب).
ALLOWED_DASHBOARD_BG_EXT = {"png", "jpg", "jpeg", "webp"}

ROLE_LABELS = {
    "admin": "Administrator",
    "reception": "Reception",
    "technician": "Lab Technician",
    "supervisor": "Lab Supervisor",
    "accountant": "Accountant",
}

# مدد الترخيص الجاهزة المعروضة بلوحة المصمم عند توليد كود لعميل — نقطة
# #12. المفتاح هو قيمة <option value="..."> بفورم designer_generate،
# والقيمة (label, days): days=None تعني ترخيص دائم فعلي (بدون تاريخ
# انتهاء إطلاقاً، يُخزَّن كـ "PERM" داخل الكود نفسه — راجع
# license_manager.generate_activation_code). التعديل هنا فقط ينعكس تلقائياً
# على القائمة المنسدلة بـ designer/panel.html.
LICENSE_DURATION_PRESETS = {
    "trial": ("تجريبي (3 أيام)", license_manager.TRIAL_DAYS),
    "1day": ("يوم واحد", 1),
    "3days": ("3 أيام", 3),
    "1month": ("شهر", 30),
    "3months": ("3 أشهر", 90),
    "6months": ("6 أشهر", 180),
    "1year": ("سنة", 365),
    "permanent": ("دائم (بدون تاريخ انتهاء)", None),
}
app.jinja_env.globals["LICENSE_DURATION_PRESETS"] = LICENSE_DURATION_PRESETS

# دكتور المختبر الفاحص — the examining lab doctor for a visit (separate from
# the external "referring" doctor already tracked via doctor_id/doctors
# table). Editable now from Management → Settings; this constant is only the
# bootstrap default used the first time (see database.get_examining_doctors).
EXAMINING_DOCTORS = ["د.خليل حمود", "د.هدى نصيف", "د.اسراء عبد الاقر"]

# ------------------------------------------------------------- print reports --
# Maps a test_definitions.code to the printable report template that
# reproduces the lab's paper letterhead for that report type.
REPORT_TEMPLATE_MAP = {
    "CBC": "reports/cbc.html",
    "BF": "reports/blood_film.html",
    "BFRETIC": "reports/blood_film.html",
    "RETIC": "reports/retic_only.html",
    "WBCDIFF": "reports/wbc_differential.html",
    "FLUIDEXAM": "reports/fluid_examination.html",
    "COAG": "reports/coagulation.html",
}

# Blood Film and Retic Count are two SEPARATE orderable tests, but when both
# are ordered for the same visit they should print as ONE Blood-Film-shaped
# report (which already has the Reticulocyte/Corrected Retic fields built
# in) instead of two separate pages. If Retic Count is ordered on its own
# (no Blood Film for that visit), it prints its own small standalone report.
BF_RETIC_LINK = {"BF": "RETIC", "RETIC": "BF"}

# Test codes whose printed report gets a blank stamp/signature space for the
# examining doctor (Blood Film, Blood Film + Retic, standalone Retic, Fluid
# examination, Hb Preparation, Sickling, BMA, BM Biopsy, WBC Differential) —
# kept in sync with EXAMINING_TEST_DEFS (the tests a doctor is paid to
# personally examine), since every one of those needs a stamp/signature line
# on its printed report. CBC and the other non-examining report types don't
# get this box.
EXAM_SIGNATURE_TEST_CODES = {"BF", "BFRETIC", "RETIC", "FLUIDEXAM", "HBPREP", "SICKLE", "BMA", "BMBIOPSY", "WBCDIFF"}

# CBC results are grouped on the printed report with a blank spacer row
# between each group, matching the reference report layout.
CBC_ROW_GROUPS = [
    ["WBCs", "NE", "Ly", "MO", "BA", "EO", "NE#", "LY#", "MO#", "BA#", "EO#"],
    ["RBC", "HGB", "HCT", "MCV", "MCH", "MCHC", "RDW", "RDW-SD"],
    ["PLT", "MPV"],
]

# الوحدة الثانية لنتيجة التحليل (مثلاً mg/dL بالإضافة لـ mmol/L) — بعض
# التحاليل تُكتب نتيجتها بوحدتين بنفس الوقت (خصوصًا الكيمياء والهرمونات).
# tp.unit2 و tp.unit2_factor يُضبطان مرة وحدة من صفحة إدارة الوحدات لكل
# باراميتر (اختياري تمامًا)، والقيمة الثانية تُحسب تلقائيًا وقت الطباعة:
# value2 = value1 * unit2_factor. لتحويل عكسي (وحدة أساسية أصغر لوحدة ثانية
# أكبر) استخدم معامل أصغر من 1 بدل قسمة منفصلة — النتيجة نفسها رياضيًا.
def format_unit2_value(value, factor):
    """يرجع نص القيمة المحوّلة للوحدة الثانية، أو None إذا ما ينطبق التحويل
    (نتيجة غير رقمية، أو ما فيه معامل محفوظ لهذا الباراميتر)."""
    if factor in (None, "") or value in (None, ""):
        return None
    try:
        converted = float(value) * float(factor)
    except (TypeError, ValueError):
        return None
    if converted == int(converted):
        return str(int(converted))
    return f"{converted:.3f}".rstrip("0").rstrip(".")


# يحوّل حقل range_text (نفس الحقل الحر الموجود أصلاً بجدول reference_ranges)
# إلى قائمة مستويات جاهزة للعرض بشكل "Normal Range" مفصّل بالتقرير، بدل
# كتابته كسطر نص واحد فقط. الصيغة المتوقعة بخانة range_text: كل مستوى
# بسطر مستقل "التسمية: القيمة" — مثال (زي Triglycerides بالنموذج المرجعي):
#   Borderline high: 150 - 199
#   Normal: Less than 150
#   High: 200 - 499
#   Very high: More than 500
# وسطر بدون ":" (أو حقل range_text بسطر وحيد فقط، مثل "166 - 507") يُعرض
# كقيمة بسيطة بلا تسمية — بنفس شكل باقي التحاليل العادية (Cortisol, TSH...).
def parse_range_tiers(text):
    if not text:
        return []
    tiers = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if ":" in line:
            label, _, value = line.partition(":")
            tiers.append({"label": label.strip(), "value": value.strip()})
        else:
            tiers.append({"label": None, "value": line})
    return tiers


ALLOWED_DOCX_EXT = {"docx"}

# نص الكليشة الجاهز لزر "Normal" بحقول وصف RBC/WBC/Platelets بتقرير الـ Blood
# film (شاشة إدخال النتائج المفردة والمجمّعة). عدّل النص هنا فقط وينعكس
# تلقائيًا بكل شاشات الإدخال دون أي تعديل إضافي بالقوالب. الزر يعبّي الحقل
# بهذا النص، ويبقى قابلاً للتعديل الكامل بعدها (نفس أي حقل نصي عادي).
NORMAL_CLICHE_TEXT = {
    "RBC_desc": "Normochromic normocytic",
    "WBC_desc": "Normal count and morphology",
    "Platelets_desc": "Within normal limits",
}
app.jinja_env.globals["NORMAL_CLICHE_TEXT"] = NORMAL_CLICHE_TEXT

# أسماء الأشهر بالتسمية العراقية/الشامية المتداولة (بدل كانون الثاني/يناير
# الرسمية أو January/February الإنكليزية) — نقطة #11. تُستخدم كـ Jinja
# filter بأي قالب: {{ "2026-08"|iraqi_month }} أو {{ 8|iraqi_month }} تعطي
# "آب"، و IRAQI_MONTHS متاح كـ global لبناء قوائم اختيار الأشهر (dropdown)
# بنفس التسمية مباشرة من القالب.
IRAQI_MONTHS = {
    1: "كانون الثاني", 2: "شباط", 3: "آذار", 4: "نيسان",
    5: "أيار", 6: "حزيران", 7: "تموز", 8: "آب",
    9: "أيلول", 10: "تشرين الأول", 11: "تشرين الثاني", 12: "كانون الأول",
}
app.jinja_env.globals["IRAQI_MONTHS"] = IRAQI_MONTHS


@app.template_filter("iraqi_month")
def iraqi_month_filter(value):
    """يحوّل شهرًا لاسمه العراقي. يقبل: رقم شهر مباشر (1-12)، أو نص تاريخ
    يبدأ بصيغة 'YYYY-MM' (مثل قيمة ?month المستخدمة بالتقارير الشهرية)،
    أو أي نص تاريخ/ISO يبدأ بنفس الصيغة. أي قيمة غير مفهومة تُرجع فاضية
    بدل ما تكسر الصفحة."""
    if value in (None, ""):
        return ""
    if isinstance(value, int):
        return IRAQI_MONTHS.get(value, "")
    s = str(value)
    try:
        if len(s) >= 7 and s[4] == "-":
            return IRAQI_MONTHS.get(int(s[5:7]), "")
        return IRAQI_MONTHS.get(int(s), "")
    except (ValueError, IndexError):
        return ""

# رموز التمييز الملوّنة بحقل الـ Conclusion (نجمة وأسهم وعلامتي استفهام/تعجب)
# — يضيفها المستخدم بنفسه من زر التنسيق فوق حقل الاستنتاج بشاشات إدخال
# النتائج (المفردة والمجمّعة). "key" هو المعرّف المستخدم بجدول settings
# و"label" هو الاسم العربي اللي يظهر بشاشة Management → Settings.
# CONCLUSION_MARKER_DEFAULT_COLORS تحت هو اللون الافتراضي فقط — المدير يكدر
# يغيّر أي لون من شاشة الإعدادات، والقيمة المحفوظة هناك (settings key:
# "conclusion_marker_colors", JSON) هي اللي تُستخدم فعليًا وقت الطباعة/العرض
# عبر get_conclusion_marker_colors(db). الأزرار تُدرج الرمز في موضع المؤشر
# بالضبط (بدون فرض سطر جديد)، فالرمز ممكن يطلع بأي مكان بالسطر — لهيك
# التلوين يفحص كل رمز بالسطر أينما وجد، مو بس إذا كان ببداية السطر.
CONCLUSION_MARKERS = [
    {"key": "star", "char": "★", "label": "نجمة حمراء", "default_color": "#D40000"},
    {"key": "arrow", "char": "→", "label": "سهم مفرد", "default_color": "#1F3B7A"},
    {"key": "arrow_left", "char": "◄", "label": "سهم يسار (اختياري)", "default_color": "#1F3B7A"},
    {"key": "arrow_double", "char": "⇒", "label": "سهم مزدوج", "default_color": "#1F3B7A"},
    {"key": "triangle", "char": "▶", "label": "سهم مثلث", "default_color": "#1F3B7A"},
    {"key": "question", "char": "?", "label": "علامة استفهام", "default_color": "#E67E22"},
    {"key": "exclaim", "char": "!", "label": "علامة تعجب", "default_color": "#8E44AD"},
]
CONCLUSION_MARKER_DEFAULT_COLORS = {m["char"]: m["default_color"] for m in CONCLUSION_MARKERS}
_CONCLUSION_MARKER_PATTERN = re.compile(
    "|".join(re.escape(m) for m in CONCLUSION_MARKER_DEFAULT_COLORS)
)
_HEX_COLOR_RE = re.compile(r"^#[0-9a-fA-F]{6}$")


def get_conclusion_marker_colors(db):
    """يرجع dict {رمز: لون hex} بعد دمج الألوان الافتراضية مع أي تعديل
    حفظه المدير من شاشة Management → Settings (مخزّن كـ JSON بجدول
    settings تحت المفتاح conclusion_marker_colors: {key: hex})."""
    colors = dict(CONCLUSION_MARKER_DEFAULT_COLORS)
    raw = get_setting(db, "conclusion_marker_colors", "")
    if raw:
        try:
            saved = json.loads(raw)
        except (ValueError, TypeError):
            saved = {}
        for m in CONCLUSION_MARKERS:
            val = saved.get(m["key"])
            if val and _HEX_COLOR_RE.match(val):
                colors[m["char"]] = val
    return colors


@app.template_filter("render_conclusion")
def render_conclusion_filter(text):
    """يحوّل نص حقل الـ Conclusion (بأسطره كما كتبها المستخدم بالضبط) إلى
    HTML آمن للطباعة: كل رمز نجمة/سهم بأي موضع بالسطر يُلوَّن الرمز نفسه
    فقط بلونه (المحفوظ من شاشة الإعدادات أو الافتراضي)، وباقي السطر يبقى
    بلونه الطبيعي، والأسطر تُفصل بـ <br> حتى يطبع كل سطر لحاله بدل ما
    ينضغط كلها بسطر وحد متل الـ HTML العادي. كل النص غير الرمز يُهرَّب
    (escape) بشكل طبيعي — ما في أي HTML يُنفَّذ من داخل النص المكتوب نفسه."""
    if not text:
        return ""
    colors = get_conclusion_marker_colors(get_db())
    lines = str(text).split("\n")
    rendered = []
    for line in lines:
        pos = 0
        parts = []
        for m in _CONCLUSION_MARKER_PATTERN.finditer(line):
            if m.start() > pos:
                parts.append(str(escape(line[pos:m.start()])))
            color = colors[m.group()]
            parts.append(
                '<span class="conclusion-marker" style="color:'
                + color + ';">' + str(escape(m.group())) + "</span>"
            )
            pos = m.end()
        if pos < len(line):
            parts.append(str(escape(line[pos:])))
        rendered.append("".join(parts))
    return Markup("<br>".join(rendered))

# The department names that get the extra "previous result per parameter"
# column on the printed report (in addition to the previous-visit date that
# every report shows), AND that get the printed report's repeating letterhead
# when a result overflows onto a second page (see repeat_header_on_print in
# print_report). Covers Chemistry, Hormones, Vitamins, Tumor Markers,
# Virology and Coagulation — the quantifiable/serology departments the lab
# wants tracked over time / laid out as multi-line result cards. Hematology
# (CBC, Blood Film, WBC Differential, Fluid examination, and the other big
# dedicated-template reports) never goes through either feature — they don't
# call find_previous_reference and print_report never sets
# repeat_header_on_print for them — so they're unaffected either way. Matches
# loosely (Arabic or English, any casing) so it keeps working no matter how
# an admin later spells a new department.
PREVIOUS_VALUE_DEPARTMENT_KEYWORDS = (
    "chem", "كيمياء",
    "coagul", "تخثر",
    "hormone", "هرمون",
    "vitamin", "فيتامين",
    "virology", "viral", "فايروس", "فيروس",
    "tumor", "tumour", "marker", "دلائل", "ورمي",
)


def department_shows_previous_values(department):
    d = (department or "").lower()
    return any(k in d for k in PREVIOUS_VALUE_DEPARTMENT_KEYWORDS)


# الترتيب الافتراضي لتحاليل شاشة "إدخال النتائج" حسب القسم — يُستخدم فقط
# لأي تحليل غير مذكور صراحةً بإعداد "ترتيب التحاليل بشاشة إدخال النتائج"
# اليدوي (results_entry_test_order)؛ لو ذاك الإعداد فاضي بالكامل، يصير هذا
# الترتيب هو الافتراضي الوحيد. مطابقة كلمات مفتاحية (عربي/انكليزي، أي حالة
# أحرف) نفس أسلوب PREVIOUS_VALUE_DEPARTMENT_KEYWORDS فوق — تشتغل بغض النظر
# شلون كتب الأدمن اسم القسم بالضبط بكتالوج التحاليل. أي قسم ما يطابق ولا
# مجموعة (مثلاً قسم نادر أو Other) يظهر بالأخير.
# الترتيب المطلوب: كيمياء ← فايروسات ← هرمونات ← فيتامينات ← تخثر ← أمراض الدم.
DEPARTMENT_PRIORITY_KEYWORDS = [
    ("chem", "كيمياء"),                              # 1) Chemistry
    ("virology", "viral", "فايروس", "فيروس"),        # 2) Virology
    ("hormone", "هرمون"),                             # 3) Hormones
    ("vitamin", "فيتامين"),                           # 4) Vitamins
    ("coagul", "تخثر"),                               # 5) Coagulation
    ("hemat", "دم"),                                  # 6) Hematology / blood diseases
]


def department_priority_rank(department):
    d = (department or "").lower()
    for i, keywords in enumerate(DEPARTMENT_PRIORITY_KEYWORDS):
        if any(k in d for k in keywords):
            return i
    return len(DEPARTMENT_PRIORITY_KEYWORDS)  # قسم غير معروف/غير مذكور أعلاه — يظهر بالأخير


def get_visit_hct(db, visit_id):
    """يجيب آخر قيمة HCT مُدخلة ضمن تحليل CBC لنفس الزيارة (إن وجدت)، تُستخدم
    لحساب Corrected Retic count تلقائيًا من Reticulocyte count. يرجع None إذا
    ما كان فيه CBC بعد أو ما دخلت قيمة HCT."""
    row = db.execute(
        "SELECT r.value_numeric FROM results r "
        "JOIN test_parameters tp ON tp.id = r.test_parameter_id "
        "JOIN order_tests ot ON ot.id = r.order_test_id "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id "
        "WHERE o.visit_id = ? AND td.code = 'CBC' AND tp.name = 'HCT' "
        "AND r.value_numeric IS NOT NULL "
        "ORDER BY r.id DESC LIMIT 1",
        (visit_id,),
    ).fetchone()
    return row["value_numeric"] if row else None

def get_normal_hct(db, gender, age, age_unit):
    """يجيب متوسط المدى الطبيعي لـ HCT (من جدول القيم المرجعية reference_ranges)
    حسب عمر وجنس المريض — يُستخدم أساسًا لحساب Corrected Retic count بدل رقم
    ثابت واحد لكل الأعمار. يرجع None إذا ما كان فيه مدى مرجعي مطابق."""
    param = db.execute(
        "SELECT tp.id FROM test_parameters tp "
        "JOIN test_definitions td ON td.id = tp.test_definition_id "
        "WHERE td.code = 'CBC' AND tp.name = 'HCT' LIMIT 1"
    ).fetchone()
    if not param:
        return None
    rng = find_reference_range(db, param["id"], gender, age, age_unit)
    if not rng or rng["low"] is None or rng["high"] is None:
        return None
    return (rng["low"] + rng["high"]) / 2.0


def find_previous_reference(db, full_name, age, gender, test_definition_id, department,
                             exclude_visit_id, before_created_at):
    """Find this patient's most recent EARLIER visit that has a completed
    result for the same test — matched strictly on full name + age + gender
    (not just patients.id, since reception may have re-registered a
    returning patient as a new row) and dated before the current visit.

    Returns (previous_visit_date_str_or_None, {param_name: previous_value}).
    The values dict is only populated for Chemistry/Coagulation-type
    departments; every other department gets the date only.
    """
    row = db.execute(
        "SELECT ot.id as order_test_id, v.created_at as visit_created_at "
        "FROM order_tests ot "
        "JOIN orders o ON o.id = ot.order_id "
        "JOIN visits v ON v.id = o.visit_id "
        "JOIN patients p ON p.id = v.patient_id "
        "WHERE ot.test_definition_id = ? "
        "AND LOWER(TRIM(p.full_name)) = LOWER(TRIM(?)) "
        "AND p.age = ? AND p.gender = ? "
        "AND v.id != ? AND v.created_at < ? "
        "AND ot.status IN ('Completed', 'Verified') "
        "ORDER BY v.created_at DESC LIMIT 1",
        (test_definition_id, full_name or "", age, gender, exclude_visit_id, before_created_at or ""),
    ).fetchone()
    if not row:
        return None, {}

    try:
        dt = datetime.fromisoformat(row["visit_created_at"])
        prev_date = f"{dt.day}/{dt.month}/{dt.year}"
    except (TypeError, ValueError):
        prev_date = row["visit_created_at"] or ""

    prev_values = {}
    if department_shows_previous_values(department):
        prev_results = db.execute(
            "SELECT r.*, tp.name as param_name FROM results r "
            "JOIN test_parameters tp ON tp.id = r.test_parameter_id WHERE r.order_test_id=?",
            (row["order_test_id"],),
        ).fetchall()
        for r in prev_results:
            val = r["value_text"] if r["value_text"] not in (None, "") else r["value_numeric"]
            prev_values[r["param_name"]] = "" if val is None else val
    return prev_date, prev_values


# ---------------------------------------------------------------- helpers --
def current_lang():
    return session.get("lang", "en")


@app.context_processor
def inject_globals():
    db = get_db()
    lang = current_lang()
    brand_name = get_setting(db, "app_name_ar" if lang == "ar" else "app_name", t(lang, "app_name"))
    logo_path = get_setting(db, "logo_path", "")
    logo_url = url_for("static", filename=logo_path) if logo_path else None
    # خلفية شاشة الترحيب (dashboard) — اختيارية، تُرفع من Management → الإعدادات
    # (بطاقة "خلفية شاشة الترحيب")، نفس أسلوب الشعار logo_path أعلاه بالضبط.
    # محقونة هنا عالمياً حتى تتوفر بـ dashboard.html دون تمريرها يدويًا من route.
    dashboard_bg_path = get_setting(db, "dashboard_bg_path", "")
    dashboard_bg_url = url_for("static", filename=dashboard_bg_path) if dashboard_bg_path else None
    # عنوان وهاتف المختبر — اختياريان، يُضبطان مرة وحدة من الإعدادات (بطاقة
    # العلامة التجارية) ويظهران تلقائيًا بأي قالب يحتاجهم (خصوصًا فاتورة
    # A5 — نقطة #10) بدون تمريرهما يدويًا من كل route.
    lab_address = get_setting(db, "lab_address", "")
    lab_phone = get_setting(db, "lab_phone", "")
    # ارتفاع/عرض خلايا جدول نتائج التقرير (Test/Result/Control/Unit...) — يتحكم
    # بيها الأدمن من Management → Settings، تنطبق تلقائيًا على كل التقارير
    # لأنها محقونة هنا بدل تمريرها يدويًا بكل route.
    report_row_pad = get_setting(db, "report_row_pad", "5")
    report_col_pad = get_setting(db, "report_col_pad", "12")
    # دكاترة الفحص المُفعَّل لهم "إظهار بترويسة التقرير" — تُحقن هنا تلقائيًا
    # حتى تنعرض بترويسة أي تقرير مطبوع (reports/*.html) بدون تمريرها يدويًا
    # من كل route. راجع Management → الإعدادات → إدارة قائمة الدكاترة.
    letterhead_doctors = get_letterhead_doctors(db)
    letterhead_font_size = get_setting(db, "letterhead_font_size", "14")
    letterhead_font_family = get_setting(db, "letterhead_font_family", "Segoe UI, Tahoma, Arial, sans-serif")
    # حجم خط اسم التحليل وحجم خط النتيجة بكل التقارير المطبوعة — نفس أسلوب
    # letterhead_font_size أعلاه: محقونة هنا مرة وحدة فتنطبق تلقائيًا على كل
    # قوالب reports/* (وbase_report.html) دون تمريرها يدويًا من كل route.
    test_name_font_size = get_setting(db, "test_name_font_size", "16")
    result_value_font_size = get_setting(db, "result_value_font_size", "16")
    # موضع/حجم الشعار بترويسة كل تقرير مطبوع — إعدادان عامان محقونان هنا
    # مرة وحدة، بنفس أسلوب باقي إعدادات التقرير.
    logo_position = get_setting(db, "logo_position", "right")
    logo_width = get_setting(db, "logo_width", "100")
    # حالة "التحديث التلقائي" الحالية — تُحقن هنا حتى يظهر زر التبديل بشريط
    # الأعلى (base.html، بجانب اسم المستخدم) بأي صفحة بدون تمريرها يدويًا.
    auto_update_enabled = get_setting(db, "auto_update_enabled", "1") == "1"
    license_banner = None
    if "user_id" in session:
        lic = license_manager.check_license(db)
        # الشريط التحذيري "متبقي كم يوم" يظهر فقط للترخيص التجريبي، وليس
        # لأي ترخيص عادي (حتى لو كان له تاريخ انتهاء بعد سنة مثلاً).
        if lic["status"] == "active" and lic.get("is_trial") and lic.get("days_left") is not None:
            license_banner = f"متبقي {lic['days_left']} يوم على انتهاء الفترة التجريبية"
    db.close()
    return dict(t=lambda key: t(lang, key), lang=lang,
                current_user=session.get("full_name"), current_role=session.get("role"),
                brand_name=brand_name, logo_url=logo_url,
                dashboard_bg_url=dashboard_bg_url,
                lab_address=lab_address, lab_phone=lab_phone,
                report_row_pad=report_row_pad, report_col_pad=report_col_pad,
                letterhead_doctors=letterhead_doctors,
                letterhead_font_size=letterhead_font_size,
                letterhead_font_family=letterhead_font_family,
                test_name_font_size=test_name_font_size,
                result_value_font_size=result_value_font_size,
                logo_position=logo_position,
                logo_width=logo_width,
                auto_update_enabled=auto_update_enabled,
                license_banner=license_banner,
                # أيقونة المصمم العائمة: تظهر فقط لمن سجّل دخوله فعلاً من
                # /designer/login بنفس المتصفح (session['designer_id']).
                # المستخدمين العاديين (admin, reception...) ما عندهم هذا
                # المفتاح بالسيشن أبداً، فالأيقونة ما تظهر عندهم إطلاقاً.
                is_designer=bool(session.get("designer_id")))


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return view(*args, **kwargs)
    return wrapped


def roles_required(*roles):
    def decorator(view):
        @wraps(view)
        def wrapped(*args, **kwargs):
            if "user_id" not in session:
                return redirect(url_for("login"))
            if session.get("role") not in roles and session.get("role") != "admin":
                flash("You do not have permission to access this page.")
                return redirect(url_for("dashboard"))
            return view(*args, **kwargs)
        return wrapped
    return decorator


def designer_required(view):
    """يحمي لوحة المصمم فقط — منفصلة تماماً عن نظام المستخدمين/الأدوار
    العادي في البرنامج (users/session['role']). لا أحد غير من يعرف اسم
    المستخدم وكلمة المرور الخاصين بالمصمم (المُنشَأين مرة واحدة عبر
    /designer/setup) يقدر يدخل هذه اللوحة."""
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get("designer_id"):
            return redirect(url_for("designer_login"))
        return view(*args, **kwargs)
    return wrapped


@app.before_request
def enforce_license():
    # المسارات المستثناة من فحص الترخيص: الملفات الثابتة، ولوحة المصمم نفسها
    # (المصمم لازم يقدر يدخلها حتى لو الترخيص منتهي عشان يولّد كود جديد)،
    # وشاشة القفل/التفعيل نفسها (وإلا صار حلقة تحويل لا نهائية).
    endpoint = request.endpoint or ""

    # لو صار تحديث تلقائي بالخلفية وأعاد تشغيل البرنامج، هذا يعرض رسالة
    # نجاح لأول مستخدم يفتح أي صفحة بعدها (بدون ما يحتاج يدخل لوحة
    # المصمم إطلاقاً)، ثم يمسح العلامة فوراً حتى ما تتكرر بكل صفحة.
    if endpoint != "static":
        _db = get_db()
        _pending = get_setting(_db, "auto_update_pending_banner", "")
        if _pending:
            set_setting(_db, "auto_update_pending_banner", "")
            _db.commit()
            flash(f"✅ تم تحديث البرنامج تلقائياً إلى الإصدار {_pending} بنجاح.")
        _db.close()

    if (endpoint == "static"
            or endpoint.startswith("designer_")
            or endpoint in ("license_locked", "license_activate", "set_lang")):
        return
    db = get_db()
    lic = license_manager.check_license(db)
    db.close()
    if lic["status"] in ("expired", "hardware_mismatch", "pending", "revoked"):
        session.pop("user_id", None)
        return redirect(url_for("license_locked"))


def log_action(action, entity, entity_id, details=""):
    db = get_db()
    db.execute(
        "INSERT INTO audit_logs (user_id, action, entity, entity_id, details, created_at) "
        "VALUES (?, ?, ?, ?, ?, ?)",
        (session.get("user_id"), action, entity, entity_id, details, datetime.now().isoformat(timespec="seconds")),
    )
    db.commit()


def next_registration_number(db):
    row = db.execute("SELECT MAX(registration_number) AS m FROM visits").fetchone()
    return (row["m"] or 18999) + 1


def doctor_pricing_context(db):
    """Everything the new-visit / edit-visit pages need to auto-fill the
    referring doctor field and switch prices live in the browser:
    - doctors: list of {id, full_name} for the autocomplete list
    - doctor_prices: {doctor_id: {test_id: price}} overrides per doctor
    - doctor_name_to_id: {lowercased full_name: id} so the page can match
      what reception typed to an existing doctor without another request
    """
    doctors = db.execute("SELECT id, full_name FROM doctors ORDER BY full_name").fetchall()
    doctors = [{"id": d["id"], "full_name": d["full_name"]} for d in doctors]

    overrides = db.execute("SELECT doctor_id, test_definition_id, price FROM doctor_test_prices").fetchall()
    doctor_prices = {}
    for row in overrides:
        doctor_prices.setdefault(str(row["doctor_id"]), {})[str(row["test_definition_id"])] = row["price"]

    doctor_name_to_id = {d["full_name"].strip().lower(): d["id"] for d in doctors}

    return doctors, doctor_prices, doctor_name_to_id


# --------------------------------------------------------------- auth ------
@app.route("/set-lang/<lang>")
def set_lang(lang):
    if lang in ("en", "ar"):
        session["lang"] = lang
    return redirect(request.referrer or url_for("dashboard"))


@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        db = get_db()
        user = db.execute("SELECT * FROM users WHERE username = ? AND is_active = 1", (username,)).fetchone()
        if user and user["password_hash"] == hash_password(password):
            session["user_id"] = user["id"]
            session["full_name"] = user["full_name"]
            session["role"] = user["role"]
            session["branch_id"] = user["branch_id"]
            session.permanent = bool(request.form.get("remember"))
            log_action("Login", "user", user["id"])
            return redirect(url_for("dashboard"))
        error = t(current_lang(), "invalid_login")
    return render_template("login.html", error=error)


@app.route("/logout")
def logout():
    if "user_id" in session:
        log_action("Logout", "user", session["user_id"])
    session.clear()
    return redirect(url_for("login"))


# ---------------------------------------------------- license (client side) --
@app.route("/license/locked")
def license_locked():
    db = get_db()
    lic = license_manager.check_license(db)
    db.close()
    return render_template("license_locked.html", lic=lic)


@app.route("/license/activate", methods=["POST"])
def license_activate():
    username = request.form.get("username", "")
    code = request.form.get("code", "")
    db = get_db()
    ok, err = license_manager.apply_activation(db, username, code)
    db.close()
    if ok:
        flash("تم تفعيل الترخيص بنجاح، يمكنك تسجيل الدخول الآن")
        return redirect(url_for("login"))
    flash(err or "فشل التفعيل")
    return redirect(url_for("license_locked"))


# --------------------------------------------------------- designer panel --
@app.route("/designer/setup", methods=["GET", "POST"])
def designer_setup():
    db = get_db()
    if license_manager.designer_exists(db):
        db.close()
        return redirect(url_for("designer_login"))
    error = None
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        confirm = request.form.get("confirm", "")
        if not username or len(password) < 6:
            error = "أدخل اسم مستخدم وكلمة مرور لا تقل عن 6 أحرف"
        elif password != confirm:
            error = "كلمتا المرور غير متطابقتين"
        else:
            license_manager.create_designer_account(db, username, password)
            db.close()
            flash("تم إنشاء حساب المصمم — سجّل الدخول الآن")
            return redirect(url_for("designer_login"))
    db.close()
    return render_template("designer/setup.html", error=error)


@app.route("/designer/login", methods=["GET", "POST"])
def designer_login():
    db = get_db()
    if not license_manager.designer_exists(db):
        db.close()
        return redirect(url_for("designer_setup"))
    error = None
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        if license_manager.verify_designer(db, username, password):
            session["designer_id"] = 1
            session["designer_username"] = username
            db.close()
            return redirect(url_for("designer_panel"))
        error = "بيانات الدخول غير صحيحة"
    db.close()
    return render_template("designer/login.html", error=error)


@app.route("/designer/logout")
def designer_logout():
    session.pop("designer_id", None)
    session.pop("designer_username", None)
    return redirect(url_for("designer_login"))


@app.route("/designer")
@designer_required
def designer_panel():
    db = get_db()
    lic = license_manager.check_license(db)
    issue_log = db.execute("SELECT * FROM license_issue_log ORDER BY id DESC LIMIT 50").fetchall()
    update_info = {
        "local_version": auto_updater.get_local_version(),
        "configured": auto_updater.is_configured(),
        "enabled": get_setting(db, "auto_update_enabled", "1") == "1",
        "last_check": get_setting(db, "auto_update_last_check", ""),
        "last_status": get_setting(db, "auto_update_last_status", ""),
        "branch": auto_updater.get_update_branch(db),
        "default_branch": auto_updater.GITHUB_BRANCH,
        "custom_channel": get_setting(db, "update_channel", ""),
    }
    github_write_token = get_setting(db, "github_write_token", "")
    github_read_token = get_setting(db, "github_read_token", "")
    revoked_licenses = {}
    revocation_error = None
    if github_write_token and auto_updater.is_configured():
        try:
            revoked_licenses = auto_updater.fetch_revocation_list()
        except Exception as e:
            revocation_error = str(e)
    branding = {
        "app_name": get_setting(db, "app_name", ""),
        "app_name_ar": get_setting(db, "app_name_ar", ""),
        "logo_path": get_setting(db, "logo_path", ""),
    }
    db.close()
    return render_template(
        "designer/panel.html", lic=lic, issue_log=issue_log,
        this_hw=license_manager.get_hardware_id(), this_ip=license_manager.get_local_ip(),
        update_info=update_info, github_write_token=github_write_token,
        github_read_token=github_read_token,
        revoked_licenses=revoked_licenses, revocation_error=revocation_error,
        branding=branding,
    )


@app.route("/designer/github-token", methods=["POST"])
@designer_required
def designer_save_github_token():
    """يحفظ توكن GitHub بصلاحية كتابة محلياً بجدول settings بجهاز المصمم
    فقط — منفصل تماماً عن توكن القراءة (settings.github_read_token، راجع
    /designer/github-read-token تحت) المشحون مع نسخة كل عميل. هذا التوكن
    لا يُشحن أبداً مع أي نسخة تُسلَّم لعميل، فيبقى فقط بقاعدة بيانات جهاز
    المصمم نفسه."""
    db = get_db()
    token = request.form.get("github_write_token", "").strip()
    set_setting(db, "github_write_token", token)
    db.commit()
    db.close()
    flash("تم حفظ توكن الكتابة." if token else "تم مسح توكن الكتابة.")
    return redirect(url_for("designer_panel"))


@app.route("/designer/github-read-token", methods=["POST"])
@designer_required
def designer_save_github_read_token():
    """يحفظ توكن القراءة (اللي يستخدمه هذا الجهاز بالذات لفحص/تنزيل
    التحديثات من GitHub) بجدول settings المحلي فقط — أبداً لا يُكتب بأي
    ملف كود يدخل بـ git push، حتى لا يكتشفه GitHub تلقائياً كسر مسرّب
    ويُلغيه (هذا بالضبط سبب انقطاع التحديث المتكرر قبل هذا التعديل).
    يُطبَّق فوراً بذاكرة هذا التشغيل عبر configure_token() بدون أي حاجة
    لإعادة تشغيل البرنامج."""
    db = get_db()
    token = request.form.get("github_read_token", "").strip()
    set_setting(db, "github_read_token", token)
    db.commit()
    db.close()
    auto_updater.configure_token(token)
    flash("تم حفظ توكن القراءة وتفعيله فوراً." if token else "تم مسح توكن القراءة.")
    return redirect(url_for("designer_panel"))


# تعديل سريع لاسم المختبر (عربي/إنكليزي) والشعار مباشرة من لوحة المصمم —
# بعد رفع تحديث لعميل معيّن، يضبط المصمم هويته الصحيحة بنفس الصفحة اللي
# رفع منها التحديث، بدون ما يحتاج يفتح "الإدارة ← الإعدادات" كخطوة منفصلة.
# يستخدم بالضبط نفس مفاتيح settings (app_name, app_name_ar, logo_path)
# ونفس منطق حفظ الشعار المستخدم بصفحة الإعدادات العادية (app_settings)،
# فالقيمتين مصدرهما واحد بغض النظر من وين تُعدَّل.
@app.route("/designer/branding", methods=["POST"])
@designer_required
def designer_save_branding():
    db = get_db()
    name_en = request.form.get("app_name", "").strip()
    name_ar = request.form.get("app_name_ar", "").strip()
    if name_en:
        set_setting(db, "app_name", name_en)
    if name_ar:
        set_setting(db, "app_name_ar", name_ar)

    file = request.files.get("logo")
    if file and file.filename:
        ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        if ext in ALLOWED_LOGO_EXT:
            filename = secure_filename(f"logo.{ext}")
            for old_ext in ALLOWED_LOGO_EXT:
                old_path = os.path.join(UPLOAD_DIR, f"logo.{old_ext}")
                if os.path.exists(old_path):
                    os.remove(old_path)
            file.save(os.path.join(UPLOAD_DIR, filename))
            set_setting(db, "logo_path", f"uploads/{filename}")
        else:
            flash("امتداد الشعار غير مدعوم. استخدم PNG, JPG, GIF, SVG أو WEBP.")

    db.commit()
    db.close()
    flash("تم حفظ اسم المختبر/الشعار.")
    return redirect(url_for("designer_panel"))


@app.route("/designer/revoke", methods=["POST"])
@designer_required
def designer_revoke_license():
    db = get_db()
    write_token = get_setting(db, "github_write_token", "")
    hardware_id = request.form.get("hardware_id", "").strip()
    reason = request.form.get("reason", "").strip()
    if not hardware_id:
        flash("أدخل معرّف جهاز العميل المطلوب إلغاء ترخيصه")
    else:
        try:
            auto_updater.push_revocation(write_token, hardware_id, reason, revoke=True)
            flash(f"تم إرسال إلغاء الترخيص لجهاز {hardware_id} — راح يوصله أول ما يتصل بالنت (لازم يكون \"مربوط بالإنترنت\" مفعّل عنده).")
        except Exception as e:
            flash(f"⚠️ تعذر إلغاء الترخيص عن بُعد: {e}")
    db.close()
    return redirect(url_for("designer_panel"))


@app.route("/designer/revoke/<path:hardware_id>/restore", methods=["POST"])
@designer_required
def designer_restore_license(hardware_id):
    db = get_db()
    write_token = get_setting(db, "github_write_token", "")
    try:
        auto_updater.push_revocation(write_token, hardware_id, None, revoke=False)
        flash(f"تم رفع الإلغاء عن جهاز {hardware_id} — يقدر يفعّل ببرنامجه بكود جديد أول ما يتصل بالنت.")
    except Exception as e:
        flash(f"⚠️ تعذر رفع الإلغاء: {e}")
    db.close()
    return redirect(url_for("designer_panel"))


@app.route("/designer/update/push", methods=["POST"])
@designer_required
def designer_push_update():
    """يرسل أمر تحديث فوري لجهاز عميل واحد بالذات (مستقل عن بقية العملاء)
    — لا ينتظر دورة الفحص التلقائي (كل CHECK_INTERVAL_HOURS)، بل يُطبَّق أول
    ما يتصل ذلك الجهاز بالنت (أو فوراً لو ضغط عنده زر 'تحقق من تحديث
    الآن'). يعتمد فعلياً على أن يكون ذلك الجهاز 'مربوط بالإنترنت' من لوحة
    المصمم عنده، ويستخدم نفس توكن الكتابة المستخدَم لإلغاء التراخيص عن بُعد."""
    db = get_db()
    write_token = get_setting(db, "github_write_token", "")
    hardware_id = request.form.get("hardware_id", "").strip()
    if not hardware_id:
        flash("أدخل معرّف جهاز العميل المطلوب إرسال التحديث له")
    else:
        try:
            auto_updater.push_update_signal(write_token, hardware_id)
            flash(f"تم إرسال أمر تحديث فوري لجهاز {hardware_id} — راح يُطبَّق أول ما يتصل بالنت "
                  f"(لازم يكون \"مربوط بالإنترنت\" مفعّل عنده، أو يضغط هو زر \"تحقق من تحديث الآن\").")
        except Exception as e:
            flash(f"⚠️ تعذر إرسال أمر التحديث: {e}")
    db.close()
    return redirect(url_for("designer_panel"))


@app.route("/designer/update/channel", methods=["POST"])
@designer_required
def designer_update_channel():
    """يضبط 'قناة التحديث' (اسم الفرع بمستودع GitHub) لهذا الجهاز بالذات.
    اتركه فاضي حتى يرجع يتابع الفرع العام (GITHUB_BRANCH، افتراضياً main)
    زي باقي العملاء. عبّي اسم فرع خاص (مثلاً client-alkut) حتى يستلم هذا
    الجهاز فقط تحديثات ذلك الفرع دون بقية العملاء."""
    db = get_db()
    channel = request.form.get("update_channel", "").strip()
    set_setting(db, "update_channel", channel)
    db.commit()
    db.close()
    if channel:
        flash(f"تم ضبط هذا الجهاز على قناة تحديث خاصة: {channel} — لن يتأثر بتحديثات main العامة إلا هذا الفرع.")
    else:
        flash("تم إرجاع هذا الجهاز لمتابعة قناة التحديث العامة (main) مثل بقية العملاء.")
    return redirect(url_for("designer_panel"))


@app.route("/designer/update/toggle", methods=["POST"])
@designer_required
def designer_update_toggle():
    """يشغّل/يوقف "ربط هذا الجهاز بالإنترنت" — أي تفعيل الفحص الدوري
    التلقائي للتحديثات من GitHub. هذا الأمر خاص بلوحة المصمم فقط ولا
    يظهر أبداً للمستخدمين العاديين (admin/reception)."""
    db = get_db()
    currently_on = get_setting(db, "auto_update_enabled", "0") == "1"
    set_setting(db, "auto_update_enabled", "0" if currently_on else "1")
    db.commit()
    db.close()
    flash("تم فصل البرنامج عن التحديث التلقائي" if currently_on else "تم ربط البرنامج بالإنترنت — سيتحقق تلقائياً من التحديثات دورياً")
    return redirect(url_for("designer_panel"))


@app.route("/designer/update/check-now", methods=["POST"])
@designer_required
def designer_update_check_now():
    """فحص فوري يدوي (زر احتياطي) — يشتغل بغض النظر عن حالة الربط
    التلقائي، مفيد لتجربة الاتصال أو لتنزيل تحديث فوراً بدون انتظار.
    يفحص أيضاً قائمة الإلغاء البعيدة لهذا الجهاز نفسه بنفس الفحص (مفيد
    لو المصمم يجرّب الميزة على جهازه هو قبل ما يعتمد عليها مع عميل)."""
    db = get_db()
    if not auto_updater.is_configured():
        flash("⚠️ لم يتم إعداد بيانات GitHub بعد داخل auto_updater.py")
    else:
        result = auto_updater.check_and_apply(db, force_apply=True)
        auto_updater.check_revocation(db)
        auto_updater.check_update_signal(db)
        flash(result["message"])
    db.close()
    return redirect(url_for("designer_panel"))


@app.route("/designer/generate", methods=["POST"])
@designer_required
def designer_generate():
    hardware_id = request.form.get("hardware_id", "").strip().upper()
    username = request.form.get("username", "").strip()
    expiry_mode = request.form.get("expiry_mode")

    if not hardware_id or not username:
        flash("أدخل معرّف الجهاز واسم المستخدم")
        return redirect(url_for("designer_panel"))

    preset = LICENSE_DURATION_PRESETS.get(expiry_mode)
    if not preset:
        flash("مدة ترخيص غير صالحة — اختر مدة من القائمة")
        return redirect(url_for("designer_panel"))

    _label, days = preset
    is_trial = expiry_mode == "trial"
    # days=None يعني ترخيص دائم فعلي (بدون تاريخ انتهاء) — license_manager
    # يخزّنه كـ "PERM" داخل الكود نفسه ويتحقق منه محلياً دون أي اعتماد على
    # تاريخ. الإلغاء عن بُعد (designer_revoke_license) يبقى شغّال بنفس
    # الطريقة حتى مع ترخيص دائم، فهو ليس "بلا رجعة".
    expiry_date = None if days is None else (date.today() + timedelta(days=days)).isoformat()

    code = license_manager.generate_activation_code(hardware_id, expiry_date, is_trial=is_trial)
    db = get_db()
    license_manager.log_issued_code(db, hardware_id, username, expiry_date, code)
    db.close()
    return render_template(
        "designer/generated.html", username=username, code=code,
        hardware_id=hardware_id, expiry_date=expiry_date,
    )


@app.route("/designer/issue-log/<int:log_id>/delete", methods=["POST"])
@designer_required
def designer_delete_issue_log(log_id):
    """حذف سطر واحد من سجل أكواد التفعيل المولَّدة — لا يلغي الترخيص نفسه
    عند العميل (الكود المفعّل عنده يبقى شغّال)، فقط ينظّف السجل المعروض
    بلوحة المصمم."""
    db = get_db()
    db.execute("DELETE FROM license_issue_log WHERE id=?", (log_id,))
    db.commit()
    db.close()
    flash("تم حذف السطر من السجل.")
    return redirect(url_for("designer_panel"))


@app.route("/designer/change-password", methods=["POST"])
@designer_required
def designer_change_password():
    new_password = request.form.get("new_password", "")
    confirm = request.form.get("confirm", "")
    if len(new_password) < 6 or new_password != confirm:
        flash("تحقق من كلمة المرور الجديدة (6 أحرف على الأقل ومتطابقة)")
        return redirect(url_for("designer_panel"))
    db = get_db()
    license_manager.change_designer_password(db, new_password)
    db.close()
    flash("تم تغيير كلمة مرور المصمم")
    return redirect(url_for("designer_panel"))


# ---------------------------------------------------------- dashboard ------
@app.route("/")
@login_required
def dashboard():
    db = get_db()
    today = date.today().isoformat()
    visits_today = db.execute(
        "SELECT COUNT(*) c FROM visits WHERE substr(created_at,1,10)=?", (today,)
    ).fetchone()["c"]
    revenue_today = db.execute(
        "SELECT COALESCE(SUM(amount),0) s FROM payments WHERE substr(paid_at,1,10)=?", (today,)
    ).fetchone()["s"]
    pending_results = db.execute(
        "SELECT COUNT(*) c FROM order_tests WHERE status IN ('Accepted','In-progress')"
    ).fetchone()["c"]
    critical_count = db.execute(
        "SELECT COUNT(*) c FROM results WHERE flag='Critical' AND verified_at IS NULL"
    ).fetchone()["c"]
    recent_visits = db.execute(
        "SELECT v.registration_number, p.full_name, v.status, v.created_at "
        "FROM visits v JOIN patients p ON p.id=v.patient_id "
        "ORDER BY v.id DESC LIMIT 8"
    ).fetchall()
    return render_template("dashboard.html", visits_today=visits_today, revenue_today=revenue_today,
                            pending_results=pending_results, critical_count=critical_count,
                            recent_visits=recent_visits)


# --------------------------------------------------------------- front desk
@app.route("/front-desk/new-visit", methods=["GET", "POST"])
@login_required
def new_visit():
    db = get_db()
    if request.method == "POST":
        # زر "✅ تم سحب العينة" إلزامي بهذي الصفحة (نقطة #9) — يمثّل تأكيد
        # الاستقبال إنه سحب العينة فعليًا وقت تسجيل الزيارة (حالة شائعة
        # بالمختبرات الصغيرة حيث الاستقبال هو نفسه يسحب العينة). لو ما
        # انضغط، نرفض الحفظ من السيرفر (مو بس تحقق JS بالواجهة) قبل أي
        # INSERT بقاعدة البيانات. تأثيره: order_tests تُنشأ مباشرة بحالة
        # 'Collected' (بدل 'Accepted' الافتراضية) فتتخطى طابور "سحب
        # العينة" وتظهر مباشرة بطابور "استلام العينة" — راجع
        # samples_collection/samples_accession لنفس منطق الحالتين.
        sample_collected = request.form.get("sample_collected") == "1"
        if not sample_collected:
            error_msg = "لازم تأكيد \"تم سحب العينة\" قبل حفظ الزيارة."
            if request.headers.get("X-LIS-Ajax") == "1":
                return jsonify({"ok": False, "error": error_msg}), 400
            flash(error_msg)
            return redirect(url_for("new_visit"))

        name = request.form.get("patient_name", "").strip()
        title = request.form.get("title", "Mr.")
        gender = request.form.get("gender", "")
        age = request.form.get("age") or None
        age_unit = request.form.get("age_unit") or "Years"
        phone = request.form.get("phone", "")
        email = request.form.get("email", "").strip()
        address = request.form.get("address", "")
        national_id = request.form.get("national_id", "").strip()
        passport_number = request.form.get("passport_number", "").strip()
        travel_certificate_number = request.form.get("travel_certificate_number", "").strip()
        lab_card_number = request.form.get("lab_card_number", "").strip()
        fasting = request.form.get("fasting", "Undefined")
        notes = request.form.get("notes", "")
        # المعلومات الصحية (Health Information) — خاصة بهذي الزيارة تحديداً.
        weight = request.form.get("weight") or None
        height = request.form.get("height") or None
        symptoms = request.form.get("symptoms", "")
        disease = request.form.get("disease", "")
        therapy = request.form.get("therapy", "")
        # الزيارة المنزلية (Home Visit)
        is_home_visit = 1 if request.form.get("is_home_visit") == "1" else 0
        home_visit_address = request.form.get("home_visit_address", "") if is_home_visit else ""
        try:
            home_visit_fee = float(request.form.get("home_visit_fee") or 0) if is_home_visit else 0.0
        except ValueError:
            home_visit_fee = 0.0
        examining_doctor = request.form.get("examining_doctor", "")
        expenses = request.form.get("expenses") or 0
        # حقل "الدكتور الفاحص" (attending_doctor) صار نفس "دكتور المختبر
        # الفاحص" (examining_doctor) — كانا حقلين منفصلين بالشاشة سابقًا
        # وصار واحد بس بناءً على طلب المستخدم، وهذا العمود يُملأ تلقائيًا
        # بنفس القيمة حتى يبقى عمود "الزيارات" شغّالاً بدون أي تغيير.
        attending_doctor = examining_doctor
        contact_method = request.form.get("contact_method", "None")
        test_ids = request.form.getlist("tests")
        now = datetime.now().isoformat(timespec="seconds")

        # الطبيب المُحيل (خارجي) — إذا هذا أول مرة يذكر اسمه، ينحفظ تلقائيًا
        # بجدول الأطباء حتى يظهر بالاقتراحات بالمرات الجاية.
        referring_doctor_id = find_or_create_doctor(db, request.form.get("doctor_name", ""))

        # المختبر المُرسِل لهذه العينة (لو النموذج وارد من مختبر ثاني مو من
        # مراجع مباشر) — تُستخدم فقط للمحاسبة الداخلية بين المختبرين؛ اسم
        # هذا المختبر ما ينطبع بأي تقرير أبدًا (راجع print_report/from_other_lab).
        # يُكتب كاسم حر (نفس أسلوب حقل الدكتور المُحيل) بدل قائمة منسدلة
        # ثابتة، فيُنشأ المختبر تلقائيًا بأول مرة يُذكر اسمه — ما تحتاج
        # الاستقبال تضيفه مسبقًا من صفحة إدارة المختبرات المُحيلة.
        referral_center_id = find_or_create_referral_center(db, request.form.get("referral_lab_name", ""))

        # إذا اختار موظف الاستقبال مريضًا سبق أن زار المختبر (من نتائج البحث
        # الفوري)، تُستخدم بطاقته الحالية بدل إنشاء بطاقة مريض مكررة جديدة.
        existing_patient_id = request.form.get("existing_patient_id") or None
        patient_id = None
        if existing_patient_id:
            existing = db.execute(
                "SELECT id FROM patients WHERE id=?", (existing_patient_id,)
            ).fetchone()
            if existing:
                patient_id = existing["id"]

        if patient_id is None:
            cur = db.execute(
                "INSERT INTO patients (full_name, gender, age, age_unit, phone, address, contact_method, "
                "title, email, national_id, passport_number, travel_certificate_number, lab_card_number, "
                "branch_id, created_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (name, gender, age, age_unit, phone, address, contact_method,
                 title, email, national_id, passport_number, travel_certificate_number, lab_card_number,
                 session.get("branch_id"), now),
            )
            patient_id = cur.lastrowid

        examining_doctor_fee = compute_examining_doctor_fee(db, examining_doctor, test_ids)

        reg_number = next_registration_number(db)
        visit_type = "home-visit" if is_home_visit else "walk-in"
        cur = db.execute(
            "INSERT INTO visits (registration_number, patient_id, doctor_id, referral_center_id, visit_type, fasting, notes, "
            "examining_doctor, expenses, examining_doctor_fee, attending_doctor, weight, height, symptoms, disease, therapy, "
            "is_home_visit, home_visit_address, home_visit_fee, status, branch_id, created_by, created_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'Open', ?, ?, ?)",
            (reg_number, patient_id, referring_doctor_id, referral_center_id, visit_type, fasting, notes, examining_doctor, expenses,
             examining_doctor_fee, attending_doctor, weight, height, symptoms, disease, therapy,
             is_home_visit, home_visit_address, home_visit_fee, session.get("branch_id"), session["user_id"], now),
        )
        visit_id = cur.lastrowid

        order_cur = db.execute("INSERT INTO orders (visit_id, status, created_at) VALUES (?, 'Open', ?)",
                                (visit_id, now))
        order_id = order_cur.lastrowid

        total = 0.0
        for tid in test_ids:
            test = db.execute("SELECT * FROM test_definitions WHERE id=?", (tid,)).fetchone()
            if not test:
                continue
            price = get_test_price(db, tid, referring_doctor_id)
            barcode = f"{reg_number}{tid.zfill(3)}"
            # الحالة تبدأ 'Collected' مباشرة (مو 'Accepted') لأن زر "تم سحب
            # العينة" الإلزامي فوق تأكد إنها انسحبت فعليًا هذي اللحظة —
            # فتتخطى طابور "سحب العينة" وتظهر مباشرة بطابور "استلام العينة".
            db.execute(
                "INSERT INTO order_tests (order_id, test_definition_id, status, barcode, price, doctor_id, "
                "collected_at, created_at) VALUES (?, ?, 'Collected', ?, ?, ?, ?, ?)",
                (order_id, tid, barcode, price, referring_doctor_id, now, now),
            )
            total += price or 0

        # أجرة الزيارة المنزلية الإضافية (إن فُعِّلت) تُضاف لمجموع الفحوصات
        # قبل حساب الفاتورة، بنفس أسلوب أي بند إضافي بالمجموع.
        total += home_visit_fee

        # "المبلغ الكلي" بصفحة الزيارة الجديدة يبدأ محسوبًا تلقائيًا من
        # مجموع أسعار التحاليل المختارة، بس موظف الاستقبال يقدر يعدّله يدويًا
        # (مثلاً لخصم أو تسوية) — الفرق بين المجموع الفعلي والمبلغ المُعدَّل
        # يُسجَّل بعمودي discount_amount/extra_charges حتى يبقى مجموع أسعار
        # التحاليل الأصلي محفوظًا للمراجعة.
        computed_total = total
        total_override_raw = request.form.get("total_amount_input")
        try:
            total_override = float(total_override_raw) if total_override_raw not in (None, "") else None
        except ValueError:
            total_override = None
        invoice_total = total_override if (total_override is not None and total_override >= 0) else computed_total
        discount_amount = max(0.0, computed_total - invoice_total)
        extra_charges = max(0.0, invoice_total - computed_total)

        # "الواصل" — أي مبلغ استلمه الاستقبال نقدًا وقت تسجيل الزيارة نفسها.
        # يُسجَّل كدفعة فعلية بجدول payments (نفس آلية /billing/pay) حتى يظهر
        # بسجلات المحاسبة والتقارير، و"الباقي" يُحسب تلقائيًا من الفرق.
        try:
            paid_amount = float(request.form.get("paid_amount") or 0)
        except ValueError:
            paid_amount = 0.0
        paid_amount = max(0.0, paid_amount)
        if invoice_total > 0:
            paid_amount = min(paid_amount, invoice_total)
            invoice_status = "Paid" if paid_amount >= invoice_total else ("Partial" if paid_amount > 0 else "Unpaid")
        else:
            invoice_status = "Paid" if paid_amount <= 0 else "Partial"

        inv_cur = db.execute(
            "INSERT INTO invoices (visit_id, total_amount, discount_amount, extra_charges, paid_amount, status, "
            "created_by, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (visit_id, invoice_total, discount_amount, extra_charges, paid_amount, invoice_status,
             session["user_id"], now),
        )
        invoice_id = inv_cur.lastrowid

        if paid_amount > 0:
            db.execute(
                "INSERT INTO payments (invoice_id, amount, method, user_id, paid_at) VALUES (?, ?, 'Cash', ?, ?)",
                (invoice_id, paid_amount, session["user_id"], now),
            )
            log_action("Payment", "invoice", invoice_id, f"amount={paid_amount} (at registration)")


        if contact_method != "None":
            db.execute(
                "INSERT INTO patient_followups (patient_id, visit_id, status, followup_date, created_at) "
                "VALUES (?, ?, 'Pending', ?, ?)",
                (patient_id, visit_id, now[:10], now),
            )

        db.commit()
        log_action("Create", "visit", visit_id, f"reg#{reg_number}")
        log_action("Collect", "visit", visit_id, f"sample_collected_at_registration reg#{reg_number}")
        flash(f"Visit #{reg_number} created successfully.")

        # إذا اختار موظف الاستقبال تضمين زيارة/زيارات سابقة مع هذي الزيارة
        # الجديدة بجدول موحّد، يروح مباشرة لصفحة الجدول الموحّد بدل قائمة الزيارات.
        include_visit_ids = request.form.get("include_visit_ids", "").strip()
        if include_visit_ids:
            all_ids = [visit_id_str for visit_id_str in include_visit_ids.split(",") if visit_id_str.strip().isdigit()]
            all_ids.append(str(visit_id))
            redirect_url = url_for("combined_visits_report", visit_ids=",".join(all_ids))
        else:
            redirect_url = url_for("visits_list")

        # صفحة "زيارة جديدة" تُرسل الفورم بـ AJAX حتى تقدر تعرض زر "طباعة
        # الباركود" وتطبع من نفس الصفحة (بدون فتح تبويب/متصفح جديد) قبل ما
        # تنتقل لقائمة الزيارات. لو الطلب اجى بالطريقة العادية (بدون JS)،
        # نكمل بنفس سلوك التحويل المباشر كما كان سابقًا.
        if request.headers.get("X-LIS-Ajax") == "1":
            return jsonify({
                "ok": True,
                "visit_id": visit_id,
                "registration_number": reg_number,
                "patient_name": name,
                "redirect_url": redirect_url,
            })

        return redirect(redirect_url)

    tests = db.execute("SELECT * FROM test_definitions WHERE is_active=1 ORDER BY department, name").fetchall()
    quick_items = db.execute(
        "SELECT td.id, td.name FROM quick_add_items q JOIN test_definitions td ON td.id = q.test_definition_id "
        "WHERE q.is_active=1 ORDER BY q.display_order"
    ).fetchall()
    doctors, doctor_prices, doctor_name_to_id = doctor_pricing_context(db)
    test_default_prices = {str(tst["id"]): tst["price"] for tst in tests}
    examining_test_ids = [row["id"] for row in get_examining_tests(db)]
    referral_labs = db.execute(
        "SELECT * FROM referral_centers WHERE name != 'Walk-in' ORDER BY name"
    ).fetchall()
    return render_template("front_desk/new_visit.html", tests=tests, quick_items=quick_items,
                            examining_doctors=get_examining_doctors(db), doctors=doctors,
                            doctor_prices=doctor_prices, doctor_name_to_id=doctor_name_to_id,
                            test_default_prices=test_default_prices,
                            examining_test_ids=examining_test_ids,
                            examining_rates=get_examining_rates_map(db),
                            referral_labs=referral_labs)


@app.route("/front-desk/patients")
@login_required
def patients_list():
    db = get_db()
    q = request.args.get("q", "").strip()
    query = "SELECT * FROM patients "
    params = []
    if q:
        query += "WHERE full_name LIKE ? OR phone LIKE ? OR national_id LIKE ? "
        params = [f"%{q}%", f"%{q}%", f"%{q}%"]
    query += "ORDER BY id DESC LIMIT 200"
    patients = db.execute(query, params).fetchall()
    return render_template("front_desk/patients.html", patients=patients, q=q)


@app.route("/front-desk/patients/new", methods=["GET", "POST"])
@login_required
def patient_new():
    db = get_db()
    if request.method == "POST":
        now = datetime.now().isoformat(timespec="seconds")
        db.execute(
            "INSERT INTO patients (full_name, gender, age, age_unit, phone, email, address, national_id, "
            "passport_number, lab_card_number, contact_method, title, travel_certificate_number, "
            "branch_id, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (request.form.get("full_name", "").strip(), request.form.get("gender"),
             request.form.get("age") or None, request.form.get("age_unit") or "Years",
             request.form.get("phone"), request.form.get("email"),
             request.form.get("address"), request.form.get("national_id"),
             request.form.get("passport_number"), request.form.get("lab_card_number"),
             request.form.get("contact_method", "None"), request.form.get("title", "Mr."),
             request.form.get("travel_certificate_number"), session.get("branch_id"), now),
        )
        db.commit()
        flash("Patient added.")
        return redirect(url_for("patients_list"))
    return render_template("front_desk/patient_form.html", patient=None)


@app.route("/front-desk/patients/<int:patient_id>/edit", methods=["GET", "POST"])
@login_required
def patient_edit(patient_id):
    db = get_db()
    patient = db.execute("SELECT * FROM patients WHERE id=?", (patient_id,)).fetchone()
    if not patient:
        return "Not found", 404
    if request.method == "POST":
        db.execute(
            "UPDATE patients SET full_name=?, gender=?, age=?, age_unit=?, phone=?, email=?, address=?, "
            "national_id=?, passport_number=?, lab_card_number=?, contact_method=?, title=?, "
            "travel_certificate_number=? WHERE id=?",
            (request.form.get("full_name", "").strip(), request.form.get("gender"),
             request.form.get("age") or None, request.form.get("age_unit") or "Years",
             request.form.get("phone"), request.form.get("email"),
             request.form.get("address"), request.form.get("national_id"),
             request.form.get("passport_number"), request.form.get("lab_card_number"),
             request.form.get("contact_method", "None"), request.form.get("title", "Mr."),
             request.form.get("travel_certificate_number"), patient_id),
        )
        db.commit()
        log_action("Update", "patient", patient_id)
        flash("Patient updated.")
        return redirect(url_for("patients_list"))
    return render_template("front_desk/patient_form.html", patient=patient)


@app.route("/front-desk/patients/<int:patient_id>/delete", methods=["POST"])
@roles_required("admin")
def delete_patient(patient_id):
    db = get_db()
    patient = db.execute("SELECT * FROM patients WHERE id=?", (patient_id,)).fetchone()
    if not patient:
        return "Not found", 404
    # لا نحذف المريض إذا عنده زيارات مسجّلة (بيها نتائج/فواتير مرتبطة)، حتى
    # لا تبقى سجلات يتيمة. نمنع الحذف بهذي الحالة ونعرض عدد الزيارات.
    linked_visits = db.execute("SELECT COUNT(*) c FROM visits WHERE patient_id=?", (patient_id,)).fetchone()["c"]
    if linked_visits:
        flash(f"لا يمكن حذف هذا المريض لأنه مرتبط بـ {linked_visits} زيارة. "
              f"احذف تلك الزيارات أولاً إذا تريد حذفه فعلاً.")
        return redirect(url_for("patients_list"))
    db.execute("DELETE FROM patients WHERE id=?", (patient_id,))
    db.commit()
    log_action("Delete", "patient", patient_id, patient["full_name"])
    flash("تم حذف المريض.")
    return redirect(url_for("patients_list"))


# بحث سريع عن مريض سبق أن زار المختبر: يعرض كل زياراته السابقة والفحوصات
# التي أجراها، ويسمح للأدمن حصراً باختيار نتائج قديمة (تحليل واحد أو أكثر،
# من أي زيارة سابقة) ونسخها إلى أحدث زيارة للمريض — وليس الزيارة كاملة.
@app.route("/front-desk/patients/<int:patient_id>/history", methods=["GET", "POST"])
@login_required
def patient_history(patient_id):
    db = get_db()
    patient = db.execute("SELECT * FROM patients WHERE id=?", (patient_id,)).fetchone()
    if not patient:
        return "Not found", 404

    if request.method == "POST":
        if session.get("role") != "admin":
            flash("نسخ نتائج قديمة إلى الزيارة الحالية صلاحية خاصة بالأدمن فقط.")
            return redirect(url_for("patient_history", patient_id=patient_id))

        selected_result_ids = request.form.getlist("copy_result")
        if not selected_result_ids:
            flash("لم تحدد أي نتيجة لنسخها.")
            return redirect(url_for("patient_history", patient_id=patient_id))

        latest_visit = db.execute(
            "SELECT v.*, o.id as order_id FROM visits v JOIN orders o ON o.visit_id = v.id "
            "WHERE v.patient_id=? ORDER BY v.id DESC LIMIT 1",
            (patient_id,),
        ).fetchone()
        if not latest_visit:
            flash("لا توجد زيارة حالية لهذا المريض لنسخ النتائج إليها.")
            return redirect(url_for("patient_history", patient_id=patient_id))

        now = datetime.now().isoformat(timespec="seconds")
        copied = 0
        for result_id in selected_result_ids:
            old_result = db.execute(
                "SELECT r.*, ot.test_definition_id, ot.barcode FROM results r "
                "JOIN order_tests ot ON ot.id = r.order_test_id WHERE r.id=?",
                (result_id,),
            ).fetchone()
            if not old_result:
                continue
            reg = latest_visit["registration_number"]
            tid = old_result["test_definition_id"]
            new_barcode = f"{reg}{str(tid).zfill(3)}-old"
            cur = db.execute(
                "INSERT INTO order_tests (order_id, test_definition_id, status, barcode, notes, created_at) "
                "VALUES (?, ?, 'Completed', ?, ?, ?)",
                (latest_visit["order_id"], tid, new_barcode, "نتيجة منسوخة من زيارة سابقة", now),
            )
            new_order_test_id = cur.lastrowid
            db.execute(
                "INSERT INTO results (order_test_id, test_parameter_id, value_numeric, value_text, flag, "
                "entered_by, entered_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                (new_order_test_id, old_result["test_parameter_id"], old_result["value_numeric"],
                 old_result["value_text"], old_result["flag"], session["user_id"], now),
            )
            copied += 1
        db.commit()
        log_action("CopyOldResults", "visit", latest_visit["id"], f"{copied} نتيجة من زيارات سابقة")
        flash(f"تم نسخ {copied} نتيجة إلى الزيارة الحالية #{latest_visit['registration_number']}.")
        return redirect(url_for("visit_edit", visit_id=latest_visit["id"]))

    visits = db.execute(
        "SELECT v.* FROM visits v WHERE v.patient_id=? ORDER BY v.id DESC", (patient_id,),
    ).fetchall()
    visit_data = []
    for v in visits:
        order_tests = db.execute(
            "SELECT ot.id, ot.barcode, td.name as test_name, td.id as test_definition_id "
            "FROM order_tests ot JOIN orders o ON o.id = ot.order_id "
            "JOIN test_definitions td ON td.id = ot.test_definition_id WHERE o.visit_id=?",
            (v["id"],),
        ).fetchall()
        tests_with_results = []
        for ot in order_tests:
            results = db.execute(
                "SELECT r.*, tp.name as param_name, tp.unit FROM results r "
                "JOIN test_parameters tp ON tp.id = r.test_parameter_id WHERE r.order_test_id=?",
                (ot["id"],),
            ).fetchall()
            tests_with_results.append({"order_test": ot, "results": results})
        visit_data.append({"visit": v, "tests": tests_with_results})

    return render_template("front_desk/patient_history.html", patient=patient, visit_data=visit_data)


# بحث موحّد بالصفحة الرئيسية: يبحث بثلاث فئات مرة وحدة — اسم مريض (يظهر
# تفاصيله عبر رابط سجله)، اسم طبيب مُحيل، واسم مختبر مُرسِل (الأخيرين
# يوديان لقائمة زياراتهما مرتبة بالأحدث تاريخًا ووقتًا عبر فلتر
# doctor_id/referral_center_id بصفحة "الزيارات" — راجع visits_list أعلاه).
@app.route("/api/dashboard/search")
@login_required
def api_dashboard_search():
    q = (request.args.get("q") or "").strip()
    if len(q) < 2:
        return jsonify({"patients": [], "doctors": [], "referral_labs": []})
    db = get_db()
    patients = db.execute(
        "SELECT p.id, p.full_name, p.gender, p.age, p.age_unit, p.phone, "
        "(SELECT MAX(v.created_at) FROM visits v WHERE v.patient_id = p.id) as last_visit_at "
        "FROM patients p WHERE p.full_name LIKE ? ORDER BY p.full_name LIMIT 6",
        (f"%{q}%",),
    ).fetchall()
    doctors = db.execute(
        "SELECT d.id, d.full_name, d.phone, "
        "(SELECT MAX(v.created_at) FROM visits v WHERE v.doctor_id = d.id) as last_visit_at "
        "FROM doctors d WHERE d.full_name LIKE ? "
        "ORDER BY last_visit_at IS NULL, last_visit_at DESC LIMIT 6",
        (f"%{q}%",),
    ).fetchall()
    referral_labs = db.execute(
        "SELECT rc.id, rc.name, rc.phone, "
        "(SELECT MAX(v.created_at) FROM visits v WHERE v.referral_center_id = rc.id) as last_visit_at "
        "FROM referral_centers rc WHERE rc.name LIKE ? AND rc.name != 'Walk-in' "
        "ORDER BY last_visit_at IS NULL, last_visit_at DESC LIMIT 6",
        (f"%{q}%",),
    ).fetchall()
    return jsonify({
        "patients": [dict(r) for r in patients],
        "doctors": [dict(r) for r in doctors],
        "referral_labs": [dict(r) for r in referral_labs],
    })


# بحث فوري (Live search) يُستخدم من شاشة "زيارة جديدة": بمجرد كتابة اسم
# المريض يبحث عن أي مطابقة سابقة بجدول المرضى، حتى لا تنفتح بطاقة مريض
# مكررة لشخص سبق أن راجع المختبر.
@app.route("/api/patients/search")
@login_required
def api_patients_search():
    q = (request.args.get("q") or "").strip()
    if len(q) < 2:
        return jsonify([])
    db = get_db()
    rows = db.execute(
        "SELECT p.id, p.full_name, p.gender, p.age, p.age_unit, p.phone, p.address, "
        "(SELECT MAX(v.created_at) FROM visits v WHERE v.patient_id = p.id) as last_visit_at "
        "FROM patients p WHERE p.full_name LIKE ? ORDER BY p.full_name LIMIT 8",
        (f"%{q}%",),
    ).fetchall()
    return jsonify([dict(r) for r in rows])


# ملخص زيارات مريض سبق أن راجع (تُستدعى من شاشة "زيارة جديدة" بعد اختيار
# مطابقة من البحث الفوري) — تاريخ كل زيارة والتحاليل التي أُجريت بها، حتى
# تظهر أزرار مشاهدة / طباعة / تضمين لكل زيارة قديمة.
@app.route("/api/patients/<int:patient_id>/visits-summary")
@login_required
def api_patient_visits_summary(patient_id):
    db = get_db()
    patient = db.execute("SELECT * FROM patients WHERE id=?", (patient_id,)).fetchone()
    if not patient:
        return jsonify({"error": "not found"}), 404
    visits = db.execute(
        "SELECT v.id, v.registration_number, v.created_at FROM visits v "
        "WHERE v.patient_id=? ORDER BY v.id DESC LIMIT 10",
        (patient_id,),
    ).fetchall()
    visits_out = []
    for v in visits:
        tests = db.execute(
            "SELECT td.name FROM order_tests ot JOIN orders o ON o.id = ot.order_id "
            "JOIN test_definitions td ON td.id = ot.test_definition_id WHERE o.visit_id=? ORDER BY td.name",
            (v["id"],),
        ).fetchall()
        has_results = db.execute(
            "SELECT COUNT(*) c FROM results r JOIN order_tests ot ON ot.id = r.order_test_id "
            "JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=?",
            (v["id"],),
        ).fetchone()["c"] > 0
        visits_out.append({
            "id": v["id"],
            "registration_number": v["registration_number"],
            "created_at": v["created_at"],
            "tests": [t["name"] for t in tests],
            "has_results": has_results,
        })
    return jsonify({
        "patient": {
            "id": patient["id"], "full_name": patient["full_name"], "gender": patient["gender"],
            "age": patient["age"], "age_unit": patient["age_unit"], "phone": patient["phone"],
            "address": patient["address"],
        },
        "visits": visits_out,
    })


@app.route("/front-desk/results")
@login_required
def results_list():
    db = get_db()
    # فلترة اختيارية قادمة من بطاقات الرئيسية:
    # - "pending": زيارات فيها تحليل واحد على الأقل لسا حالته Accepted أو
    #   In-progress — نفس المعيار بالضبط المستخدم بعدّاد "نتائج قيد الإنجاز"
    #   بالرئيسية (dashboard()), حتى يطابق العدد المعروض هناك عدد الصفوف هنا.
    # - "critical": زيارات فيها نتيجة واحدة على الأقل flag='Critical' ولسا
    #   verified_at فاضي — نفس معيار عدّاد "تنبيهات حرجة" بالرئيسية بالضبط.
    filter_type = request.args.get("filter", "")
    query = (
        "SELECT v.id as visit_id, v.registration_number, v.created_at, p.full_name as patient_name, "
        "COUNT(ot.id) as tests_count, "
        "SUM(CASE WHEN ot.status IN ('Completed', 'Verified') THEN 1 ELSE 0 END) as done_count, "
        "SUM(CASE WHEN ot.status = 'Verified' THEN 1 ELSE 0 END) as verified_count, "
        "GROUP_CONCAT(ot.id || ':' || td.name, '||') as tests_list "
        "FROM order_tests ot "
        "JOIN orders o ON o.id = ot.order_id "
        "JOIN visits v ON v.id = o.visit_id "
        "JOIN patients p ON p.id = v.patient_id "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
    )
    filter_label = None
    if filter_type == "pending":
        query += (
            "WHERE v.id IN ("
            "  SELECT o2.visit_id FROM order_tests ot2 "
            "  JOIN orders o2 ON o2.id = ot2.order_id "
            "  WHERE ot2.status IN ('Accepted','In-progress')"
            ") "
        )
        filter_label = "زيارات فيها نتائج قيد الإنجاز"
    elif filter_type == "critical":
        query += (
            "WHERE v.id IN ("
            "  SELECT o2.visit_id FROM order_tests ot2 "
            "  JOIN orders o2 ON o2.id = ot2.order_id "
            "  JOIN results r2 ON r2.order_test_id = ot2.id "
            "  WHERE r2.flag='Critical' AND r2.verified_at IS NULL"
            ") "
        )
        filter_label = "زيارات فيها تنبيهات حرجة غير مُصادَق عليها"
    query += "GROUP BY v.id ORDER BY v.id DESC LIMIT 200"
    rows = db.execute(query).fetchall()
    return render_template("front_desk/results.html", rows=rows, filter_type=filter_type,
                            filter_label=filter_label)


@app.route("/front-desk/followups")
@login_required
def followups_list():
    db = get_db()
    status_filter = request.args.get("status", "")
    query = (
        "SELECT f.id, f.status, f.notes, f.followup_date, p.full_name, p.phone, "
        "v.registration_number "
        "FROM patient_followups f JOIN patients p ON p.id = f.patient_id "
        "LEFT JOIN visits v ON v.id = f.visit_id "
    )
    params = []
    if status_filter:
        query += "WHERE f.status = ? "
        params.append(status_filter)
    query += "ORDER BY f.id DESC LIMIT 200"
    rows = db.execute(query, params).fetchall()
    return render_template("front_desk/followups.html", rows=rows, status_filter=status_filter)


@app.route("/front-desk/followups/<int:followup_id>/update", methods=["POST"])
@login_required
def followup_update(followup_id):
    db = get_db()
    status = request.form.get("status", "Pending")
    notes = request.form.get("notes", "")
    db.execute("UPDATE patient_followups SET status=?, notes=? WHERE id=?", (status, notes, followup_id))
    db.commit()
    log_action("UpdateFollowup", "patient_followup", followup_id, status)
    flash("Followup updated.")
    return redirect(url_for("followups_list"))



@app.route("/management/doctors", methods=["GET", "POST"])
@roles_required("admin")
def doctors_list():
    db = get_db()
    if request.method == "POST":
        db.execute(
            "INSERT INTO doctors (full_name, specialty, phone, email, commission_percent) "
            "VALUES (?, ?, ?, ?, ?)",
            (request.form.get("full_name", "").strip(), request.form.get("specialty"),
             request.form.get("phone"), request.form.get("email"),
             float(request.form.get("commission_percent") or 0)),
        )
        db.commit()
        flash("Doctor added.")
        return redirect(url_for("doctors_list"))
    doctors = db.execute("SELECT * FROM doctors ORDER BY full_name").fetchall()
    return render_template("management/doctors.html", doctors=doctors)


@app.route("/management/doctors/<int:doctor_id>/delete", methods=["POST"])
@roles_required("admin")
def delete_doctor(doctor_id):
    db = get_db()
    doctor = db.execute("SELECT * FROM doctors WHERE id=?", (doctor_id,)).fetchone()
    if not doctor:
        return "Not found", 404
    # لا نحذف الطبيب إذا مرتبط بزيارات أو طلبات موجودة فعلاً، حتى لا تبقى
    # سجلات يتيمة (orphaned) بالزيارات القديمة. بهذي الحالة نمنع الحذف
    # ونعرض للمستخدم عدد السجلات المرتبطة.
    linked_visits = db.execute("SELECT COUNT(*) c FROM visits WHERE doctor_id=?", (doctor_id,)).fetchone()["c"]
    linked_orders = db.execute("SELECT COUNT(*) c FROM order_tests WHERE doctor_id=?", (doctor_id,)).fetchone()["c"]
    if linked_visits or linked_orders:
        flash(f"لا يمكن حذف هذا الطبيب لأنه مرتبط بـ {linked_visits} زيارة و {linked_orders} طلب تحليل. "
              f"احذف/عدّل تلك السجلات أولاً إذا تريد حذفه فعلاً.")
        return redirect(url_for("doctors_list"))
    db.execute("DELETE FROM doctor_test_prices WHERE doctor_id=?", (doctor_id,))
    db.execute("DELETE FROM doctors WHERE id=?", (doctor_id,))
    db.commit()
    log_action("Delete", "doctor", doctor_id, doctor["full_name"])
    flash("تم حذف الطبيب.")
    return redirect(url_for("doctors_list"))


@app.route("/management/doctors/<int:doctor_id>/rates", methods=["GET", "POST"])
@roles_required("admin", "accountant")
def doctor_rates(doctor_id):
    db = get_db()
    doctor = db.execute("SELECT * FROM doctors WHERE id=?", (doctor_id,)).fetchone()
    if not doctor:
        return "Not found", 404
    if request.method == "POST":
        test_id = request.form.get("test_id")
        price = request.form.get("price", "").strip()
        if price == "":
            # حقل فارغ = رجّع هذا التحليل لسعر المختبر الافتراضي لهذا الطبيب
            db.execute("DELETE FROM doctor_test_prices WHERE doctor_id=? AND test_definition_id=?",
                       (doctor_id, test_id))
            flash("Price reset to the default lab rate for this test.")
        else:
            db.execute(
                "INSERT INTO doctor_test_prices (doctor_id, test_definition_id, price) VALUES (?, ?, ?) "
                "ON CONFLICT(doctor_id, test_definition_id) DO UPDATE SET price=excluded.price",
                (doctor_id, test_id, float(price)),
            )
            flash("Doctor's price updated.")
        db.commit()
        return redirect(url_for("doctor_rates", doctor_id=doctor_id))

    tests = db.execute("SELECT * FROM test_definitions WHERE is_active=1 ORDER BY department, name").fetchall()
    overrides = {
        row["test_definition_id"]: row["price"]
        for row in db.execute(
            "SELECT test_definition_id, price FROM doctor_test_prices WHERE doctor_id=?", (doctor_id,)
        ).fetchall()
    }
    return render_template("management/doctor_rates.html", doctor=doctor, tests=tests, overrides=overrides)


@app.route("/management/doctors/<int:doctor_id>/statement")
@roles_required("admin")
def doctor_statement(doctor_id):
    db = get_db()
    doctor = db.execute("SELECT * FROM doctors WHERE id=?", (doctor_id,)).fetchone()
    month = request.args.get("month", datetime.now().strftime("%Y-%m"))
    rows = db.execute(
        "SELECT ot.id, td.name as test_name, COALESCE(ot.price, td.price) as price, "
        "p.full_name as patient_name, ot.created_at "
        "FROM order_tests ot JOIN test_definitions td ON td.id=ot.test_definition_id "
        "JOIN orders o ON o.id=ot.order_id JOIN visits v ON v.id=o.visit_id "
        "JOIN patients p ON p.id=v.patient_id "
        "WHERE ot.doctor_id=? AND substr(ot.created_at,1,7)=? ORDER BY ot.created_at",
        (doctor_id, month),
    ).fetchall()
    total = sum((r["price"] or 0) for r in rows)
    commission_due = total * ((doctor["commission_percent"] or 0) / 100)
    return render_template("management/doctor_statement.html", doctor=doctor, rows=rows, month=month,
                            total=total, commission_due=commission_due)



@app.route("/management/examining-doctor-rates")
@roles_required("admin", "accountant")
def examining_doctor_rates_list():
    db = get_db()
    names = get_examining_doctors(db)
    rates_map = get_examining_rates_map(db)
    eligible_count = len(get_examining_tests(db))
    doctors_summary = [
        {"name": n, "configured_count": len(rates_map.get(n, {}))} for n in names
    ]
    return render_template("management/examining_doctor_rates_list.html",
                            doctors_summary=doctors_summary, eligible_count=eligible_count)


@app.route("/management/examining-doctor-rates/<doctor_name>", methods=["GET", "POST"])
@roles_required("admin", "accountant")
def examining_doctor_rates(doctor_name):
    db = get_db()
    if doctor_name not in get_examining_doctors(db):
        return "Not found", 404
    if request.method == "POST":
        test_id = request.form.get("test_id")
        rate = request.form.get("rate", "").strip()
        set_examining_doctor_rate(db, doctor_name, int(test_id), float(rate or 0))
        flash("تم تحديث أجر الدكتور الفاحص لهذا الفحص.")
        return redirect(url_for("examining_doctor_rates", doctor_name=doctor_name))

    tests = get_examining_tests(db)
    rates = get_examining_rates_map(db).get(doctor_name, {})
    return render_template("management/examining_doctor_rates.html", doctor_name=doctor_name,
                            tests=tests, rates=rates)


@app.route("/management/examining-doctor-rates/<doctor_name>/statement")
@roles_required("admin", "accountant")
def examining_doctor_statement(doctor_name):
    db = get_db()
    month = request.args.get("month", datetime.now().strftime("%Y-%m"))
    rows = db.execute(
        "SELECT v.registration_number, v.examining_doctor_fee, v.created_at, p.full_name "
        "FROM visits v JOIN patients p ON p.id = v.patient_id "
        "WHERE v.examining_doctor=? AND substr(v.created_at,1,7)=? AND v.examining_doctor_fee > 0 "
        "ORDER BY v.created_at",
        (doctor_name, month),
    ).fetchall()
    total = sum((r["examining_doctor_fee"] or 0) for r in rows)
    return render_template("management/examining_doctor_statement.html", doctor_name=doctor_name,
                            rows=rows, month=month, total=total)


@app.route("/management/referral-labs", methods=["GET", "POST"])
@roles_required("admin")
def referral_labs_list():
    db = get_db()
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        phone = request.form.get("phone", "").strip()
        if not name:
            flash("اسم المختبر إلزامي.")
            return redirect(url_for("referral_labs_list"))
        exists = db.execute(
            "SELECT id FROM referral_centers WHERE LOWER(TRIM(name))=LOWER(?)", (name,)
        ).fetchone()
        if exists:
            flash("هذا المختبر مضاف أصلاً.")
        else:
            db.execute("INSERT INTO referral_centers (name, type, phone) VALUES (?, 'Lab', ?)",
                       (name, phone or None))
            db.commit()
            flash("تمت إضافة المختبر.")
        return redirect(url_for("referral_labs_list"))

    labs = db.execute(
        "SELECT * FROM referral_centers WHERE name != 'Walk-in' ORDER BY name"
    ).fetchall()
    return render_template("management/referral_labs.html", labs=labs)


@app.route("/management/referral-labs/<int:center_id>/phone", methods=["POST"])
@roles_required("admin")
def referral_lab_update_phone(center_id):
    db = get_db()
    phone = request.form.get("phone", "").strip()
    db.execute("UPDATE referral_centers SET phone=? WHERE id=?", (phone or None, center_id))
    db.commit()
    flash("تم تحديث رقم واتساب المختبر.")
    return redirect(url_for("referral_labs_list"))


def _referral_lab_statement_data(db, center_id, month):
    """يرجّع (lab, rows, by_day, by_test, total) لكشف حساب مختبر مُرسِل معيّن
    عن شهر معيّن — كل تحليل بكل نموذج وارد منه بهذا الشهر، مجمّع يوميًا
    (لمتابعة الحساب أول بأول) ومجمّع حسب نوع التحليل (لملخص نهاية الشهر)."""
    lab = db.execute("SELECT * FROM referral_centers WHERE id=?", (center_id,)).fetchone()
    if not lab:
        return None, [], [], [], 0.0

    rows = db.execute(
        "SELECT ot.id, td.name as test_name, COALESCE(ot.price, td.price) as price, "
        "p.full_name as patient_name, v.registration_number, v.created_at "
        "FROM order_tests ot "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id "
        "JOIN visits v ON v.id = o.visit_id "
        "JOIN patients p ON p.id = v.patient_id "
        "WHERE v.referral_center_id=? AND substr(v.created_at,1,7)=? "
        "ORDER BY v.created_at",
        (center_id, month),
    ).fetchall()

    day_map = {}
    test_map = {}
    total = 0.0
    for r in rows:
        price = r["price"] or 0
        total += price
        day_key = (r["created_at"] or "")[:10]
        d = day_map.setdefault(day_key, {"date": day_key, "count": 0, "subtotal": 0.0})
        d["count"] += 1
        d["subtotal"] += price
        tname = r["test_name"]
        tm = test_map.setdefault(tname, {"test_name": tname, "count": 0, "subtotal": 0.0})
        tm["count"] += 1
        tm["subtotal"] += price

    by_day = sorted(day_map.values(), key=lambda x: x["date"])
    by_test = sorted(test_map.values(), key=lambda x: -x["subtotal"])
    return lab, rows, by_day, by_test, total


@app.route("/management/referral-labs/<int:center_id>/statement")
@roles_required("admin", "accountant")
def referral_lab_statement(center_id):
    db = get_db()
    month = request.args.get("month", datetime.now().strftime("%Y-%m"))
    lab, rows, by_day, by_test, total = _referral_lab_statement_data(db, center_id, month)
    if not lab:
        return "Not found", 404
    return render_template("management/referral_lab_statement.html", lab=lab, rows=rows,
                            by_day=by_day, by_test=by_test, total=total, month=month)


@app.route("/management/referral-labs/<int:center_id>/statement/pdf")
@roles_required("admin", "accountant")
def referral_lab_statement_pdf(center_id):
    import pdf_export

    db = get_db()
    month = request.args.get("month", datetime.now().strftime("%Y-%m"))
    lab, rows, by_day, by_test, total = _referral_lab_statement_data(db, center_id, month)
    if not lab:
        return "Not found", 404
    html_content = render_template("management/print_referral_lab_statement.html", lab=lab,
                                    by_day=by_day, by_test=by_test, total=total, month=month)
    pdf_path = pdf_export.make_temp_pdf_path(f"reflab{center_id}_{month}")
    try:
        pdf_export.html_to_pdf(html_content, request.url_root, pdf_path)
    except Exception as exc:
        flash(f"❌ تعذّر توليد ملف PDF لكشف الحساب: {exc}")
        return redirect(url_for("referral_lab_statement", center_id=center_id, month=month))
    log_action("PDF", "referral_center", center_id, f"statement {month}")
    safe_name = re.sub(r"[^\w\-]+", "_", lab["name"])
    return send_file(pdf_path, as_attachment=True, download_name=f"{safe_name}_{month}.pdf")


@app.route("/management/referral-labs/<int:center_id>/statement/whatsapp", methods=["POST"])
@roles_required("admin", "accountant")
def referral_lab_statement_whatsapp(center_id):
    import pdf_export
    import whatsapp_bridge

    db = get_db()
    month = request.args.get("month", datetime.now().strftime("%Y-%m"))
    lab, rows, by_day, by_test, total = _referral_lab_statement_data(db, center_id, month)
    if not lab:
        return "Not found", 404
    if not lab["phone"]:
        flash("لا يوجد رقم واتساب مسجّل لهذا المختبر — أضِفه أولًا من نفس هذي الصفحة.")
        return redirect(url_for("referral_labs_list"))

    html_content = render_template("management/print_referral_lab_statement.html", lab=lab,
                                    by_day=by_day, by_test=by_test, total=total, month=month)
    pdf_path = pdf_export.make_temp_pdf_path(f"reflab{center_id}_{month}")
    try:
        pdf_export.html_to_pdf(html_content, request.url_root, pdf_path)
    except Exception as exc:
        flash(f"❌ تعذّر توليد ملف PDF لإرساله: {exc}")
        return redirect(url_for("referral_lab_statement", center_id=center_id, month=month))

    try:
        whatsapp_bridge.send_pdf(
            lab["phone"], pdf_path,
            caption=f"كشف حساب {lab['name']} — {month}",
            country_code=get_setting(db, "whatsapp_country_code", "964"),
        )
        flash(f"✅ تم إرسال كشف حساب {month} عبر واتساب لمدير {lab['name']}.")
        log_action("WhatsAppSend", "referral_center", center_id, f"statement {month} OK")
    except whatsapp_bridge.WhatsAppSendError as exc:
        flash(f"❌ تعذّر الإرسال: {exc}")
        log_action("WhatsAppSend", "referral_center", center_id, f"statement {month} failed: {exc}")
    return redirect(url_for("referral_lab_statement", center_id=center_id, month=month))


@app.route("/front-desk/visits")
@login_required
def visits_list():
    db = get_db()
    q = request.args.get("q", "").strip()
    show_all = request.args.get("all") == "1"
    # فلترة اختيارية بطبيب مُحيل أو مختبر مُرسِل — تُستخدم من نتائج بحث
    # الرئيسية (بحث باسم طبيب/مختبر) لعرض كل زياراته مرتبة بالأحدث تاريخًا
    # ووقتًا، بغض النظر عن يوم الزيارة (بعكس الوضع الافتراضي المقتصر على
    # اليوم الحالي فقط).
    doctor_id = request.args.get("doctor_id", type=int)
    referral_center_id = request.args.get("referral_center_id", type=int)
    filter_label = None
    if doctor_id:
        row = db.execute("SELECT full_name FROM doctors WHERE id=?", (doctor_id,)).fetchone()
        if row:
            filter_label = f"زيارات الدكتور المُحيل: {row['full_name']}"
    elif referral_center_id:
        row = db.execute("SELECT name FROM referral_centers WHERE id=?", (referral_center_id,)).fetchone()
        if row:
            filter_label = f"زيارات المختبر المُرسِل: {row['name']}"

    query = (
        "SELECT v.id, v.registration_number, p.full_name, p.gender, p.age, p.phone, "
        "v.status, v.created_at, v.attending_doctor, "
        "(SELECT COALESCE(SUM(ot.price),0) FROM order_tests ot "
        " JOIN orders o ON o.id=ot.order_id WHERE o.visit_id=v.id) as total, "
        # الواصل (المبلغ المدفوع فعليًا) والإجمالي الفعلي بالفاتورة (بعد أي
        # خصم/رسوم إضافية) — يُستخدمان بالقالب لعرض "الواصل" و"الباقي" بدل
        # حقل "المدفوع" القديم. LEFT JOIN لأن بعض الزيارات القديمة قد لا
        # تملك صف فاتورة.
        "COALESCE(i.paid_amount, 0) as paid_amount, i.total_amount as invoice_total "
        "FROM visits v JOIN patients p ON p.id = v.patient_id "
        "LEFT JOIN invoices i ON i.visit_id = v.id "
    )
    params = []
    conditions = []
    if doctor_id:
        conditions.append("v.doctor_id=?")
        params.append(doctor_id)
    elif referral_center_id:
        conditions.append("v.referral_center_id=?")
        params.append(referral_center_id)
    elif q:
        conditions.append("(p.full_name LIKE ? OR p.phone LIKE ? OR v.registration_number LIKE ?)")
        params += [f"%{q}%", f"%{q}%", f"%{q}%"]
    elif not show_all:
        # الصفحة تعرض زيارات اليوم الحالي فقط افتراضيًا — تُصفَّر تلقائيًا كل
        # يوم جديد. السجل الكامل لكل الأيام السابقة يبقى متوفرًا دائمًا عبر
        # صفحة "السجلات" (التقرير اليومي)، أو بالبحث بالاسم/الرقم هنا.
        today = date.today().isoformat()
        conditions.append("substr(v.created_at,1,10)=?")
        params.append(today)
    if conditions:
        query += "WHERE " + " AND ".join(conditions) + " "
    query += "ORDER BY v.created_at DESC, v.id DESC LIMIT 100"
    visits = db.execute(query, params).fetchall()
    return render_template("front_desk/visits.html", visits=visits, q=q, show_all=show_all,
                            filter_label=filter_label)

@app.route("/front-desk/visits/<int:visit_id>/delete", methods=["POST"])
@roles_required("admin")
def delete_visit(visit_id):
    """حذف زيارة بالكامل مع كل ما يتبعها (طلبات، نتائج، فواتير، مدفوعات،
    متابعات، إرسالات واتساب)، ثم — إذا صار المريض بلا أي زيارة أخرى — يُحذف
    المريض نفسه فينحذف من صفحة/أيقونة المرضى، وكذلك إذا صار الدكتور المرسل
    (doctor_id) غير مرتبط بأي زيارة أو طلب تحليل آخر يُحذف هو الآخر تلقائيًا."""
    db = get_db()
    visit = db.execute("SELECT * FROM visits WHERE id=?", (visit_id,)).fetchone()
    if not visit:
        return "Not found", 404

    patient_id = visit["patient_id"]
    doctor_id = visit["doctor_id"]

    order_ids = [r["id"] for r in db.execute("SELECT id FROM orders WHERE visit_id=?", (visit_id,)).fetchall()]
    if order_ids:
        placeholders = ",".join("?" * len(order_ids))
        order_test_ids = [r["id"] for r in db.execute(
            f"SELECT id FROM order_tests WHERE order_id IN ({placeholders})", order_ids
        ).fetchall()]
        if order_test_ids:
            ot_placeholders = ",".join("?" * len(order_test_ids))
            db.execute(f"DELETE FROM result_history WHERE order_test_id IN ({ot_placeholders})", order_test_ids)
            db.execute(f"DELETE FROM results WHERE order_test_id IN ({ot_placeholders})", order_test_ids)
        db.execute(f"DELETE FROM order_tests WHERE order_id IN ({placeholders})", order_ids)
    db.execute("DELETE FROM orders WHERE visit_id=?", (visit_id,))
    db.execute(
        "DELETE FROM payments WHERE invoice_id IN (SELECT id FROM invoices WHERE visit_id=?)",
        (visit_id,),
    )
    db.execute("DELETE FROM invoices WHERE visit_id=?", (visit_id,))
    db.execute("DELETE FROM removed_order_tests WHERE visit_id=?", (visit_id,))
    db.execute("DELETE FROM patient_followups WHERE visit_id=?", (visit_id,))
    db.execute("DELETE FROM whatsapp_sends WHERE visit_id=?", (visit_id,))
    db.execute("DELETE FROM visits WHERE id=?", (visit_id,))
    db.commit()
    log_action("Delete", "visit", visit_id, visit["registration_number"])

    # المريض: يُحذف فقط إذا ما عنده أي زيارة أخرى، حتى ما نحذف مريض له سجل
    # فعلي بزيارات ثانية.
    remaining_visits = db.execute(
        "SELECT COUNT(*) c FROM visits WHERE patient_id=?", (patient_id,)
    ).fetchone()["c"]
    if remaining_visits == 0:
        patient = db.execute("SELECT * FROM patients WHERE id=?", (patient_id,)).fetchone()
        if patient:
            db.execute("DELETE FROM patient_followups WHERE patient_id=?", (patient_id,))
            db.execute("DELETE FROM patients WHERE id=?", (patient_id,))
            db.commit()
            log_action("Delete", "patient", patient_id, patient["full_name"])

    # الدكتور المرسل: يُحذف فقط إذا ما بقى مرتبط بأي زيارة أو طلب تحليل آخر.
    if doctor_id:
        linked_visits = db.execute(
            "SELECT COUNT(*) c FROM visits WHERE doctor_id=?", (doctor_id,)
        ).fetchone()["c"]
        linked_orders = db.execute(
            "SELECT COUNT(*) c FROM order_tests WHERE doctor_id=?", (doctor_id,)
        ).fetchone()["c"]
        if not linked_visits and not linked_orders:
            doctor = db.execute("SELECT * FROM doctors WHERE id=?", (doctor_id,)).fetchone()
            if doctor:
                db.execute("DELETE FROM doctor_test_prices WHERE doctor_id=?", (doctor_id,))
                db.execute("DELETE FROM doctors WHERE id=?", (doctor_id,))
                db.commit()
                log_action("Delete", "doctor", doctor_id, doctor["full_name"])

    flash("تم حذف الزيارة، والمريض والدكتور المرسل أيضًا إذا صاروا بلا أي زيارات أخرى مرتبطة.")
    return redirect(url_for("visits_list"))


@app.route("/front-desk/visits/<int:visit_id>/edit", methods=["GET", "POST"])
@login_required
def visit_edit(visit_id):
    db = get_db()
    visit = db.execute(
        "SELECT v.*, p.full_name, p.gender, p.age, p.age_unit, p.phone, p.title, p.email, "
        "p.national_id, p.passport_number, p.travel_certificate_number, p.lab_card_number "
        "FROM visits v JOIN patients p ON p.id = v.patient_id WHERE v.id=?", (visit_id,),
    ).fetchone()
    if not visit:
        return "Not found", 404
    order = db.execute("SELECT * FROM orders WHERE visit_id=?", (visit_id,)).fetchone()
    invoice = db.execute("SELECT * FROM invoices WHERE visit_id=?", (visit_id,)).fetchone()

    if request.method == "POST":
        if invoice and invoice["is_locked"]:
            flash("This invoice is locked and cannot be edited.")
            return redirect(url_for("visits_list"))

        db.execute(
            "UPDATE patients SET full_name=?, gender=?, age=?, age_unit=?, phone=?, title=?, email=?, "
            "national_id=?, passport_number=?, travel_certificate_number=?, lab_card_number=? WHERE id=?",
            (request.form.get("patient_name", "").strip(), request.form.get("gender"),
             request.form.get("age") or None, request.form.get("age_unit") or "Years",
             request.form.get("phone"), request.form.get("title", "Mr."), request.form.get("email", "").strip(),
             request.form.get("national_id", "").strip(), request.form.get("passport_number", "").strip(),
             request.form.get("travel_certificate_number", "").strip(), request.form.get("lab_card_number", "").strip(),
             visit["patient_id"]),
        )
        referring_doctor_id = find_or_create_doctor(db, request.form.get("doctor_name", ""))
        doctor_changed = referring_doctor_id != visit["doctor_id"]

        examining_doctor = request.form.get("examining_doctor", "")
        # نفس منطق شاشة "زيارة جديدة" — attending_doctor يتبع examining_doctor
        # تلقائيًا الآن بعد ما صار حقلاً واحدًا بالشاشة.
        attending_doctor = examining_doctor
        referral_center_id = find_or_create_referral_center(db, request.form.get("referral_lab_name", ""))
        # المعلومات الصحية + الزيارة المنزلية — نفس حقول شاشة "زيارة جديدة".
        weight = request.form.get("weight") or None
        height = request.form.get("height") or None
        symptoms = request.form.get("symptoms", "")
        disease = request.form.get("disease", "")
        therapy = request.form.get("therapy", "")
        is_home_visit = 1 if request.form.get("is_home_visit") == "1" else 0
        home_visit_address = request.form.get("home_visit_address", "") if is_home_visit else ""
        try:
            home_visit_fee = float(request.form.get("home_visit_fee") or 0) if is_home_visit else 0.0
        except ValueError:
            home_visit_fee = 0.0
        db.execute(
            "UPDATE visits SET doctor_id=?, referral_center_id=?, fasting=?, notes=?, examining_doctor=?, "
            "expenses=?, attending_doctor=?, weight=?, height=?, symptoms=?, disease=?, therapy=?, "
            "is_home_visit=?, home_visit_address=?, home_visit_fee=? WHERE id=?",
            (referring_doctor_id, referral_center_id, request.form.get("fasting", "Undefined"), request.form.get("notes", ""),
             examining_doctor, request.form.get("expenses") or 0, attending_doctor,
             weight, height, symptoms, disease, therapy,
             is_home_visit, home_visit_address, home_visit_fee, visit_id))

        new_test_id = request.form.get("add_test")
        if new_test_id:
            test = db.execute("SELECT * FROM test_definitions WHERE id=?", (new_test_id,)).fetchone()
            if test:
                reg = visit["registration_number"]
                barcode = f"{reg}{new_test_id.zfill(3)}"
                price = get_test_price(db, new_test_id, referring_doctor_id)
                db.execute(
                    "INSERT INTO order_tests (order_id, test_definition_id, status, barcode, price, doctor_id, created_at) "
                    "VALUES (?, ?, 'Accepted', ?, ?, ?, ?)",
                    (order["id"], new_test_id, barcode, price, referring_doctor_id,
                     datetime.now().isoformat(timespec="seconds")),
                )

        # نعيد حساب أجر دكتور المختبر الفاحص حسب كل الفحوصات الحالية بالزيارة
        # (بعد أي حذف/إضافة) واسم الدكتور الفاحص المختار حاليًا.
        current_test_ids = [
            row["test_definition_id"]
            for row in db.execute("SELECT test_definition_id FROM order_tests WHERE order_id=?", (order["id"],)).fetchall()
        ]
        examining_doctor_fee = compute_examining_doctor_fee(db, examining_doctor, current_test_ids)
        db.execute("UPDATE visits SET examining_doctor_fee=? WHERE id=?", (examining_doctor_fee, visit_id))

        # إذا انتغيّر الطبيب المُحيل، نعيد تسعير كل التحاليل الموجودة أصلاً
        # بالزيارة حسب جدول أسعار الطبيب الجديد (أو السعر الافتراضي إذا ماكو
        # طبيب / ماكو سعر خاص إله).
        if doctor_changed:
            existing = db.execute(
                "SELECT id, test_definition_id FROM order_tests WHERE order_id=?", (order["id"],)
            ).fetchall()
            for row in existing:
                new_price = get_test_price(db, row["test_definition_id"], referring_doctor_id)
                db.execute("UPDATE order_tests SET price=?, doctor_id=? WHERE id=?",
                           (new_price, referring_doctor_id, row["id"]))

        extra_charges = float(request.form.get("extra_charges") or 0)
        discount_amount = float(request.form.get("discount_amount") or 0)
        new_total = db.execute(
            "SELECT COALESCE(SUM(ot.price),0) as total FROM order_tests ot WHERE ot.order_id=?",
            (order["id"],),
        ).fetchone()["total"]
        # أجرة الزيارة المنزلية (إن فُعِّلت) تُضاف لمجموع الفحوصات، نفس أسلوب
        # شاشة "زيارة جديدة" — منفصلة عن حقل "Extra Charges" اليدوي.
        new_total += home_visit_fee
        if invoice:
            db.execute(
                "UPDATE invoices SET total_amount=?, extra_charges=?, discount_amount=? WHERE id=?",
                (new_total + extra_charges, extra_charges, discount_amount, invoice["id"]),
            )
        db.commit()
        log_action("Update", "visit", visit_id)
        flash("Visit updated.")
        return redirect(url_for("visits_list"))

    order_tests = db.execute(
        "SELECT ot.id, ot.status, ot.test_definition_id, td.name, COALESCE(ot.price, td.price) as price "
        "FROM order_tests ot JOIN test_definitions td ON td.id = ot.test_definition_id WHERE ot.order_id=?",
        (order["id"] if order else 0,),
    ).fetchall()
    all_tests = db.execute("SELECT * FROM test_definitions WHERE is_active=1 ORDER BY department, name").fetchall()
    current_doctor_name = ""
    if visit["doctor_id"]:
        d = db.execute("SELECT full_name FROM doctors WHERE id=?", (visit["doctor_id"],)).fetchone()
        current_doctor_name = d["full_name"] if d else ""
    current_referral_lab_name = ""
    if visit["referral_center_id"]:
        rl = db.execute("SELECT name FROM referral_centers WHERE id=?", (visit["referral_center_id"],)).fetchone()
        current_referral_lab_name = rl["name"] if rl else ""
    doctors, doctor_prices, doctor_name_to_id = doctor_pricing_context(db)
    test_default_prices = {str(tst["id"]): tst["price"] for tst in all_tests}
    last_removed_row = db.execute(
        "SELECT rot.*, td.name as test_name FROM removed_order_tests rot "
        "JOIN test_definitions td ON td.id = rot.test_definition_id "
        "WHERE rot.visit_id=? AND rot.restored=0 ORDER BY rot.id DESC LIMIT 1",
        (visit_id,),
    ).fetchone()
    last_removed = dict(last_removed_row) if last_removed_row else None

    examining_test_ids = [row["id"] for row in get_examining_tests(db)]
    referral_labs = db.execute(
        "SELECT * FROM referral_centers WHERE name != 'Walk-in' ORDER BY name"
    ).fetchall()
    return render_template("front_desk/visit_edit.html", visit=visit, order_tests=order_tests,
                            all_tests=all_tests, invoice=invoice, examining_doctors=get_examining_doctors(db),
                            doctors=doctors, doctor_prices=doctor_prices, doctor_name_to_id=doctor_name_to_id,
                            current_doctor_name=current_doctor_name,
                            current_referral_lab_name=current_referral_lab_name,
                            test_default_prices=test_default_prices,
                            last_removed=last_removed, examining_test_ids=examining_test_ids,
                            examining_rates=get_examining_rates_map(db), referral_labs=referral_labs)


@app.route("/front-desk/visits/<int:visit_id>/order-tests/<int:order_test_id>/price", methods=["POST"])
@login_required
def visit_update_test_price(visit_id, order_test_id):
    db = get_db()
    invoice = db.execute("SELECT * FROM invoices WHERE visit_id=?", (visit_id,)).fetchone()
    if invoice and invoice["is_locked"]:
        flash("This invoice is locked and cannot be edited.")
        return redirect(url_for("visit_edit", visit_id=visit_id))
    try:
        new_price = float(request.form.get(f"price_{order_test_id}"))
    except (TypeError, ValueError):
        flash("Invalid price.")
        return redirect(url_for("visit_edit", visit_id=visit_id))
    order_test = db.execute("SELECT * FROM order_tests WHERE id=?", (order_test_id,)).fetchone()
    if not order_test:
        return "Not found", 404
    db.execute("UPDATE order_tests SET price=? WHERE id=?", (new_price, order_test_id))
    if invoice:
        order = db.execute("SELECT * FROM orders WHERE visit_id=?", (visit_id,)).fetchone()
        new_total = db.execute(
            "SELECT COALESCE(SUM(ot.price),0) as total FROM order_tests ot WHERE ot.order_id=?",
            (order["id"],),
        ).fetchone()["total"]
        db.execute(
            "UPDATE invoices SET total_amount=? WHERE id=?",
            (new_total + (invoice["extra_charges"] or 0), invoice["id"]),
        )
    db.commit()
    log_action("Update", "order_test_price", order_test_id)
    flash("Price updated.")
    return redirect(url_for("visit_edit", visit_id=visit_id))


@app.route("/front-desk/visits/<int:visit_id>/remove-test/<int:order_test_id>", methods=["POST"])
@login_required
def visit_remove_test(visit_id, order_test_id):
    db = get_db()
    invoice = db.execute("SELECT * FROM invoices WHERE visit_id=?", (visit_id,)).fetchone()
    if invoice and invoice["is_locked"]:
        flash("This invoice is locked.")
        return redirect(url_for("visit_edit", visit_id=visit_id))

    ot_row = db.execute("SELECT * FROM order_tests WHERE id=?", (order_test_id,)).fetchone()
    if not ot_row:
        return redirect(url_for("visit_edit", visit_id=visit_id))
    result_rows = db.execute("SELECT * FROM results WHERE order_test_id=?", (order_test_id,)).fetchall()
    snapshot = {
        "order_test": dict(ot_row),
        "results": [dict(r) for r in result_rows],
    }
    now = datetime.now().isoformat(timespec="seconds")
    db.execute(
        "INSERT INTO removed_order_tests (visit_id, order_id, test_definition_id, snapshot, removed_by, removed_at) "
        "VALUES (?, ?, ?, ?, ?, ?)",
        (visit_id, ot_row["order_id"], ot_row["test_definition_id"], json.dumps(snapshot), session["user_id"], now),
    )

    db.execute("DELETE FROM results WHERE order_test_id=?", (order_test_id,))
    db.execute("DELETE FROM order_tests WHERE id=?", (order_test_id,))
    order = db.execute("SELECT * FROM orders WHERE visit_id=?", (visit_id,)).fetchone()
    new_total = db.execute(
        "SELECT COALESCE(SUM(ot.price),0) as total FROM order_tests ot WHERE ot.order_id=?",
        (order["id"],),
    ).fetchone()["total"]
    if invoice:
        db.execute("UPDATE invoices SET total_amount=? WHERE id=?",
                   (new_total + (invoice["extra_charges"] or 0), invoice["id"]))
    db.commit()
    log_action("RemoveTest", "order_test", order_test_id)
    flash("Test removed from visit. You can undo this from the visit edit page.")
    return redirect(url_for("visit_edit", visit_id=visit_id))


@app.route("/front-desk/visits/<int:visit_id>/undo-remove-test", methods=["POST"])
@login_required
def visit_undo_remove_test(visit_id):
    db = get_db()
    invoice = db.execute("SELECT * FROM invoices WHERE visit_id=?", (visit_id,)).fetchone()
    if invoice and invoice["is_locked"]:
        flash("This invoice is locked.")
        return redirect(url_for("visit_edit", visit_id=visit_id))

    last_removed = db.execute(
        "SELECT * FROM removed_order_tests WHERE visit_id=? AND restored=0 ORDER BY id DESC LIMIT 1",
        (visit_id,),
    ).fetchone()
    if not last_removed:
        flash("Nothing to undo.")
        return redirect(url_for("visit_edit", visit_id=visit_id))

    snapshot = json.loads(last_removed["snapshot"])
    ot = snapshot["order_test"]
    cur = db.execute(
        "INSERT INTO order_tests (order_id, test_definition_id, status, barcode, notes, doctor_id, "
        "collected_at, accessioned_at, price, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (ot["order_id"], ot["test_definition_id"], ot["status"], ot["barcode"], ot["notes"],
         ot["doctor_id"], ot["collected_at"], ot["accessioned_at"], ot.get("price"), ot["created_at"]),
    )
    new_order_test_id = cur.lastrowid
    for r in snapshot["results"]:
        db.execute(
            "INSERT INTO results (order_test_id, test_parameter_id, value_numeric, value_text, flag, "
            "entered_by, entered_at, verified_by, verified_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (new_order_test_id, r["test_parameter_id"], r["value_numeric"], r["value_text"], r["flag"],
             r["entered_by"], r["entered_at"], r["verified_by"], r["verified_at"]),
        )
    db.execute("UPDATE removed_order_tests SET restored=1 WHERE id=?", (last_removed["id"],))

    order = db.execute("SELECT * FROM orders WHERE visit_id=?", (visit_id,)).fetchone()
    new_total = db.execute(
        "SELECT COALESCE(SUM(ot.price),0) as total FROM order_tests ot WHERE ot.order_id=?",
        (order["id"],),
    ).fetchone()["total"]
    if invoice:
        db.execute("UPDATE invoices SET total_amount=? WHERE id=?",
                   (new_total + (invoice["extra_charges"] or 0), invoice["id"]))
    db.commit()
    log_action("UndoRemoveTest", "order_test", new_order_test_id)
    flash("Test restored.")
    return redirect(url_for("visit_edit", visit_id=visit_id))



@app.route("/barcode/<code>.png")
def barcode_image(code):
    img = generate_code39(code)
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return send_file(buf, mimetype="image/png")


def _visit_barcode_context(db, visit_id):
    """Visit + its ordered tests, used to build the visit barcode label
    (both the print page and the Code128/QR image endpoints)."""
    visit = db.execute(
        "SELECT v.*, p.full_name, p.age, p.age_unit, p.gender FROM visits v "
        "JOIN patients p ON p.id = v.patient_id WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        return None, None
    tests = db.execute(
        "SELECT td.name, td.name_ar, td.code FROM order_tests ot "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=?",
        (visit_id,),
    ).fetchall()
    # ساعة السحب: أوّل collected_at مسجّل بين تحاليل الزيارة، وإن ما كانت
    # العيّنة مسحوبة بعد نرجع لوقت إنشاء الزيارة كبديل.
    draw_row = db.execute(
        "SELECT MIN(ot.collected_at) as draw_time FROM order_tests ot "
        "JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=? AND ot.collected_at IS NOT NULL",
        (visit_id,),
    ).fetchone()
    draw_time = (draw_row["draw_time"] if draw_row else None) or visit["created_at"]
    return visit, tests, draw_time


@app.route("/front-desk/visits/<int:visit_id>/barcode/code128.png")
@login_required
def visit_barcode_code128(visit_id):
    # Code128 (English/Latin only) carries the registration number plus the
    # English test codes -- this is what lab scanners/instruments read.
    db = get_db()
    visit, tests, draw_time = _visit_barcode_context(db, visit_id)
    if not visit:
        return "Not found", 404
    test_codes = [row["code"] for row in tests if row["code"]]
    reg = visit["registration_number"]
    payload = f"{reg}|{','.join(test_codes)}" if test_codes else str(reg)
    img = generate_code128(payload, caption=str(reg))
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return send_file(buf, mimetype="image/png")


@app.route("/front-desk/visits/<int:visit_id>/barcode/qr.png")
@login_required
def visit_barcode_qr(visit_id):
    # QR carries the Arabic content (Code128 cannot represent it): the
    # patient's full (triple) name plus the Arabic names of their tests.
    db = get_db()
    visit, tests, draw_time = _visit_barcode_context(db, visit_id)
    if not visit:
        return "Not found", 404
    test_names_ar = [(row["name_ar"] or row["name"]) for row in tests]
    lines = [visit["full_name"]]
    if test_names_ar:
        lines.append("، ".join(test_names_ar))
    img = generate_qr("\n".join(lines))
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return send_file(buf, mimetype="image/png")


@app.route("/front-desk/visits/<int:visit_id>/print/barcode")
@login_required
def print_visit_barcode(visit_id):
    db = get_db()
    visit, tests, draw_time = _visit_barcode_context(db, visit_id)
    if not visit:
        return "Not found", 404
    return render_template("front_desk/print_visit_barcode.html", visit=visit, tests=tests, draw_time=draw_time)


@app.route("/front-desk/visits/<int:visit_id>/print/samples")
@login_required
def print_sample_barcodes(visit_id):
    db = get_db()
    visit = db.execute(
        "SELECT v.*, p.full_name FROM visits v JOIN patients p ON p.id = v.patient_id WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        return "Not found", 404
    samples = db.execute(
        "SELECT ot.barcode, td.name as test_name, td.sample_type FROM order_tests ot "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=?",
        (visit_id,),
    ).fetchall()
    return render_template("front_desk/print_sample_barcodes.html", visit=visit, samples=samples)


@app.route("/front-desk/visits/<int:visit_id>/print/invoice")
@login_required
def print_invoice(visit_id):
    db = get_db()
    visit = db.execute(
        "SELECT v.*, p.full_name, p.age, p.gender, p.phone, "
        "rc.name as referral_center_name "
        "FROM visits v JOIN patients p ON p.id = v.patient_id "
        "LEFT JOIN referral_centers rc ON rc.id = v.referral_center_id "
        "WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        return "Not found", 404
    items = db.execute(
        "SELECT td.name, td.name_ar, COALESCE(ot.price, td.price) as price FROM order_tests ot "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=?",
        (visit_id,),
    ).fetchall()
    invoice = db.execute("SELECT * FROM invoices WHERE visit_id=?", (visit_id,)).fetchone()
    doctor = None
    if visit["doctor_id"]:
        doctor = db.execute("SELECT * FROM doctors WHERE id=?", (visit["doctor_id"],)).fetchone()
    return render_template("front_desk/print_invoice.html", visit=visit, items=items,
                            invoice=invoice, doctor=doctor)


# طباعة تحاليل ونتائج زيارة واحدة فقط (تُستخدم من زر "طباعة" أمام أي زيارة
# سابقة تظهر بشاشة "زيارة جديدة" عند العثور على المريض بالبحث الفوري).
@app.route("/front-desk/visits/<int:visit_id>/print/results")
@login_required
def print_visit_results(visit_id):
    db = get_db()
    visit = db.execute(
        "SELECT v.*, p.full_name, p.age, p.age_unit, p.gender, p.phone, "
        "rc.name as referral_center_name FROM visits v "
        "JOIN patients p ON p.id = v.patient_id "
        "LEFT JOIN referral_centers rc ON rc.id = v.referral_center_id WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        return "Not found", 404
    # نفس منطق from_other_lab بصفحة print_report: نعرض العبارة فقط، وما
    # نطبع اسم المختبر المرسل بأي مكان بهذي الصفحة (الجدول الموحّد للزيارة).
    from_other_lab = bool(visit["referral_center_id"]) and visit["referral_center_name"] not in (None, "", "Walk-in")
    # completed_only=1: يُستخدم من زرّي "طباعة/إرسال كل النتائج المكتملة" أعلى
    # شاشة إدخال النتائج، حتى لا تُطبع/تُرسل تحاليل لسا بانتظار الإدخال أو
    # الاعتماد ضمن نفس الملف.
    completed_only = request.args.get("completed_only") == "1"
    query = (
        "SELECT ot.id, ot.barcode, ot.status, td.name as test_name, td.name_ar as test_name_ar "
        "FROM order_tests ot JOIN orders o ON o.id = ot.order_id "
        "JOIN test_definitions td ON td.id = ot.test_definition_id WHERE o.visit_id=?"
    )
    order_tests = db.execute(query, (visit_id,)).fetchall()
    if completed_only:
        order_tests = [ot for ot in order_tests if ot["status"] in ("Completed", "Verified")]
    tests_with_results = []
    for ot in order_tests:
        results = db.execute(
            "SELECT r.*, tp.name as param_name, tp.unit FROM results r "
            "JOIN test_parameters tp ON tp.id = r.test_parameter_id WHERE r.order_test_id=?",
            (ot["id"],),
        ).fetchall()
        tests_with_results.append({"order_test": ot, "results": results})
    return render_template("front_desk/print_visit_results.html", visit=visit,
                            tests_with_results=tests_with_results, from_other_lab=from_other_lab,
                            completed_only=completed_only)


# "طباعة كل النتائج المكتملة" بشاشة إدخال النتائج — بدل الجدول الموحّد
# العام (print_visit_results فوق، الذي لا يستخدم تصميم التقرير الخاص بكل
# تحليل)، هذا الزر يفتح لكل تحليل مكتمل تقريره المخصّص الحقيقي (نفس
# print_report المستخدم من زر "🖨 طباعة التقرير" أسفل كل تحليل مباشرة) —
# كل واحد بنافذة/تبويب منفصلة، حتى يطابق شكل التقرير المطبوع تماماً شكل
# تقرير نفس التحليل لو طبعته لحاله. صفحة "مُشغّل" خفيفة فقط تفتح الروابط
# تلقائياً + تعرضها كروابط احتياطية لو المتصفح منع النوافذ المنبثقة.
@app.route("/front-desk/visits/<int:visit_id>/print/results-bundle")
@login_required
def print_visit_results_bundle(visit_id):
    db = get_db()
    visit = db.execute(
        "SELECT v.*, p.full_name FROM visits v JOIN patients p ON p.id = v.patient_id WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        return "Not found", 404

    order_tests = db.execute(
        "SELECT ot.id, ot.status, td.code as test_code, td.name as test_name, td.department as test_department "
        "FROM order_tests ot JOIN orders o ON o.id = ot.order_id "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "WHERE o.visit_id=? AND ot.status IN ('Completed', 'Verified') ORDER BY ot.id",
        (visit_id,),
    ).fetchall()

    # نفس منطق الترتيب المستخدم بشاشة "إدخال النتائج" بالضبط (ترتيب يدوي من
    # الإعدادات إن وُجد، وإلا حسب أولوية القسم: كيمياء ← فايروسات ← هرمونات ←
    # فيتامينات ← تخثر ← أمراض الدم) — كانت هذي الصفحة تستخدم ترتيب الإضافة
    # الخام (ot.id) فقط بدون أي فرز، فتفتح التقارير بترتيب عشوائي لا علاقة له
    # بترتيب شاشة الإدخال.
    order_pref_raw = get_setting(db, "results_entry_test_order", "")
    order_pref = [ln.strip() for ln in order_pref_raw.splitlines() if ln.strip()]
    if order_pref:
        rank = {name: i for i, name in enumerate(order_pref)}
        order_tests = sorted(
            order_tests,
            key=lambda row: (
                rank.get(row["test_name"], len(order_pref)),
                department_priority_rank(row["test_department"]),
                row["id"],
            ),
        )
    else:
        order_tests = sorted(
            order_tests,
            key=lambda row: (department_priority_rank(row["test_department"]), row["id"]),
        )

    # تجنّب فتح تقرير "لطاخة الدم" مرتين لو الزيارة فيها BF وRETIC كتحليلين
    # منفصلين (بدل BFRETIC الموحّد) — print_report أصلاً يدمجهم بتقرير واحد
    # لو فتحت أي وحدة منهم (انظر BF_RETIC_LINK)، فنكتفي هنا برابط واحد بس.
    seen_codes = set()
    reports = []
    for ot in order_tests:
        code = ot["test_code"]
        if code in BF_RETIC_LINK and BF_RETIC_LINK[code] in seen_codes:
            continue
        seen_codes.add(code)
        reports.append({"order_test_id": ot["id"], "test_name": ot["test_name"]})

    return render_template("front_desk/print_results_bundle.html", visit=visit, reports=reports)


# الجدول الموحّد: يدمج نتائج زيارة أو أكثر (عادة زيارة قديمة + الزيارة الجديدة)
# بصفحة واحدة، صف لكل باراميتر وعمود لكل زيارة، حتى تتضح مقارنة النتائج عبر
# الزمن دون الحاجة لنسخ أي بيانات فعليًا داخل قاعدة البيانات.
@app.route("/front-desk/visits/combined")
@login_required
def combined_visits_report():
    ids_param = request.args.get("visit_ids", "")
    visit_ids = [int(x) for x in ids_param.split(",") if x.strip().isdigit()]
    if not visit_ids:
        return "لم يتم تحديد أي زيارات لعرضها بجدول موحّد.", 400
    db = get_db()
    placeholders = ",".join("?" * len(visit_ids))
    visits = db.execute(
        f"SELECT v.*, p.full_name, p.age, p.age_unit, p.gender, rc.name as referral_center_name FROM visits v "
        f"JOIN patients p ON p.id = v.patient_id "
        f"LEFT JOIN referral_centers rc ON rc.id = v.referral_center_id "
        f"WHERE v.id IN ({placeholders}) "
        f"ORDER BY v.created_at ASC",
        visit_ids,
    ).fetchall()
    if not visits:
        return "Not found", 404
    # إذا أي زيارة من الزيارات المدموجة بهذا الجدول أصلها نموذج وارد من
    # مختبر آخر، نعرض نفس العبارة الانكليزية أسفل الجدول (بدون ذكر اسم
    # المختبر بأي مكان).
    from_other_lab = any(
        v["referral_center_id"] and v["referral_center_name"] not in (None, "", "Walk-in") for v in visits
    )

    param_order = []
    param_rows = {}
    for v in visits:
        order_tests = db.execute(
            "SELECT ot.id FROM order_tests ot JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=?",
            (v["id"],),
        ).fetchall()
        for ot in order_tests:
            results = db.execute(
                "SELECT r.*, tp.name as param_name, tp.unit FROM results r "
                "JOIN test_parameters tp ON tp.id = r.test_parameter_id WHERE r.order_test_id=?",
                (ot["id"],),
            ).fetchall()
            for r in results:
                key = r["param_name"]
                if key not in param_rows:
                    param_rows[key] = {"unit": r["unit"], "values": {}}
                    param_order.append(key)
                display_value = r["value_text"] if r["value_text"] else r["value_numeric"]
                param_rows[key]["values"][v["id"]] = {"value": display_value, "flag": r["flag"]}

    return render_template("front_desk/combined_report.html", visits=visits,
                            param_order=param_order, param_rows=param_rows, patient=visits[0],
                            from_other_lab=from_other_lab)


def period_breakdown(db, group_len, base_where, base_params):
    """يجمع الزيارات حسب فترة زمنية (يوم أو شهر) — يرجع صفوف فيها الدخل
    والصرفيات والصافي لكل فترة فرعية، بالإضافة إلى إجمالي الفترة كاملة.
    group_len: طول substr(created_at) المستخدم للتجميع (10=يوم، 7=شهر).
    base_where/base_params: شرط SQL لتحديد الفترة الكبيرة (مثلاً شهر معين أو سنة معينة)."""
    revenue_rows = db.execute(
        f"SELECT substr(v.created_at,1,{group_len}) as period, "
        f"COALESCE(SUM(COALESCE(ot.price, td.price)),0) as revenue, COUNT(DISTINCT v.id) as visits_count "
        f"FROM visits v JOIN orders o ON o.visit_id=v.id JOIN order_tests ot ON ot.order_id=o.id "
        f"JOIN test_definitions td ON td.id=ot.test_definition_id "
        f"WHERE {base_where} GROUP BY period", base_params,
    ).fetchall()
    expense_rows = db.execute(
        f"SELECT substr(v.created_at,1,{group_len}) as period, "
        f"COALESCE(SUM(v.expenses + COALESCE(v.examining_doctor_fee,0)),0) as expenses "
        f"FROM visits v WHERE {base_where} GROUP BY period", base_params,
    ).fetchall()
    data = {}
    for r in revenue_rows:
        data[r["period"]] = {"period": r["period"], "revenue": r["revenue"],
                              "visits_count": r["visits_count"], "expenses": 0.0}
    for r in expense_rows:
        data.setdefault(r["period"], {"period": r["period"], "revenue": 0.0,
                                       "visits_count": 0, "expenses": 0.0})
        data[r["period"]]["expenses"] = r["expenses"]
    rows = sorted(data.values(), key=lambda x: x["period"])
    for r in rows:
        r["net"] = r["revenue"] - r["expenses"]
    totals = {
        "revenue": sum(r["revenue"] for r in rows),
        "expenses": sum(r["expenses"] for r in rows),
        "visits_count": sum(r["visits_count"] for r in rows),
    }
    totals["net"] = totals["revenue"] - totals["expenses"]
    return rows, totals


@app.route("/reports/daily")
@login_required
def daily_report():
    db = get_db()
    report_date = request.args.get("date", date.today().isoformat())
    visits = db.execute(
        "SELECT v.id, v.registration_number, v.examining_doctor, v.expenses, v.examining_doctor_fee, "
        "v.notes, p.full_name, p.gender, p.age FROM visits v "
        "JOIN patients p ON p.id = v.patient_id "
        "WHERE substr(v.created_at,1,10)=? ORDER BY v.registration_number",
        (report_date,),
    ).fetchall()

    rows = []
    grand_total = 0
    grand_expenses = 0
    for v in visits:
        items = db.execute(
            "SELECT td.name, COALESCE(ot.price, td.price) as price, ot.doctor_id FROM order_tests ot "
            "JOIN test_definitions td ON td.id = ot.test_definition_id "
            "JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=?",
            (v["id"],),
        ).fetchall()
        test_names = ", ".join(i["name"] for i in items)
        total = sum((i["price"] or 0) for i in items)
        grand_total += total
        row_expenses = (v["expenses"] or 0) + (v["examining_doctor_fee"] or 0)
        grand_expenses += row_expenses
        doctor_names = set()
        for i in items:
            if i["doctor_id"]:
                d = db.execute("SELECT full_name FROM doctors WHERE id=?", (i["doctor_id"],)).fetchone()
                if d:
                    doctor_names.add(d["full_name"])
        rows.append({
            "reg": v["registration_number"], "name": v["full_name"],
            "gender": v["gender"] or "-", "age": v["age"] if v["age"] is not None else "-",
            "tests": test_names, "total": total, "doctors": ", ".join(doctor_names),
            "examining_doctor": v["examining_doctor"] or "-",
            "expenses": row_expenses,
            "notes": v["notes"] or "",
        })

    return render_template("front_desk/daily_report.html", rows=rows, report_date=report_date,
                            grand_total=grand_total, grand_expenses=grand_expenses,
                            grand_net=grand_total - grand_expenses)


@app.route("/reports/monthly")
@roles_required("admin", "accountant", "supervisor")
def monthly_report():
    db = get_db()
    month = request.args.get("month", date.today().strftime("%Y-%m"))
    rows, totals = period_breakdown(db, 10, "substr(v.created_at,1,7)=?", (month,))
    return render_template("front_desk/monthly_report.html", rows=rows, totals=totals, month=month)


@app.route("/reports/half-year")
@roles_required("admin", "accountant", "supervisor")
def half_year_report():
    db = get_db()
    year = request.args.get("year", str(date.today().year))
    half = request.args.get("half", "1" if date.today().month <= 6 else "2")
    start_month, end_month = (1, 6) if half == "1" else (7, 12)
    rows, totals = period_breakdown(
        db, 7,
        "substr(v.created_at,1,4)=? AND CAST(substr(v.created_at,6,2) AS INTEGER) BETWEEN ? AND ?",
        (year, start_month, end_month),
    )
    return render_template("front_desk/half_year_report.html", rows=rows, totals=totals, year=year, half=half)


@app.route("/reports/annual")
@roles_required("admin", "accountant", "supervisor")
def annual_report():
    db = get_db()
    year = request.args.get("year", str(date.today().year))
    rows, totals = period_breakdown(db, 7, "substr(v.created_at,1,4)=?", (year,))
    return render_template("front_desk/annual_report.html", rows=rows, totals=totals, year=year)



@app.route("/workbench/samples-collection")
@login_required
def samples_collection():
    db = get_db()
    order_tests = db.execute(
        "SELECT ot.id, ot.barcode, ot.status, td.name as test_name, td.sample_type, "
        "p.full_name as patient_name, v.registration_number "
        "FROM order_tests ot JOIN test_definitions td ON td.id=ot.test_definition_id "
        "JOIN orders o ON o.id=ot.order_id JOIN visits v ON v.id=o.visit_id "
        "JOIN patients p ON p.id=v.patient_id WHERE ot.status='Accepted' ORDER BY ot.id DESC"
    ).fetchall()
    return render_template("workbench/samples_collection.html", order_tests=order_tests)


@app.route("/workbench/samples-collection/<int:order_test_id>/collect", methods=["POST"])
@login_required
def collect_sample(order_test_id):
    db = get_db()
    now = datetime.now().isoformat(timespec="seconds")
    db.execute("UPDATE order_tests SET status='Collected', collected_at=? WHERE id=?", (now, order_test_id))
    db.commit()
    log_action("Collect", "order_test", order_test_id)
    flash("Sample marked as collected.")
    return redirect(url_for("samples_collection"))


@app.route("/workbench/samples-accession")
@login_required
def samples_accession():
    db = get_db()
    order_tests = db.execute(
        "SELECT ot.id, ot.barcode, ot.status, td.name as test_name, td.sample_type, td.department, "
        "p.full_name as patient_name, v.registration_number "
        "FROM order_tests ot JOIN test_definitions td ON td.id=ot.test_definition_id "
        "JOIN orders o ON o.id=ot.order_id JOIN visits v ON v.id=o.visit_id "
        "JOIN patients p ON p.id=v.patient_id WHERE ot.status='Collected' ORDER BY ot.id DESC"
    ).fetchall()
    return render_template("workbench/samples_accession.html", order_tests=order_tests)


@app.route("/workbench/samples-accession/<int:order_test_id>/accept", methods=["POST"])
@login_required
def accept_sample(order_test_id):
    db = get_db()
    now = datetime.now().isoformat(timespec="seconds")
    db.execute("UPDATE order_tests SET status='Accessioned', accessioned_at=? WHERE id=?", (now, order_test_id))
    db.commit()
    log_action("Accession", "order_test", order_test_id)
    flash("Sample accepted for testing.")
    return redirect(url_for("samples_accession"))



@app.route("/workbench/orders")
@login_required
def orders_list():
    db = get_db()
    status = request.args.get("status", "")
    query = (
        "SELECT ot.id, ot.barcode, ot.status, ot.created_at, "
        "td.name as test_name, td.code as test_code, td.department, td.sample_type, "
        "p.full_name as patient_name, v.registration_number "
        "FROM order_tests ot "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id "
        "JOIN visits v ON v.id = o.visit_id "
        "JOIN patients p ON p.id = v.patient_id "
    )
    params = []
    if status:
        query += "WHERE ot.status = ? "
        params.append(status)
    query += "ORDER BY ot.id DESC LIMIT 150"
    order_tests = db.execute(query, params).fetchall()
    return render_template("workbench/orders.html", order_tests=order_tests, status=status)


def save_order_test_results(db, ot, parameters, form, user_id, field_prefix=""):
    """يحفظ نتائج تحليل واحد (order_test) من بيانات نموذج مُرسل، بنفس منطق
    الأعلام (Flag) والمرجعيات وتسجيل التأريخ المستخدم بشاشة إدخال النتائج —
    مشتركة بين شاشة الإدخال المفردة وشاشة الإدخال المُجمّع لكل تحاليل الزيارة
    معًا، حتى لا يتكرر نفس المنطق الحساس بمكانين. تعيد True إذا أُدخلت/عُدّلت
    أي قيمة فعليًا لهذا التحليل."""
    now = datetime.now().isoformat(timespec="seconds")
    touched = False
    for param in parameters:
        field = f"{field_prefix}param_{param['id']}"
        value = form.get(field, "").strip()
        if value == "":
            continue
        touched = True
        rng = find_reference_range(db, param["id"], ot["gender"], ot["age"], ot["age_unit"])
        flag = "Normal"
        value_numeric = None
        value_text = None
        if param["result_type"] == "Numeric":
            try:
                value_numeric = float(value)
                if rng and rng["low"] is not None and value_numeric < rng["low"]:
                    flag = "Low"
                elif rng and rng["high"] is not None and value_numeric > rng["high"]:
                    flag = "High"
                    if rng["high"] and value_numeric > rng["high"] * 2:
                        flag = "Critical"
            except ValueError:
                value_text = value
        else:
            value_text = value

        existing = db.execute(
            "SELECT id, value_numeric, value_text, flag FROM results WHERE order_test_id=? AND test_parameter_id=?",
            (ot["id"], param["id"]),
        ).fetchone()
        if existing:
            if existing["value_numeric"] != value_numeric or existing["value_text"] != value_text:
                db.execute(
                    "INSERT INTO result_history (result_id, order_test_id, test_parameter_id, "
                    "prev_value_numeric, prev_value_text, prev_flag, changed_by, changed_at) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (existing["id"], ot["id"], param["id"], existing["value_numeric"],
                     existing["value_text"], existing["flag"], user_id, now),
                )
            db.execute(
                "UPDATE results SET value_numeric=?, value_text=?, flag=?, entered_by=?, entered_at=? WHERE id=?",
                (value_numeric, value_text, flag, user_id, now, existing["id"]),
            )
        else:
            db.execute(
                "INSERT INTO results (order_test_id, test_parameter_id, value_numeric, value_text, "
                "flag, entered_by, entered_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
                (ot["id"], param["id"], value_numeric, value_text, flag, user_id, now),
            )
    return touched


@app.route("/workbench/result/<int:order_test_id>", methods=["GET", "POST"])
@login_required
def result_entry(order_test_id):
    db = get_db()
    ot = db.execute(
        "SELECT ot.*, td.name as test_name, td.code as test_code, p.full_name as patient_name, p.gender, p.age, p.age_unit, "
        "v.id as visit_id, v.registration_number "
        "FROM order_tests ot JOIN test_definitions td ON td.id=ot.test_definition_id "
        "JOIN orders o ON o.id=ot.order_id JOIN visits v ON v.id=o.visit_id "
        "JOIN patients p ON p.id=v.patient_id WHERE ot.id=?",
        (order_test_id,),
    ).fetchone()
    if not ot:
        return "Not found", 404

    parameters = db.execute(
        "SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id", (ot["test_definition_id"],)
    ).fetchall()

    if request.method == "POST":
        save_order_test_results(db, ot, parameters, request.form, session["user_id"])
        # التحليل المعتمد (Verified) يبقى معتمَد بعد التعديل — الحقول صارت
        # قابلة للتعديل دائمًا حتى بعد الاعتماد (بطلب المستخدم)، وكل تعديل
        # يبقى مسجّل بجدول result_history (القيمة القديمة + مين عدّلها ووقتها)
        # حتى لو ما تغيّرت حالة الاعتماد ظاهريًا.
        if ot["status"] != "Verified":
            db.execute("UPDATE order_tests SET status='Completed' WHERE id=?", (order_test_id,))
        db.commit()
        _maybe_auto_whatsapp_send(db, ot["visit_id"])
        _maybe_archive_visit_pdf(db, ot["visit_id"])
        log_action("EnterResult", "order_test", order_test_id)
        flash("Results saved.")
        return redirect(url_for("orders_list"))

    results = db.execute(
        "SELECT r.*, tp.name as param_name FROM results r "
        "JOIN test_parameters tp ON tp.id=r.test_parameter_id WHERE order_test_id=?",
        (order_test_id,),
    ).fetchall()
    results_by_param = {r["test_parameter_id"]: r for r in results}

    ranges = {}
    suggestions_map = {}
    history_map = {}
    for p in parameters:
        r = find_reference_range(db, p["id"], ot["gender"], ot["age"], ot["age_unit"])
        ranges[p["id"]] = r
        sugg = db.execute("SELECT content FROM suggestions WHERE test_parameter_id=?", (p["id"],)).fetchall()
        suggestions_map[p["id"]] = [s["content"] for s in sugg]
        last_hist = db.execute(
            "SELECT * FROM result_history WHERE order_test_id=? AND test_parameter_id=? "
            "ORDER BY id DESC LIMIT 1", (order_test_id, p["id"]),
        ).fetchone()
        history_map[p["id"]] = last_hist

    return render_template("workbench/result_entry.html", ot=ot, parameters=parameters,
                            results_by_param=results_by_param, ranges=ranges,
                            suggestions_map=suggestions_map, history_map=history_map,
                            patient_hct=get_visit_hct(db, ot["visit_id"]))


# شاشة إدخال النتائج المُجمّعة لكل تحاليل الزيارة معًا: تُظهر كل تحليل طُلب
# لهذا المريض كمستطيل (بطاقة) منفصل بنفس صفوفه/حقوله المعتادة، وزر "حفظ"
# واحد يحفظ كل ما تمت تعبئته دفعة واحدة. فور الحفظ تصبح النتائج متاحة مباشرة
# بريبورت كل تحليل (زر طباعة أمام كل بطاقة)، مع بقاء إمكانية التعديل والحفظ
# مجددًا بنفس الصفحة قبل الطباعة الفعلية.
@app.route("/front-desk/visits/<int:visit_id>/results", methods=["GET", "POST"])
@login_required
def visit_results_entry(visit_id):
    db = get_db()
    visit = db.execute(
        "SELECT v.*, p.full_name as patient_name, p.gender, p.age, p.age_unit, p.phone "
        "FROM visits v JOIN patients p ON p.id = v.patient_id WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        return "Not found", 404

    order_tests = db.execute(
        "SELECT ot.*, td.name as test_name, td.code as test_code, td.department as test_department, "
        "p.gender as gender, p.age as age, p.age_unit as age_unit "
        "FROM order_tests ot JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id JOIN visits v2 ON v2.id = o.visit_id "
        "JOIN patients p ON p.id = v2.patient_id "
        "WHERE o.visit_id=? ORDER BY ot.id",
        (visit_id,),
    ).fetchall()
    if not order_tests:
        flash("لا توجد تحاليل مطلوبة بهذي الزيارة.")
        return redirect(url_for("results_list"))

    # ترتيب ثابت لعرض/طباعة/إدخال التحاليل — بغض النظر عن ترتيب طلبها الفعلي
    # بهذي الزيارة (ot.id الافتراضي). الأدمن يحدد الترتيب المطلوب من
    # Management → الإعدادات → "ترتيب التحاليل بشاشة إدخال النتائج" (اسم
    # التحليل بالضبط كما يظهر بالشاشة، مثلاً "HbA1c (BioChemstry)"، سطر لكل
    # تحليل). أي تحليل غير مذكور هناك يُرتَّب افتراضيًا حسب أولوية القسم
    # (DEPARTMENT_PRIORITY_KEYWORDS: كيمياء ← فايروسات ← هرمونات ← فيتامينات
    # ← تخثر ← أمراض الدم)، وبينهم حسب ترتيبهم الأصلي (ot.id) — فلو المريض
    # ماله بعض الأقسام تلقائيًا تنتخطى وتنطبع/تنعرض بس الأقسام الموجودة فعلاً
    # بنفس التسلسل المطلوب.
    order_pref_raw = get_setting(db, "results_entry_test_order", "")
    order_pref = [ln.strip() for ln in order_pref_raw.splitlines() if ln.strip()]
    if order_pref:
        rank = {name: i for i, name in enumerate(order_pref)}
        order_tests = sorted(
            order_tests,
            key=lambda row: (
                rank.get(row["test_name"], len(order_pref)),
                department_priority_rank(row["test_department"]),
                row["id"],
            ),
        )
    else:
        order_tests = sorted(
            order_tests,
            key=lambda row: (department_priority_rank(row["test_department"]), row["id"]),
        )

    if request.method == "POST":
        any_saved = False
        for ot in order_tests:
            # التحاليل المعتمدة (Verified) صارت قابلة للتعديل دائمًا هيه
            # بعد، وتبقى معتمَدة بعد الحفظ (ما تحتاج فتح/إلغاء اعتماد أولًا)
            # — كل تعديل يبقى مسجّل بجدول result_history (القيمة القديمة +
            # مين عدّلها ووقتها) بغض النظر عن حالة الاعتماد.
            parameters = db.execute(
                "SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id", (ot["test_definition_id"],)
            ).fetchall()
            touched = save_order_test_results(
                db, ot, parameters, request.form, session["user_id"], field_prefix=f"ot{ot['id']}_"
            )
            if touched:
                any_saved = True
                if ot["status"] != "Verified":
                    db.execute("UPDATE order_tests SET status='Completed' WHERE id=?", (ot["id"],))
        db.commit()
        if any_saved:
            _maybe_auto_whatsapp_send(db, visit_id)
            _maybe_archive_visit_pdf(db, visit_id)
            log_action("EnterResultsBulk", "visit", visit_id)
            flash("تم حفظ النتائج.")
        else:
            flash("لم تُدخل أي قيمة جديدة.")
        return redirect(url_for("visit_results_entry", visit_id=visit_id))

    boxes = []
    for ot in order_tests:
        parameters = db.execute(
            "SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id", (ot["test_definition_id"],)
        ).fetchall()
        results = db.execute(
            "SELECT r.*, tp.name as param_name FROM results r "
            "JOIN test_parameters tp ON tp.id=r.test_parameter_id WHERE order_test_id=?",
            (ot["id"],),
        ).fetchall()
        results_by_param = {r["test_parameter_id"]: r for r in results}
        ranges = {}
        suggestions_map = {}
        history_map = {}
        for p in parameters:
            ranges[p["id"]] = find_reference_range(db, p["id"], visit["gender"], visit["age"], visit["age_unit"])
            sugg = db.execute("SELECT content FROM suggestions WHERE test_parameter_id=?", (p["id"],)).fetchall()
            suggestions_map[p["id"]] = [s["content"] for s in sugg]
            history_map[p["id"]] = db.execute(
                "SELECT * FROM result_history WHERE order_test_id=? AND test_parameter_id=? "
                "ORDER BY id DESC LIMIT 1",
                (ot["id"], p["id"]),
            ).fetchone()
        boxes.append({
            "ot": ot, "parameters": parameters, "results_by_param": results_by_param,
            "ranges": ranges, "suggestions_map": suggestions_map, "history_map": history_map,
        })

    return render_template("front_desk/visit_results_entry.html", visit=visit, boxes=boxes,
                            patient_hct=get_visit_hct(db, visit_id),
                            patient_hct_normal=get_normal_hct(db, visit["gender"], visit["age"], visit["age_unit"]))


# ------------------------------------------------------------------------
# "اللوحة المجمّعة" — تجمع كل التحاليل "البسيطة" لنفس الزيارة (أي تحليل
# ماله تقرير مخصص كبير بـ REPORT_TEMPLATE_MAP، يعني كل شي غير CBC/Blood
# Film/WBC Differential/Fluid exam/Coagulation) بورقة A4 وحدة، مرتّبة
# حسب القسم (Department) اللي حدده الأدمن لكل تحليل بكتالوج التحاليل —
# مثلاً كل تحاليل قسم "Biochemical Test" تحت عنوان وحد، وتحتها "Virology
# screen"، وتحتها "Hormones"، وتحتها "Vitamins"، وهكذا لأي قسم إضافي،
# بنفس ترويسة الشعار والأطباء المشتركة (reports/base_report.html) بدون
# أي تكرار لها. تحليل بدون أي نتيجة مُدخلة بعد يُستبعد من اللوحة تلقائيًا
# حتى ما تطلع أسطر فاضية.
# ------------------------------------------------------------------------
@app.route("/front-desk/visits/<int:visit_id>/print/combined-panel")
@login_required
def print_combined_panel(visit_id):
    db = get_db()
    visit = db.execute(
        "SELECT v.*, p.full_name as patient_name, p.age, p.age_unit, p.gender, "
        "d.full_name as referring_doctor_name, rc.name as referral_center_name "
        "FROM visits v JOIN patients p ON p.id = v.patient_id "
        "LEFT JOIN doctors d ON d.id = v.doctor_id "
        "LEFT JOIN referral_centers rc ON rc.id = v.referral_center_id WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        return "Not found", 404

    order_tests = db.execute(
        "SELECT ot.*, td.code as test_code, td.name as test_name, td.department as department, "
        "td.report_group as report_group "
        "FROM order_tests ot JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id "
        "WHERE o.visit_id=? AND ot.status IN ('Completed', 'Verified') "
        "ORDER BY COALESCE(NULLIF(TRIM(td.report_group), ''), td.department), td.name",
        (visit_id,),
    ).fetchall()

    groups_by_dept = {}
    dept_order = []
    for ot in order_tests:
        if ot["test_code"] in REPORT_TEMPLATE_MAP:
            continue  # لهذا التحليل تقريره الكبير الخاص، ما يدخل باللوحة المجمّعة
        params = db.execute(
            "SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id", (ot["test_definition_id"],)
        ).fetchall()
        if not params:
            continue
        results = db.execute(
            "SELECT r.*, tp.name as pname FROM results r "
            "JOIN test_parameters tp ON tp.id = r.test_parameter_id WHERE order_test_id=?",
            (ot["id"],),
        ).fetchall()
        results_by_name = {r["pname"]: r for r in results}
        multi_param = len(params) > 1
        rows = []
        # نفس منطق التقرير الفردي (print_report/find_previous_reference):
        # يجيب آخر زيارة سابقة لنفس هذا التحليل مطابقة بالاسم الثلاثي +
        # العمر + الجنس، وتُضاف قيمها لكل صف فقط إذا كان قسم التحليل من
        # الأقسام المفعّلة بـ PREVIOUS_VALUE_DEPARTMENT_KEYWORDS. تُستدعى
        # مرة وحدة لكل تحليل (مو لكل باراميتر) لتفادي تكرار الاستعلام.
        previous_visit_date, previous_values = find_previous_reference(
            db, visit["patient_name"], visit["age"], visit["gender"],
            ot["test_definition_id"], ot["department"], visit_id, visit["created_at"],
        )
        for p in params:
            r = results_by_name.get(p["name"])
            if r is None:
                continue
            value = r["value_text"] if r["value_text"] not in (None, "") else r["value_numeric"]
            if value in (None, ""):
                continue
            rng = find_reference_range(db, p["id"], visit["gender"], visit["age"], visit["age_unit"])
            rows.append({
                "name": p["name"] if multi_param else ot["test_name"],
                "result": value,
                "unit": p["unit"] or "",
                "low": rng["low"] if rng else None,
                "high": rng["high"] if rng else None,
                "range_text": rng["range_text"] if rng else None,
                "previous_value": previous_values.get(p["name"]),
                "previous_date": previous_visit_date,
                # لون مخصص + فاصل صفحة اختياريان — الأولوية دائماً للباراميتر
                # المفرد (test_parameters.panel_color/panel_page_break، يُضبط
                # من "مصمم التقارير" لكل باراميتر لحاله)؛ لو غير مضبوط له
                # تحديداً، يرجع لإعداد التحليل كامل (test_definitions) كسلوك
                # احتياطي قديم — هذا يسمح بتلوين باراميتر وحدة بس (زي NRBC)
                # داخل تحليل متعدد الباراميترات (زي Blood film) بدون ما
                # يلوّن باقي باراميتراته.
                "color": p["panel_color"] or ot["panel_color"] or None,
                "_own_page_break": bool(p["panel_page_break"]),
            })
        if not rows:
            continue
        # فاصل الصفحة: كل صف يحمل فاصله المفرد (لو مضبوط لباراميتره تحديداً)،
        # وفاصل التحليل كامل (لو مفعّل) يُطبَّق فقط على أول صف من صفوفه.
        for row in rows:
            if row.pop("_own_page_break", False):
                row["page_break_before"] = True
        if ot["panel_page_break"]:
            rows[0]["page_break_before"] = True
        # الأولوية دائماً لاسم "الريبورت المجمّع" المخصص (report_group) إذا
        # الأدمن حدده لهذا التحليل من كتالوج التحاليل — وإلا نرجع لاسم
        # القسم (department) القديم كما كان الوضع قبل هذي الميزة.
        dept = (ot["report_group"] or "").strip() or ot["department"] or "Other"
        if dept not in groups_by_dept:
            groups_by_dept[dept] = []
            dept_order.append(dept)
        groups_by_dept[dept].extend(rows)

    # ترتيب طباعة أقسام اللوحة المجمّعة — افتراضيًا أبجدي (كما هو بالاستعلام
    # فوق)؛ لو الأدمن حدد ترتيبًا مخصصًا من الإعدادات (combined_panel_group_order)
    # نقدّم الأقسام المذكورة هناك بنفس ترتيبها، وأي قسم غير مذكور يبقى بعدها
    # بترتيبه الأبجدي الأصلي دون أي تغيير.
    group_order_raw = get_setting(db, "combined_panel_group_order", "")
    preferred_order = [ln.strip() for ln in group_order_raw.splitlines() if ln.strip()]
    if preferred_order:
        remaining = [d for d in dept_order if d not in preferred_order]
        dept_order = [d for d in preferred_order if d in dept_order] + remaining

    panel_groups = [{"department": d, "rows": groups_by_dept[d]} for d in dept_order]

    logo_path = get_setting(db, "logo_path", "")
    logo_url = url_for("static", filename=logo_path) if logo_path else None
    from_other_lab = bool(visit["referral_center_id"]) and visit["referral_center_name"] not in (None, "", "Walk-in")

    try:
        dt = datetime.fromisoformat(visit["created_at"])
        visit_date = f"{dt.day}/{dt.month}/{dt.year}"
    except (TypeError, ValueError):
        visit_date = visit["created_at"] or ""

    age_unit_abbr = {"Hours": "H", "Days": "D", "Weeks": "W", "Months": "M", "Years": "Y"}
    age_display = f"{visit['age']}{age_unit_abbr.get(visit['age_unit'] or 'Years', 'Y')}" if visit["age"] not in (None, "") else ""

    return render_template(
        "reports/combined_panel.html",
        panel_groups=panel_groups, logo_url=logo_url, from_other_lab=from_other_lab,
        visit_date=visit_date, sex=visit["gender"] or "", age=age_display,
        patient_name=visit["patient_name"], patient_id=visit["registration_number"],
        referring_doctor_name=visit["referring_doctor_name"] or "",
        show_exam_signature=False,
    )


@app.route("/reports/print/<int:order_test_id>")
@login_required
def print_report(order_test_id):
    db = get_db()
    ot = db.execute(
        "SELECT ot.*, td.code as test_code, td.name as test_name, td.department as test_department, "
        "td.is_examining_test as is_examining_test, "
        "p.id as patient_id, p.full_name as patient_name, p.gender, p.age, p.age_unit, "
        "v.id as visit_id, v.created_at as visit_created_at, v.registration_number, "
        "v.doctor_id, v.referral_center_id, "
        "d.full_name as referring_doctor_name, rc.name as referral_center_name "
        "FROM order_tests ot "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id "
        "JOIN visits v ON v.id = o.visit_id "
        "JOIN patients p ON p.id = v.patient_id "
        "LEFT JOIN doctors d ON d.id = v.doctor_id "
        "LEFT JOIN referral_centers rc ON rc.id = v.referral_center_id "
        "WHERE ot.id = ?",
        (order_test_id,),
    ).fetchone()
    if not ot:
        return "Not found", 404

    template_name = REPORT_TEMPLATE_MAP.get(ot["test_code"])

    sibling_ot = None
    if ot["test_code"] in BF_RETIC_LINK:
        sibling_code = BF_RETIC_LINK[ot["test_code"]]
        sibling_ot = db.execute(
            "SELECT ot2.*, td2.code as test_code FROM order_tests ot2 "
            "JOIN test_definitions td2 ON td2.id = ot2.test_definition_id "
            "WHERE ot2.order_id = ? AND td2.code = ?",
            (ot["order_id"], sibling_code),
        ).fetchone()
        if sibling_ot:
            template_name = REPORT_TEMPLATE_MAP["BF"]

    custom_template = None
    if not template_name:
        custom_template = get_report_template(db, ot["test_definition_id"])
        if not custom_template:
            if session.get("role") == "admin":
                flash("لا يوجد تصميم طباعة لهذا التحليل بعد. صممه من: الإدارة ← مصمم التقارير.")
                return redirect(url_for("report_designer", test_definition_id=ot["test_definition_id"]))
            flash("No printable report layout is defined for this test yet.")
            return redirect(url_for("orders_list"))
        template_name = "reports/custom.html"

    # When Blood Film and Retic Count are both ordered for this visit, pull
    # parameters/results from BOTH order_tests so the one combined report
    # has everything — regardless of which of the two fields (Retic count,
    # Corrected Retic count) happen to live under which test definition.
    test_def_ids = [ot["test_definition_id"]]
    order_test_ids = [order_test_id]
    if sibling_ot:
        test_def_ids.append(sibling_ot["test_definition_id"])
        order_test_ids.append(sibling_ot["id"])

    parameters = []
    seen_param_names = set()
    for tdid in test_def_ids:
        for p in db.execute("SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id", (tdid,)).fetchall():
            if p["name"] not in seen_param_names:
                parameters.append(p)
                seen_param_names.add(p["name"])

    results = []
    for otid in order_test_ids:
        results.extend(db.execute(
            "SELECT r.*, tp.name as param_name FROM results r "
            "JOIN test_parameters tp ON tp.id = r.test_parameter_id WHERE order_test_id=?",
            (otid,),
        ).fetchall())
    results_by_name = {r["param_name"]: r for r in results}

    params = {}
    ranges = {}
    # units: يُبنى بنفس حلقة params/ranges — يغذّي القوالب الستة الكبيرة
    # (Blood Film, WBC Diff, Coagulation, Fluid exam, Retic) بوحدة كل
    # باراميتر من قاعدة البيانات (test_parameters.unit) بدل ما تكون نصاً
    # ثابتاً مكتوباً داخل القالب نفسه — فتعديلها من صفحة "📏 تعديل
    # الوحدات" ينعكس فعلياً على الورقة المطبوعة لهذي التقارير أيضاً، تماماً
    # متل ما يصير أصلاً مع custom.html وcbc.html.
    units = {}
    for p in parameters:
        r = results_by_name.get(p["name"])
        if r is not None:
            value = r["value_text"] if r["value_text"] not in (None, "") else r["value_numeric"]
        else:
            value = None
        params[p["name"]] = "" if value is None else value
        ranges[p["name"]] = find_reference_range(db, p["id"], ot["gender"], ot["age"], ot["age_unit"])
        units[p["name"]] = p["unit"] or ""

    logo_path = get_setting(db, "logo_path", "")
    logo_url = url_for("static", filename=logo_path) if logo_path else None

    # Only show the "received from another lab" note when the visit was
    # actually referred in via a referral center, not a plain walk-in.
    from_other_lab = bool(ot["referral_center_id"]) and ot["referral_center_name"] not in (None, "", "Walk-in")

    font_size = 14 if ot["test_code"] == "CBC" else 16
    # مربع "توقيع وختم الدكتور الفاحص" مُلغى بكل التقارير حسب الطلب — الشرط
    # تعطّل عمدًا بدل حذف EXAM_SIGNATURE_TEST_CODES نفسها، فلو احتجتوها
    # ترجع بالمستقبل تكفي إعادة هذا السطر لحالته: ot["test_code"] in
    # EXAM_SIGNATURE_TEST_CODES or bool(ot["is_examining_test"])
    show_exam_signature = False

    try:
        dt = datetime.fromisoformat(ot["visit_created_at"])
        visit_date = f"{dt.day}/{dt.month}/{dt.year}"
    except (TypeError, ValueError):
        visit_date = ot["visit_created_at"] or ""

    # زيارة سابقة لنفس الشخص (تطابق كامل بالاسم + العمر + الجنس) لنفس هذا
    # التحليل: تُذكر دائمًا كتاريخ فقط أسفل التقرير، وتُذكر مع القيم السابقة
    # لكل باراميتر فقط إذا كان قسم التحليل كيمياء أو تخثر — وهذا الافتراضي
    # يمديك تتجاوزه لحظة الطباعة بمربع "إظهار النتيجة السابقة" بشاشة إدخال
    # النتائج (؟show_prev=1/0)، لأن عرضها بالتقرير المطبوع اختياري دائماً
    # وليس شرطاً حتى لو القسم كيمياء/تخثر.
    show_prev_values = department_shows_previous_values(ot["test_department"])
    show_prev_override = request.args.get("show_prev")
    if show_prev_override in ("0", "1"):
        show_prev_values = show_prev_override == "1"
    previous_visit_date, previous_values = find_previous_reference(
        db, ot["patient_name"], ot["age"], ot["gender"], ot["test_definition_id"],
        ot["test_department"], ot["visit_id"], ot["visit_created_at"],
    )

    cbc_groups = None
    if ot["test_code"] == "CBC":
        units_by_name = {p["name"]: p["unit"] for p in parameters}
        highlight_by_name = {p["name"]: bool(p["highlight"]) for p in parameters}
        cbc_groups = []
        for group in CBC_ROW_GROUPS:
            rows = []
            for name in group:
                rng = ranges.get(name)
                rows.append({
                    "name": name,
                    "result": params.get(name, ""),
                    "unit": units_by_name.get(name, ""),
                    "low": rng["low"] if rng else "",
                    "high": rng["high"] if rng else "",
                    "highlight": highlight_by_name.get(name, False),
                    "previous": previous_values.get(name, "") if show_prev_values else "",
                })
            cbc_groups.append(rows)

    custom_rows = None
    custom_heading = None
    custom_heading_align = "center"
    custom_rows_align = "right"
    custom_unit_column = False
    if custom_template:
        units_by_name = {p["name"]: p["unit"] for p in parameters}
        unit2_by_name = {p["name"]: p["unit2"] for p in parameters}
        unit2_factor_by_name = {p["name"]: p["unit2_factor"] for p in parameters}
        custom_heading = custom_template["heading"] or ot["test_name"]
        custom_heading_align = custom_template["heading_align"] or "center"
        custom_rows_align = custom_template["rows_align"] or "right"
        custom_unit_column = bool(custom_template["unit_column"])
        row_defs = json.loads(custom_template["rows_json"] or "[]")
        custom_rows = []
        for rd in row_defs:
            pname = rd.get("param_name", "")
            rng = ranges.get(pname)
            normal_range = ""
            if rng:
                if rng["range_text"]:
                    normal_range = rng["range_text"]
                elif rng["low"] is not None and rng["high"] is not None:
                    normal_range = f"{rng['low']} - {rng['high']}"
                elif rng["low"] is not None:
                    normal_range = f"> {rng['low']}"
                elif rng["high"] is not None:
                    normal_range = f"< {rng['high']}"
            result_val = params.get(pname, "")

            # حقل "This test done by ... (FDA Approved)" (source_note) أُلغي
            # عرضه نهائيًا بالتقرير المطبوع بناءً على الطلب — يبقى العمود
            # موجود بجدول reference_ranges (ما يُحذف) لكن لا يُقرأ ولا يُمرَّر
            # للقالب أبداً بعد الآن، فلا يظهر إطلاقًا مهما كانت قيمته بقاعدة
            # البيانات.

            # مستويات Normal Range المفصّلة (زي Triglycerides بالنموذج
            # المرجعي) — من range_text لو موجود، وإلا من normal_range
            # الجاهزة (سطر واحد بلا تسمية، زي أغلب التحاليل العادية).
            if rng and rng["range_text"]:
                range_tiers = parse_range_tiers(rng["range_text"])
            elif normal_range:
                range_tiers = [{"label": None, "value": normal_range}]
            else:
                range_tiers = []

            # مدى الوحدة الثانية (unit2) — يُحسب تلقائيًا بنفس معامل تحويل
            # النتيجة (unit2_factor) على low/high، فقط لمدى بسيط (سطر واحد
            # رقمي بلا مستويات)؛ لا ينطبق على مدى متدرّج (range_tiers متعدد).
            normal_range2 = None
            if rng and not (rng["range_text"] and len(range_tiers) > 1) and rng["low"] is not None and rng["high"] is not None:
                factor = unit2_factor_by_name.get(pname)
                low2 = format_unit2_value(rng["low"], factor)
                high2 = format_unit2_value(rng["high"], factor)
                if low2 is not None and high2 is not None:
                    normal_range2 = f"{low2} - {high2}"

            custom_rows.append({
                "label": rd.get("label") or pname,
                "result": result_val,
                "unit": units_by_name.get(pname, ""),
                "normal_range": normal_range,
                "range_tiers": range_tiers,
                "normal_range2": normal_range2,
                "previous": previous_values.get(pname, "") if show_prev_values else "",
                "result2": format_unit2_value(result_val, unit2_factor_by_name.get(pname)),
                "unit2": unit2_by_name.get(pname) or "",
                # حرية تحريك موضع الحقول لكل صف على حدة (يُضبط من مصمم
                # التقرير مستقبلاً، يُقرأ من rows_json): name_side='right'
                # ينقل اسم التحليل فيزيائيًا لعمود القيمة بدل عمود الاسم،
                # range_position='below' يرجع المدى الطبيعي لسطر مستقل تحت
                # اسم التحليل بدل عمود محاذي لعمود النتيجة (الافتراضي).
                "name_side": rd.get("name_side") or "left",
                "range_position": rd.get("range_position") or "inline",
                "name_align": rd.get("name_align"),
                "color": rd.get("color"),
                "page_break_before": rd.get("page_break_before", False),
            })

    age_unit_abbr = {"Hours": "H", "Days": "D", "Weeks": "W", "Months": "M", "Years": "Y"}
    age_display = f"{ot['age']}{age_unit_abbr.get(ot['age_unit'] or 'Years', 'Y')}" if ot["age"] not in (None, "") else ""

    # ترويسة متكرّرة تلقائيًا بأعلى كل صفحة إضافية عند الطباعة (نفس الشعار،
    # أسماء الأطباء، الدكتور المرسل، اسم المريض، التاريخ...) — فقط لتقارير
    # الكيمياء/الهرمونات/الفيتامينات/الدلائل الورمية/التخثر/الفايروسات (اللي
    # تُطبع كبطاقات وقد تطول لأكثر من صفحة)، وليس لتحاليل أمراض الدم إطلاقًا
    # (CBC وغيره من التقارير الجاهزة أبدًا ما يُمرَّر لها هذا المتغيّر، فتبقى
    # بسلوكها القديم تمامًا). راجع base_report.html لآلية التكرار الفعلية.
    repeat_header_on_print = department_shows_previous_values(ot["test_department"])

    return render_template(
        template_name,
        ot=ot, params=params, ranges=ranges, units=units, cbc_groups=cbc_groups,
        custom_rows=custom_rows, custom_heading=custom_heading,
        custom_heading_align=custom_heading_align, custom_rows_align=custom_rows_align,
        custom_unit_column=custom_unit_column,
        show_prev_values=show_prev_values, previous_visit_date=previous_visit_date,
        previous_values=previous_values, repeat_header_on_print=repeat_header_on_print,
        logo_url=logo_url, from_other_lab=from_other_lab, font_size=font_size,
        show_exam_signature=show_exam_signature,
        visit_date=visit_date, sex=ot["gender"] or "", age=age_display,
        patient_name=ot["patient_name"], patient_id=ot["registration_number"],
        referring_doctor_name=ot["referring_doctor_name"] or "",
    )


def _whatsapp_flush_pdf_dir():
    d = os.path.join(app.static_folder, "whatsapp_pdfs")
    os.makedirs(d, exist_ok=True)
    return d


def _whatsapp_generate_and_queue(db, visit_row, patient_id, patient_name, phone,
                                  label, html_content, order_test_id=None):
    """يولّد PDF من الـHTML الجاهز (نفس صفحة الطباعة)، يسجّل صف بطابور
    whatsapp_sends، ويحاول الإرسال فورًا إذا الإنترنت متوفر؛ وإلا يبقى الصف
    pending وتاخذه المهمة الخلفية لاحقًا أو زر "إعادة المحاولة" اليدوي."""
    import pdf_export
    import whatsapp_bridge

    now = datetime.now().isoformat(timespec="seconds")
    pdf_path = pdf_export.make_temp_pdf_path(f"visit{visit_row['id']}")

    # توليد PDF نفسه كان بدون أي حماية — أي خطأ يصير أثناء التحويل (خط
    # ناقص، مسار شعار كسران، أي عطل بمكتبة التحويل) كان يطيح الطلب كامل
    # بصفحة 500 بيضاء بدون أي رسالة توضح شنو صار، ويقفل الصفحة على
    # المستخدم بدل ما يعطيه رسالة مفهومة أو يحفظ المحاولة بالطابور. الآن
    # أي فشل هنا يُسجَّل كصف "failed" بنفس الطابور (تكدر تعيد المحاولة من
    # "طابور واتساب" لاحقاً) بدل ما يكسر الصفحة بالكامل.
    try:
        pdf_export.html_to_pdf(html_content, request.url_root, pdf_path)
    except Exception as exc:
        cur = db.execute(
            "INSERT INTO whatsapp_sends (visit_id, order_test_id, patient_id, patient_name, "
            "phone, label, pdf_path, status, attempts, error, requested_by, created_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, 'failed', 1, ?, ?, ?)",
            (visit_row["id"], order_test_id, patient_id, patient_name, phone, label,
             pdf_path, f"فشل توليد PDF: {exc}", session.get("user_id"), now),
        )
        db.commit()
        log_action("WhatsAppSend", "order_test", order_test_id or visit_row["id"], f"PDF generation failed: {exc}")
        return False, f"فشل توليد ملف PDF: {exc}"

    cur = db.execute(
        "INSERT INTO whatsapp_sends (visit_id, order_test_id, patient_id, patient_name, "
        "phone, label, pdf_path, status, attempts, requested_by, created_at) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, 'pending', 0, ?, ?)",
        (visit_row["id"], order_test_id, patient_id, patient_name, phone, label,
         pdf_path, session.get("user_id"), now),
    )
    db.commit()
    send_id = cur.lastrowid

    try:
        whatsapp_bridge.send_pdf(
            phone, pdf_path,
            caption=f"نتيجة {label} — {patient_name}",
            country_code=get_setting(db, "whatsapp_country_code", "964"),
        )
        db.execute(
            "UPDATE whatsapp_sends SET status='sent', sent_at=?, attempts=attempts+1 WHERE id=?",
            (datetime.now().isoformat(timespec="seconds"), send_id),
        )
        db.commit()
        return True, None
    except whatsapp_bridge.WhatsAppSendError as exc:
        db.execute(
            "UPDATE whatsapp_sends SET status='failed', error=?, attempts=attempts+1 WHERE id=?",
            (str(exc), send_id),
        )
        db.commit()
        return False, str(exc)


def _maybe_auto_whatsapp_send(db, visit_id):
    """إرسال تلقائي عبر واتساب لكل نتائج الزيارة دفعة واحدة (بملف PDF واحد،
    نفس أسلوب زر 'إرسال كل النتائج' اليدوي) — بس إذا تحققت الشروط الثلاثة:
    (1) رقم هاتف المريض مسجّل بملفه، (2) كل تحليل مطلوب بهذي الزيارة صار
    Completed أو Verified (ولا تحليل واحد لسا ناقص نتيجة)، و(3) ما سبق
    إرسال (تلقائي أو يدوي) لكل نتائج هذي الزيارة مجموعة من قبل — تفاديًا
    لتكرار الإرسال في كل مرة تُعدَّل فيها نتيجة بعد اكتمال الزيارة.
    تُستدعى بعد أي حفظ نتيجة (مفردة من result_entry، أو مُجمّعة من
    visit_results_entry). أي خطأ بتوليد الـPDF أو الإرسال لا يوقف حفظ
    النتائج أبدًا — يُسجَّل بطابور واتساب كـ pending/failed مثل أي محاولة
    إرسال يدوية عادية، وتلتقطه المهمة الخلفية أو زر 'إعادة المحاولة' لاحقًا."""
    visit = db.execute(
        "SELECT v.*, p.id as patient_id, p.full_name, p.phone FROM visits v "
        "JOIN patients p ON p.id = v.patient_id WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit or not (visit["phone"] or "").strip():
        return False

    statuses = [row["status"] for row in db.execute(
        "SELECT ot.status FROM order_tests ot JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=?",
        (visit_id,),
    ).fetchall()]
    if not statuses or any(s not in ("Completed", "Verified") for s in statuses):
        return False  # لسا فيه تحليل ناقص نتيجة — ما نرسل شي بعد

    already = db.execute(
        "SELECT id FROM whatsapp_sends WHERE visit_id=? AND order_test_id IS NULL "
        "AND status IN ('pending', 'sent')",
        (visit_id,),
    ).fetchone()
    if already:
        return False  # سبق إرسال/طابور كل نتائج هذي الزيارة مجموعة من قبل

    try:
        html_content = print_visit_results(visit_id)
    except Exception:
        return False
    if not isinstance(html_content, str):
        return False  # print_visit_results رجّع redirect/خطأ بدل HTML — نتجاهل بصمت

    ok, error = _whatsapp_generate_and_queue(
        db, visit, visit["patient_id"], visit["full_name"], visit["phone"],
        "كل نتائج الزيارة (إرسال تلقائي)", html_content, order_test_id=None,
    )
    log_action("WhatsAppAutoSend", "visit", visit_id, "OK" if ok else f"queued: {error}")
    if ok:
        flash("📱 النتائج مكتملة — تم إرسالها تلقائيًا عبر واتساب للمريض.")
    else:
        flash(f"📱 النتائج مكتملة لكن تعذّر الإرسال الفوري — انضافت لطابور واتساب وسترسل تلقائيًا عند توفر الإنترنت. ({error})")
    return True


def _get_pdf_archive_dir(db):
    """مجلد أرشفة الـPDF الدائم كما ظبطه المدير من Management → Settings
    (settings.pdf_archive_dir). يرجّع None إذا لسا فاضي (غير مُعد بعد) —
    وبهذي الحالة الأرشفة التلقائية تُتجاهل بصمت بدل ما تكسر حفظ النتائج."""
    d = (get_setting(db, "pdf_archive_dir", "") or "").strip()
    if not d:
        return None
    try:
        os.makedirs(d, exist_ok=True)
    except OSError:
        return None
    return d


def _maybe_archive_visit_pdf(db, visit_id):
    """يحفظ نسخة PDF دائمة (موحّدة لكل نتائج الزيارة) بمجلد الأرشيف الثابت،
    بس إذا: (1) المدير ظبط مجلد أرشفة من الإعدادات، و(2) كل تحليل مطلوب
    بهذي الزيارة صار Completed أو Verified. تُستدعى من نفس نقطتي حفظ
    النتائج اللي تستدعي _maybe_auto_whatsapp_send (مفردة ومُجمّعة)، وتُعيد
    توليد/استبدال نفس الملف (upsert بـ visit_id) في كل مرة — حتى يبقى
    الملف المؤرشف مطابقًا دائمًا لآخر تعديل على النتائج، حتى لو تم تعديلها
    بعد الاكتمال. أي خطأ بالتوليد لا يوقف حفظ النتائج أبدًا."""
    archive_dir = _get_pdf_archive_dir(db)
    if not archive_dir:
        return False

    visit = db.execute(
        "SELECT v.*, p.id as patient_id, p.full_name, p.registration_number as _unused, "
        "d.full_name as referring_doctor_name, rc.name as referral_center_name "
        "FROM visits v JOIN patients p ON p.id = v.patient_id "
        "LEFT JOIN doctors d ON d.id = v.doctor_id "
        "LEFT JOIN referral_centers rc ON rc.id = v.referral_center_id "
        "WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        return False

    statuses = [row["status"] for row in db.execute(
        "SELECT ot.status FROM order_tests ot JOIN orders o ON o.id = ot.order_id WHERE o.visit_id=?",
        (visit_id,),
    ).fetchall()]
    if not statuses or any(s not in ("Completed", "Verified") for s in statuses):
        return False  # لسا فيه تحليل ناقص نتيجة — ما نؤرشف بعد

    try:
        html_content = print_visit_results(visit_id)
    except Exception:
        return False
    if not isinstance(html_content, str):
        return False

    import pdf_export
    try:
        # اسم ملف واضح للبحث اليدوي بمجلد الأرشيف نفسه من خارج البرنامج:
        # رقم_التسجيل - اسم المريض - تاريخ الزيارة.pdf (كل رمز غير آمن
        # بالاسم يُستبدل بشرطة سفلية).
        safe_name = secure_filename(visit["full_name"] or "patient") or "patient"
        try:
            dt = datetime.fromisoformat(visit["created_at"])
            date_part = dt.strftime("%Y-%m-%d")
        except (TypeError, ValueError):
            date_part = datetime.now().strftime("%Y-%m-%d")
        filename = f"{visit['registration_number']}_{safe_name}_{date_part}.pdf"
        pdf_path = os.path.join(archive_dir, filename)
        pdf_export.html_to_pdf(html_content, request.url_root, pdf_path)
    except Exception:
        return False

    save_saved_report(
        db, visit_id, visit["patient_id"], visit["full_name"], visit["registration_number"],
        pdf_path, referring_doctor_name=visit["referring_doctor_name"],
        referral_center_name=visit["referral_center_name"],
    )
    return True


@app.route("/whatsapp/send-result/<int:order_test_id>", methods=["POST"])
@login_required
def whatsapp_send_result(order_test_id):
    """إرسال نتيجة تحليل واحد محدّد (مو كل نتائج الزيارة) عبر واتساب —
    نفس زر '🖨 طباعة التقرير' لكن يولّد PDF ويرسله بدل ما يفتحه بالمتصفح."""
    db = get_db()
    ot = db.execute(
        "SELECT ot.*, td.name as test_name, v.id as visit_id, "
        "p.id as patient_id, p.full_name as patient_name, p.phone "
        "FROM order_tests ot "
        "JOIN test_definitions td ON td.id = ot.test_definition_id "
        "JOIN orders o ON o.id = ot.order_id "
        "JOIN visits v ON v.id = o.visit_id "
        "JOIN patients p ON p.id = v.patient_id WHERE ot.id=?",
        (order_test_id,),
    ).fetchone()
    if not ot:
        flash("النتيجة غير موجودة.")
        return redirect(url_for("results_list"))
    if not ot["phone"]:
        flash("لا يوجد رقم هاتف مسجّل لهذا المريض — أضِفه من ملف المريض أولًا.")
        return redirect(url_for("visit_results_entry", visit_id=ot["visit_id"]))

    html_content = print_report(order_test_id)
    if not isinstance(html_content, str):
        # print_report يرجّع redirect (تنبيه) لو ما فيه تصميم طباعة لهذا
        # التحليل بعد — نفس الحالة تنطبق هنا فنعيد نفس التوجيه.
        return html_content

    ok, error = _whatsapp_generate_and_queue(
        db, ot, ot["patient_id"], ot["patient_name"], ot["phone"],
        ot["test_name"], html_content, order_test_id=order_test_id,
    )
    if ok:
        flash(f"✅ تم إرسال نتيجة {ot['test_name']} عبر واتساب للمريض.")
    else:
        flash(f"⏳ تعذّر الإرسال الآن، تمت إضافتها للطابور وستُرسل تلقائيًا عند توفر الإنترنت. ({error})")
    log_action("WhatsAppSend", "order_test", order_test_id, "OK" if ok else f"queued: {error}")
    return redirect(url_for("visit_results_entry", visit_id=ot["visit_id"]))


@app.route("/whatsapp/send-visit/<int:visit_id>", methods=["POST"])
@login_required
def whatsapp_send_visit(visit_id):
    """إرسال كل نتائج الزيارة مجموعة بملف PDF واحد (الجدول الموحّد) عبر
    واتساب — يستخدم نفس صفحة print_visit_results."""
    db = get_db()
    visit = db.execute(
        "SELECT v.*, p.id as patient_id, p.full_name, p.phone FROM visits v "
        "JOIN patients p ON p.id = v.patient_id WHERE v.id=?",
        (visit_id,),
    ).fetchone()
    if not visit:
        flash("الزيارة غير موجودة.")
        return redirect(url_for("results_list"))
    if not visit["phone"]:
        flash("لا يوجد رقم هاتف مسجّل لهذا المريض — أضِفه من ملف المريض أولًا.")
        return redirect(url_for("visit_results_entry", visit_id=visit_id))

    html_content = print_visit_results(visit_id)
    if not isinstance(html_content, str):
        return html_content

    ok, error = _whatsapp_generate_and_queue(
        db, visit, visit["patient_id"], visit["full_name"], visit["phone"],
        "كل نتائج الزيارة", html_content, order_test_id=None,
    )
    if ok:
        flash("✅ تم إرسال كل نتائج الزيارة عبر واتساب للمريض بملف واحد.")
    else:
        flash(f"⏳ تعذّر الإرسال الآن، تمت إضافتها للطابور وستُرسل تلقائيًا عند توفر الإنترنت. ({error})")
    log_action("WhatsAppSend", "visit", visit_id, "OK" if ok else f"queued: {error}")
    return redirect(url_for("visit_results_entry", visit_id=visit_id))


@app.route("/reports/archive")
@login_required
def pdf_archive():
    """بحث سريع بأرشيف الـPDF الدائم (باسم المريض أو رقم التسجيل)، من داخل
    البرنامج مباشرة — بدل ما يفتّش المستخدم يدويًا بمجلد الأرشيف بالحاسبة."""
    db = get_db()
    q = request.args.get("q", "").strip()
    archive_dir = _get_pdf_archive_dir(db)
    rows = search_saved_reports(db, q) if archive_dir else []
    return render_template("pdf_archive.html", rows=rows, q=q, archive_configured=bool(archive_dir))


@app.route("/reports/archive/<int:visit_id>/open")
@login_required
def pdf_archive_open(visit_id):
    """يفتح/يحمّل نسخة الأرشيف الدائمة مباشرة (نفس ملف مجلد الأرشفة بالضبط،
    وليس توليدًا جديدًا) — لو الملف انمسح يدويًا من مجلد الأرشيف من خارج
    البرنامج (نقل/حذف)، نرجّع رسالة واضحة بدل خطأ سيرفر غامض."""
    db = get_db()
    row = get_saved_report(db, visit_id)
    if not row or not row["pdf_path"] or not os.path.exists(row["pdf_path"]):
        flash("الملف غير موجود بمجلد الأرشيف — يمكن انحذف أو انقل يدويًا من خارج البرنامج.")
        return redirect(url_for("pdf_archive"))
    return send_file(row["pdf_path"], as_attachment=False,
                      download_name=os.path.basename(row["pdf_path"]))


@app.route("/whatsapp/queue")
@roles_required("supervisor")
def whatsapp_queue():
    db = get_db()
    rows = db.execute(
        "SELECT * FROM whatsapp_sends ORDER BY id DESC LIMIT 200"
    ).fetchall()
    return render_template("whatsapp_queue.html", rows=rows)


@app.route("/whatsapp/queue/<int:send_id>/retry", methods=["POST"])
@roles_required("supervisor")
def whatsapp_retry(send_id):
    import whatsapp_bridge

    db = get_db()
    row = db.execute("SELECT * FROM whatsapp_sends WHERE id=?", (send_id,)).fetchone()
    if not row:
        flash("العنصر غير موجود.")
        return redirect(url_for("whatsapp_queue"))
    try:
        whatsapp_bridge.send_pdf(
            row["phone"], row["pdf_path"],
            caption=f"نتيجة {row['label']} — {row['patient_name']}",
            country_code=get_setting(db, "whatsapp_country_code", "964"),
        )
        db.execute(
            "UPDATE whatsapp_sends SET status='sent', sent_at=?, attempts=attempts+1, error=NULL WHERE id=?",
            (datetime.now().isoformat(timespec="seconds"), send_id),
        )
        db.commit()
        flash("✅ تم الإرسال بنجاح.")
    except whatsapp_bridge.WhatsAppSendError as exc:
        db.execute(
            "UPDATE whatsapp_sends SET status='failed', error=?, attempts=attempts+1 WHERE id=?",
            (str(exc), send_id),
        )
        db.commit()
        flash(f"❌ فشلت المحاولة: {exc}")
    return redirect(url_for("whatsapp_queue"))


def whatsapp_background_worker():
    """تعمل بخيط منفصل طول ما البرنامج شغال: كل 3 دقائق تفحص إذا فيه اتصال
    إنترنت، وإذا موجود تحاول ترسل أي عناصر pending أو failed بالطابور —
    هذا اللي يخلّي 'البرنامج محلي بس يرسل تلقائيًا لما يوصله نت' ممكن."""
    import time as _time
    import whatsapp_bridge

    while True:
        _time.sleep(180)
        try:
            if not whatsapp_bridge.has_internet():
                continue
            with app.app_context():
                db = get_db()
                pending = db.execute(
                    "SELECT * FROM whatsapp_sends WHERE status IN ('pending','failed') "
                    "AND attempts < 5 ORDER BY id ASC LIMIT 10"
                ).fetchall()
                for row in pending:
                    try:
                        whatsapp_bridge.send_pdf(
                            row["phone"], row["pdf_path"],
                            caption=f"نتيجة {row['label']} — {row['patient_name']}",
                            country_code=get_setting(db, "whatsapp_country_code", "964"),
                        )
                        db.execute(
                            "UPDATE whatsapp_sends SET status='sent', sent_at=?, attempts=attempts+1 WHERE id=?",
                            (datetime.now().isoformat(timespec="seconds"), row["id"]),
                        )
                    except whatsapp_bridge.WhatsAppSendError as exc:
                        db.execute(
                            "UPDATE whatsapp_sends SET status='failed', error=?, attempts=attempts+1 WHERE id=?",
                            (str(exc), row["id"]),
                        )
                    db.commit()
        except Exception:
            # أي خطأ غير متوقع بالمهمة الخلفية ما يوقف السيرفر أبدًا
            continue


@app.route("/workbench/verify/<int:order_test_id>", methods=["POST"])
@roles_required("supervisor")
def verify_result(order_test_id):
    db = get_db()
    now = datetime.now().isoformat(timespec="seconds")
    db.execute("UPDATE results SET verified_by=?, verified_at=? WHERE order_test_id=?",
               (session["user_id"], now, order_test_id))
    db.execute("UPDATE order_tests SET status='Verified' WHERE id=?", (order_test_id,))
    db.commit()
    log_action("Verify", "order_test", order_test_id)
    flash("Result verified & approved.")
    next_visit_id = request.form.get("next_visit_id")
    if next_visit_id:
        return redirect(url_for("visit_results_entry", visit_id=next_visit_id))
    return redirect(url_for("orders_list"))


@app.route("/workbench/unverify/<int:order_test_id>", methods=["POST"])
@roles_required("supervisor")
def unverify_result(order_test_id):
    # Reopens a previously verified/released result for correction. Never
    # silent: it clears verified_by/verified_at (so it visibly shows as
    # un-verified again everywhere, including on any printed report) and
    # writes an audit log entry with who reopened it and why, then sends
    # the person straight back into the edit form. Re-verification is
    # required again afterward — nothing about a corrected result skips the
    # normal approval step.
    db = get_db()
    reason = request.form.get("reason", "").strip()
    db.execute("UPDATE results SET verified_by=NULL, verified_at=NULL WHERE order_test_id=?", (order_test_id,))
    db.execute("UPDATE order_tests SET status='Completed' WHERE id=?", (order_test_id,))
    db.commit()
    log_action("Unverify", "order_test", order_test_id, reason or "(no reason given)")
    flash("Result reopened for editing — re-verification will be required after saving.")
    next_visit_id = request.form.get("next_visit_id")
    if next_visit_id:
        return redirect(url_for("visit_results_entry", visit_id=next_visit_id))
    return redirect(url_for("result_entry", order_test_id=order_test_id))



# ---------------------------------------------------------------- billing --
@app.route("/billing/rates", methods=["GET", "POST"])
@roles_required("admin", "accountant")
def rates():
    db = get_db()
    if request.method == "POST":
        test_id = request.form.get("test_id")
        price = float(request.form.get("price") or 0)
        db.execute("UPDATE test_definitions SET price=? WHERE id=?", (price, test_id))
        db.commit()
        flash("Price updated.")
        return redirect(url_for("rates"))
    tests = db.execute("SELECT * FROM test_definitions ORDER BY department, name").fetchall()
    return render_template("billing/rates.html", tests=tests)


@app.route("/billing/invoices")
@login_required
def invoices_list():
    db = get_db()
    rows = db.execute(
        "SELECT i.id, i.total_amount, i.discount_amount, i.paid_amount, i.status, i.created_at, i.is_locked, "
        "p.full_name as patient_name, v.registration_number, v.id as visit_id "
        "FROM invoices i JOIN visits v ON v.id = i.visit_id JOIN patients p ON p.id = v.patient_id "
        "ORDER BY i.id DESC LIMIT 200"
    ).fetchall()
    return render_template("billing/invoices.html", rows=rows)


@app.route("/billing/invoices/<int:invoice_id>/toggle-lock", methods=["POST"])
@roles_required("admin", "accountant")
def toggle_invoice_lock(invoice_id):
    db = get_db()
    inv = db.execute("SELECT is_locked FROM invoices WHERE id=?", (invoice_id,)).fetchone()
    new_val = 0 if inv["is_locked"] else 1
    db.execute("UPDATE invoices SET is_locked=? WHERE id=?", (new_val, invoice_id))
    db.commit()
    log_action("ToggleLock", "invoice", invoice_id, str(new_val))
    flash("Invoice locked." if new_val else "Invoice reopened.")
    return redirect(url_for("invoices_list"))


@app.route("/billing/invoices/<int:invoice_id>/discount", methods=["POST"])
@roles_required("admin", "accountant")
def apply_discount(invoice_id):
    db = get_db()
    inv = db.execute("SELECT is_locked FROM invoices WHERE id=?", (invoice_id,)).fetchone()
    if inv and inv["is_locked"]:
        flash("Invoice is locked. Reopen it first to make changes.")
        return redirect(url_for("invoices_list"))
    amount = float(request.form.get("discount_amount") or 0)
    db.execute("UPDATE invoices SET discount_amount=? WHERE id=?", (amount, invoice_id))
    db.commit()
    log_action("Discount", "invoice", invoice_id, f"amount={amount}")
    flash("Discount applied.")
    return redirect(url_for("invoices_list"))



@app.route("/billing/registers")
@login_required
def registers():
    db = get_db()
    rows = db.execute(
        "SELECT u.full_name as user, "
        "COALESCE(SUM(pay.amount),0) as payments, "
        "COALESCE((SELECT SUM(discount_amount) FROM invoices WHERE created_by=u.id),0) as discounts "
        "FROM users u LEFT JOIN payments pay ON pay.user_id = u.id "
        "GROUP BY u.id ORDER BY payments DESC"
    ).fetchall()
    return render_template("billing/registers.html", rows=rows)


@app.route("/billing/pay/<int:visit_id>", methods=["POST"])
@login_required
def pay_invoice(visit_id):
    db = get_db()
    amount = float(request.form.get("amount") or 0)
    invoice = db.execute("SELECT * FROM invoices WHERE visit_id=?", (visit_id,)).fetchone()
    if invoice:
        now = datetime.now().isoformat(timespec="seconds")
        db.execute("INSERT INTO payments (invoice_id, amount, method, user_id, paid_at) "
                   "VALUES (?, ?, 'Cash', ?, ?)", (invoice["id"], amount, session["user_id"], now))
        new_paid = invoice["paid_amount"] + amount
        status = "Paid" if new_paid >= invoice["total_amount"] else "Partial"
        db.execute("UPDATE invoices SET paid_amount=?, status=? WHERE id=?",
                   (new_paid, status, invoice["id"]))
        db.commit()
        log_action("Payment", "invoice", invoice["id"], f"amount={amount}")
        flash("Payment recorded.")
    return redirect(url_for("visits_list"))


# ------------------------------------------------------- master definitions
@app.route("/master/packages", methods=["GET", "POST"])
@roles_required("supervisor")
def packages_list():
    db = get_db()
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        code = request.form.get("code", "").strip()
        test_ids = request.form.getlist("tests")
        cur = db.execute("INSERT INTO packages (name, code) VALUES (?, ?)", (name, code))
        pkg_id = cur.lastrowid
        for tid in test_ids:
            db.execute("INSERT INTO package_tests (package_id, test_definition_id) VALUES (?, ?)", (pkg_id, tid))
        db.commit()
        flash("Package created.")
        return redirect(url_for("packages_list"))

    packages = db.execute("SELECT * FROM packages ORDER BY name").fetchall()
    pkg_tests = {}
    for pkg in packages:
        tests = db.execute(
            "SELECT td.name FROM package_tests pt JOIN test_definitions td ON td.id = pt.test_definition_id "
            "WHERE pt.package_id=?", (pkg["id"],),
        ).fetchall()
        pkg_tests[pkg["id"]] = [t["name"] for t in tests]
    all_tests = db.execute("SELECT * FROM test_definitions WHERE is_active=1 ORDER BY name").fetchall()
    return render_template("master/packages.html", packages=packages, pkg_tests=pkg_tests, all_tests=all_tests)


@app.route("/master/suggestions", methods=["GET", "POST"])
@roles_required("supervisor")
def suggestions_list():
    db = get_db()
    if request.method == "POST":
        param_id = request.form.get("test_parameter_id")
        content = request.form.get("content", "").strip()
        if content:
            db.execute("INSERT INTO suggestions (test_parameter_id, content) VALUES (?, ?)", (param_id, content))
            db.commit()
            flash("Suggestion added.")
        return redirect(url_for("suggestions_list"))

    rows = db.execute(
        "SELECT s.id, s.content, tp.name as param_name, td.name as test_name FROM suggestions s "
        "JOIN test_parameters tp ON tp.id = s.test_parameter_id "
        "JOIN test_definitions td ON td.id = tp.test_definition_id ORDER BY td.name LIMIT 200"
    ).fetchall()
    parameters = db.execute(
        "SELECT tp.id, tp.name, td.name as test_name FROM test_parameters tp "
        "JOIN test_definitions td ON td.id = tp.test_definition_id ORDER BY td.name"
    ).fetchall()
    return render_template("master/suggestions.html", rows=rows, parameters=parameters)



@app.route("/master/mapcodes/bulk_add_test", methods=["POST"])
@roles_required("supervisor")
def mapcodes_bulk_add_test():
    """يضيف مرة وحدة سطر Mapcode لكل عناصر تحليل معيّن دفعة وحدة (مثلاً كل
    الـ 21 عنصر بتحليل CBC) بدل ما تضيفهم واحد واحد يدوياً من القائمة.
    الكود الافتراضي المكتوب بـ machine_code هو اسم العنصر نفسه كبداية —
    لازم بعدها تراجعه وتصححه بالكود الحقيقي اللي يرسله جهازك بالذات (شوف
    صفحة Host Interface بعد أول عينة تجريبية). العناصر النصية (Text، مثل
    Comment أو Conclusion) تُستثنى لأنه الأجهزة لا ترسل نصوص حرة عادةً،
    وأي عنصر عنده Mapcode موجود مسبقاً يُتخطّى حتى ما يتكرر."""
    db = get_db()
    test_definition_id = request.form.get("test_definition_id")
    machine_name = request.form.get("bulk_machine_name", "").strip()
    params = db.execute(
        "SELECT id, name FROM test_parameters WHERE test_definition_id=? AND result_type='Numeric' ORDER BY id",
        (test_definition_id,),
    ).fetchall()
    added = 0
    skipped = 0
    for p in params:
        exists = db.execute(
            "SELECT 1 FROM mapcodes WHERE test_parameter_id=? LIMIT 1", (p["id"],)
        ).fetchone()
        if exists:
            skipped += 1
            continue
        db.execute(
            "INSERT INTO mapcodes (test_parameter_id, machine_name, machine_code, send_enabled, receive_enabled) "
            "VALUES (?, ?, ?, 1, 1)",
            (p["id"], machine_name, p["name"]),
        )
        added += 1
    db.commit()
    flash(f"تمت إضافة {added} عنصر — راجع وصحّح الأكواد بالأكواد الحقيقية من جهازك. ({skipped} كانوا موجودين مسبقاً وتُخطّوا)" if added or skipped else "هذا التحليل ما عنده عناصر رقمية.")
    return redirect(url_for("mapcodes_list"))


@app.route("/master/mapcodes", methods=["GET", "POST"])
@roles_required("supervisor")
def mapcodes_list():
    db = get_db()
    if request.method == "POST":
        param_id = request.form.get("test_parameter_id")
        machine_name = request.form.get("machine_name", "").strip()
        machine_code = request.form.get("machine_code", "").strip()
        send_enabled = 1 if request.form.get("send_enabled") else 0
        receive_enabled = 1 if request.form.get("receive_enabled") else 0
        db.execute(
            "INSERT INTO mapcodes (test_parameter_id, machine_name, machine_code, send_enabled, receive_enabled) "
            "VALUES (?, ?, ?, ?, ?)",
            (param_id, machine_name, machine_code, send_enabled, receive_enabled),
        )
        db.commit()
        flash("Mapcode added.")
        return redirect(url_for("mapcodes_list"))

    rows = db.execute(
        "SELECT m.*, tp.name as param_name, td.name as test_name FROM mapcodes m "
        "JOIN test_parameters tp ON tp.id = m.test_parameter_id "
        "JOIN test_definitions td ON td.id = tp.test_definition_id ORDER BY td.name LIMIT 200"
    ).fetchall()
    parameters = db.execute(
        "SELECT tp.id, tp.name, td.name as test_name FROM test_parameters tp "
        "JOIN test_definitions td ON td.id = tp.test_definition_id ORDER BY td.name"
    ).fetchall()
    tests = db.execute(
        "SELECT id, name FROM test_definitions ORDER BY name"
    ).fetchall()
    return render_template("master/mapcodes.html", rows=rows, parameters=parameters, tests=tests)


@app.route("/master/mapcodes/<int:mapcode_id>/edit", methods=["POST"])
@roles_required("supervisor")
def update_mapcode(mapcode_id):
    db = get_db()
    machine_name = request.form.get("machine_name", "").strip()
    machine_code = request.form.get("machine_code", "").strip()
    send_enabled = 1 if request.form.get("send_enabled") else 0
    receive_enabled = 1 if request.form.get("receive_enabled") else 0
    db.execute(
        "UPDATE mapcodes SET machine_name=?, machine_code=?, send_enabled=?, receive_enabled=? WHERE id=?",
        (machine_name, machine_code, send_enabled, receive_enabled, mapcode_id),
    )
    db.commit()
    flash(t(session.get("lang", "en"), "mapcode_updated"))
    return redirect(url_for("mapcodes_list"))


@app.route("/master/mapcodes/<int:mapcode_id>/delete", methods=["POST"])
@roles_required("supervisor")
def delete_mapcode(mapcode_id):
    db = get_db()
    db.execute("DELETE FROM mapcodes WHERE id=?", (mapcode_id,))
    db.commit()
    flash(t(session.get("lang", "en"), "mapcode_deleted"))
    return redirect(url_for("mapcodes_list"))


@app.route("/master/host-interface", methods=["GET", "POST"])
@designer_required
def host_interface():
    db = get_db()
    if request.method == "POST":
        enabled = "1" if request.form.get("enabled") else "0"
        ip = request.form.get("ip", "0.0.0.0").strip() or "0.0.0.0"
        port = request.form.get("port", "5000").strip() or "5000"
        set_setting(db, "host_listener_enabled", enabled)
        set_setting(db, "host_listener_ip", ip)
        set_setting(db, "host_listener_port", port)
        db.commit()
        flash(t(session.get("lang", "en"), "host_settings_saved"))
        return redirect(url_for("host_interface"))

    enabled = get_setting(db, "host_listener_enabled", "0") == "1"
    ip = get_setting(db, "host_listener_ip", "0.0.0.0")
    port = get_setting(db, "host_listener_port", "5000")
    logs = db.execute(
        "SELECT * FROM host_interface_log ORDER BY id DESC LIMIT 100"
    ).fetchall()
    return render_template("master/host_interface.html", enabled=enabled, ip=ip, port=port, logs=logs)


@app.route("/master/host-interface/clear-log", methods=["POST"])
@designer_required
def clear_host_interface_log():
    db = get_db()
    db.execute("DELETE FROM host_interface_log")
    db.commit()
    return redirect(url_for("host_interface"))


@app.route("/master/quick-add-items", methods=["GET", "POST"])
@roles_required("supervisor")
def quick_add_items():
    db = get_db()
    if request.method == "POST":
        test_id = request.form.get("test_definition_id")
        exists = db.execute("SELECT id FROM quick_add_items WHERE test_definition_id=?", (test_id,)).fetchone()
        if not exists:
            db.execute("INSERT INTO quick_add_items (test_definition_id) VALUES (?)", (test_id,))
            db.commit()
            flash("Added to Quick Add.")
        return redirect(url_for("quick_add_items"))

    rows = db.execute(
        "SELECT q.id, td.name as test_name, td.department FROM quick_add_items q "
        "JOIN test_definitions td ON td.id = q.test_definition_id WHERE q.is_active=1 ORDER BY q.display_order"
    ).fetchall()
    all_tests = db.execute("SELECT * FROM test_definitions WHERE is_active=1 ORDER BY name").fetchall()
    return render_template("master/quick_add_items.html", rows=rows, all_tests=all_tests)


@app.route("/master/quick-add-items/<int:item_id>/remove", methods=["POST"])
@roles_required("supervisor")
def quick_add_item_remove(item_id):
    db = get_db()
    db.execute("UPDATE quick_add_items SET is_active=0 WHERE id=?", (item_id,))
    db.commit()
    flash("Removed from Quick Add.")
    return redirect(url_for("quick_add_items"))



@app.route("/master/reference-ranges", methods=["GET", "POST"])
@roles_required("supervisor")
def reference_ranges():
    db = get_db()
    if request.method == "POST":
        param_id = request.form.get("test_parameter_id")
        gender = request.form.get("gender") or "Both"
        age_from = request.form.get("age_from") or 0
        age_from_unit = request.form.get("age_from_unit") or "Years"
        age_to = request.form.get("age_to") or 120
        age_to_unit = request.form.get("age_to_unit") or "Years"
        low = request.form.get("low") or None
        high = request.form.get("high") or None
        range_text = request.form.get("range_text") or None
        source_note = request.form.get("source_note") or None
        db.execute(
            "INSERT INTO reference_ranges (test_parameter_id, gender, age_from, age_from_unit, age_to, age_to_unit, low, high, range_text, source_note) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (param_id, gender, age_from, age_from_unit, age_to, age_to_unit, low, high, range_text, source_note),
        )
        db.commit()
        flash("Reference range added.")
        return redirect(url_for("reference_ranges"))

    ranges = db.execute(
        "SELECT rr.*, tp.name as param_name, td.name as test_name FROM reference_ranges rr "
        "JOIN test_parameters tp ON tp.id = rr.test_parameter_id "
        "JOIN test_definitions td ON td.id = tp.test_definition_id "
        "ORDER BY td.name LIMIT 200"
    ).fetchall()
    parameters = db.execute(
        "SELECT tp.id, tp.name, tp.highlight, td.name as test_name FROM test_parameters tp "
        "JOIN test_definitions td ON td.id = tp.test_definition_id ORDER BY td.name"
    ).fetchall()
    return render_template("master/reference_ranges.html", ranges=ranges, parameters=parameters)


@app.route("/master/reference-ranges/<int:range_id>/edit", methods=["POST"])
@roles_required("supervisor")
def update_reference_range(range_id):
    db = get_db()
    gender = request.form.get("gender") or "Both"
    age_from = request.form.get("age_from") or 0
    age_from_unit = request.form.get("age_from_unit") or "Years"
    age_to = request.form.get("age_to") or 120
    age_to_unit = request.form.get("age_to_unit") or "Years"
    low = request.form.get("low") or None
    high = request.form.get("high") or None
    range_text = request.form.get("range_text") or None
    source_note = request.form.get("source_note") or None
    db.execute(
        "UPDATE reference_ranges SET gender=?, age_from=?, age_from_unit=?, age_to=?, age_to_unit=?, low=?, high=?, range_text=?, source_note=? WHERE id=?",
        (gender, age_from, age_from_unit, age_to, age_to_unit, low, high, range_text, source_note, range_id),
    )
    db.commit()
    flash(t(session.get("lang", "en"), "range_updated"))
    return redirect(url_for("reference_ranges"))


@app.route("/master/reference-ranges/<int:range_id>/delete", methods=["POST"])
@roles_required("supervisor")
def delete_reference_range(range_id):
    db = get_db()
    db.execute("DELETE FROM reference_ranges WHERE id=?", (range_id,))
    db.commit()
    flash(t(session.get("lang", "en"), "range_deleted"))
    return redirect(url_for("reference_ranges"))


@app.route("/master/test-catalog/<int:test_id>/price", methods=["POST"])
@roles_required("supervisor")
def update_test_price(test_id):
    db = get_db()
    try:
        price = float(request.form.get("price") or request.json.get("price"))
    except (TypeError, ValueError):
        return {"ok": False, "error": "invalid price"}, 400
    db.execute("UPDATE test_definitions SET price=? WHERE id=?", (price, test_id))
    db.commit()
    return {"ok": True, "price": price}


@app.route("/master/test-catalog/<int:test_id>/toggle-examining", methods=["POST"])
@roles_required("supervisor")
def toggle_examining_test(test_id):
    db = get_db()
    row = db.execute("SELECT is_examining_test FROM test_definitions WHERE id=?", (test_id,)).fetchone()
    if not row:
        return {"ok": False}, 404
    new_val = 0 if row["is_examining_test"] else 1
    db.execute("UPDATE test_definitions SET is_examining_test=? WHERE id=?", (new_val, test_id))
    db.commit()
    return {"ok": True, "is_examining_test": new_val}


# اسم "الريبورت المجمّع" (report_group) اللي يظهر عنوانًا فرعيًا للوحة
# الطباعة المجمّعة (print_combined_panel) — فاضي يرجّع التحليل يعتمد على
# اسم القسم (department) العام كالمعتاد بدل عنوان مخصص.
@app.route("/master/test-catalog/<int:test_id>/report-group", methods=["POST"])
@roles_required("supervisor")
def update_test_report_group(test_id):
    db = get_db()
    value = (request.form.get("report_group") or "").strip()
    row = db.execute("SELECT id FROM test_definitions WHERE id=?", (test_id,)).fetchone()
    if not row:
        return {"ok": False}, 404
    db.execute("UPDATE test_definitions SET report_group=? WHERE id=?", (value, test_id))
    db.commit()
    return {"ok": True, "report_group": value}


@app.route("/master/test-parameters/<int:param_id>/toggle-highlight", methods=["POST"])
@roles_required("supervisor")
def toggle_parameter_highlight(param_id):
    # Yellow-shading a result on the printed report is a manual, per-parameter
    # choice the admin makes here — nothing is highlighted automatically.
    db = get_db()
    row = db.execute("SELECT highlight FROM test_parameters WHERE id=?", (param_id,)).fetchone()
    if not row:
        return {"ok": False}, 404
    new_val = 0 if row["highlight"] else 1
    db.execute("UPDATE test_parameters SET highlight=? WHERE id=?", (new_val, param_id))
    db.commit()
    return {"ok": True, "highlight": new_val}


# ------------------------------------------------------------------------
# أداة "تعديل وحدة القياس" — تسمح للمشرف/الأدمن يغيّر وحدة أي معيار
# (test_parameter) مباشرة (مثل NRBC، أو أي معيار يحتاج تغيير وحدته لاحقًا)
# دون لمس قاعدة البيانات يدويًا. لو التغيير يمثّل تحويل قياس فعلي (مو مجرد
# إعادة تسمية)، تقدر تعطيها معامل تحويل (factor) واتجاه (ضرب/قسمة) فتتحول
# كل المدايات المرجعية (reference_ranges) المرتبطة بهذا المعيار تلقائيًا
# بنفس المعامل والاتجاه، حتى تبقى متوافقة مع الوحدة الجديدة. لإعادة تسمية
# بدون أي تغيير بالأرقام (مثل NRBC من "100/wbc" إلى "/100WBC")، اترك
# المعامل = 1.
# ------------------------------------------------------------------------
@app.route("/master/unit-converter", methods=["GET"])
@roles_required("supervisor")
def unit_converter():
    db = get_db()
    parameters = db.execute(
        "SELECT tp.id, tp.name, tp.unit, tp.unit2, tp.unit2_factor, td.name as test_name, td.department "
        "FROM test_parameters tp JOIN test_definitions td ON td.id = tp.test_definition_id "
        "WHERE tp.result_type = 'Numeric' "
        "ORDER BY td.department, td.name, tp.name"
    ).fetchall()
    return render_template("master/unit_converter.html", parameters=parameters)


# ------------------------------------------------------------------------
# الوحدة الثانية الدائمة لعرض النتيجة بوحدتين بنفس الوقت وقت الطباعة
# (مثلاً mg/dL و mmol/L لنفس الباراميتر) — بخلاف "تحويل الوحدة" أعلاه اللي
# يغيّر الوحدة الأساسية مرة وحدة ويحوّل المديات المرجعية معها، هذا الإعداد
# دائم ولا يلمس الوحدة الأساسية ولا المديات المرجعية إطلاقًا: فقط يضيف قيمة
# محسوبة تلقائيًا (value2 = value1 × factor) تُطبع جنب النتيجة الأصلية.
# اترك حقل "الوحدة الثانية" فاضي لإلغاء/تعطيل الوحدة الثانية لهذا الباراميتر.
# ------------------------------------------------------------------------
@app.route("/master/unit-converter/<int:param_id>/set-dual-unit", methods=["POST"])
@roles_required("supervisor")
def unit_converter_set_dual(param_id):
    db = get_db()
    param = db.execute("SELECT * FROM test_parameters WHERE id=?", (param_id,)).fetchone()
    if not param:
        flash("المعيار غير موجود.")
        return redirect(url_for("unit_converter"))

    unit2 = (request.form.get("unit2") or "").strip()
    if not unit2:
        db.execute("UPDATE test_parameters SET unit2=NULL, unit2_factor=NULL WHERE id=?", (param_id,))
        db.commit()
        log_action("SetDualUnit", "test_parameter", param_id, "cleared")
        flash(f"تم إلغاء الوحدة الثانية لـ \"{param['name']}\".")
        return redirect(url_for("unit_converter"))

    factor_raw = (request.form.get("unit2_factor") or "").strip()
    try:
        factor = float(factor_raw)
        if factor <= 0:
            raise ValueError
    except ValueError:
        flash("معامل التحويل للوحدة الثانية يجب أن يكون رقمًا أكبر من صفر.")
        return redirect(url_for("unit_converter"))

    db.execute("UPDATE test_parameters SET unit2=?, unit2_factor=? WHERE id=?", (unit2, factor, param_id))
    db.commit()
    log_action("SetDualUnit", "test_parameter", param_id, f"{param['unit']} -> {unit2} (x{factor})")
    flash(f"تم حفظ الوحدة الثانية لـ \"{param['name']}\": {unit2} (يُحسب تلقائيًا = القيمة × {factor}).")
    return redirect(url_for("unit_converter"))


@app.route("/master/unit-converter/<int:param_id>/apply", methods=["POST"])
@roles_required("supervisor")
def unit_converter_apply(param_id):
    db = get_db()
    param = db.execute("SELECT * FROM test_parameters WHERE id=?", (param_id,)).fetchone()
    if not param:
        flash("المعيار غير موجود.")
        return redirect(url_for("unit_converter"))

    new_unit = (request.form.get("new_unit") or "").strip()
    direction = request.form.get("direction", "multiply")
    try:
        factor = float(request.form.get("factor") or 1)
    except ValueError:
        factor = 1.0
    if factor <= 0:
        flash("معامل التحويل يجب أن يكون رقمًا أكبر من صفر.")
        return redirect(url_for("unit_converter"))

    old_unit = param["unit"]
    ranges = db.execute("SELECT id, low, high FROM reference_ranges WHERE test_parameter_id=?", (param_id,)).fetchall()
    for r in ranges:
        new_low = r["low"]
        new_high = r["high"]
        if direction == "divide":
            if new_low is not None:
                new_low = new_low / factor
            if new_high is not None:
                new_high = new_high / factor
        else:
            if new_low is not None:
                new_low = new_low * factor
            if new_high is not None:
                new_high = new_high * factor
        db.execute("UPDATE reference_ranges SET low=?, high=? WHERE id=?", (new_low, new_high, r["id"]))

    if new_unit:
        db.execute("UPDATE test_parameters SET unit=? WHERE id=?", (new_unit, param_id))

    db.commit()
    log_action(
        "UnitConvert", "test_parameter", param_id,
        f"{old_unit!r} -> {new_unit!r} ({direction} x{factor}, {len(ranges)} ranges updated)",
    )
    flash(f"تم تحديث وحدة \"{param['name']}\" وتحويل {len(ranges)} مدى مرجعي.")
    return redirect(url_for("unit_converter"))


@app.route("/master/test-catalog", methods=["GET", "POST"])
@roles_required("supervisor")
def test_catalog():
    db = get_db()
    if request.method == "POST":
        code = request.form.get("code", "").strip()
        name = request.form.get("name", "").strip()
        department = request.form.get("department", "").strip()
        report_group = request.form.get("report_group", "").strip()
        sample_type = request.form.get("sample_type", "").strip()
        price = float(request.form.get("price") or 0)
        is_examining_test = 1 if request.form.get("is_examining_test") else 0
        params_raw = request.form.get("parameters", "").strip()

        cur = db.execute(
            "INSERT INTO test_definitions (code, name, department, report_group, sample_type, price, is_examining_test) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (code, name, department, report_group, sample_type, price, is_examining_test),
        )
        test_id = cur.lastrowid
        for chunk in params_raw.split(","):
            chunk = chunk.strip()
            if not chunk:
                continue
            if "unit:" in chunk:
                pname, unit = chunk.split("unit:")
                pname, unit = pname.strip(), unit.strip()
            else:
                pname, unit = chunk, ""
            db.execute(
                "INSERT INTO test_parameters (test_definition_id, name, unit, result_type) "
                "VALUES (?, ?, ?, 'Numeric')",
                (test_id, pname, unit),
            )
        db.commit()
        flash("Test added to catalog.")
        return redirect(url_for("test_catalog"))

    tests = db.execute("SELECT * FROM test_definitions ORDER BY department, name").fetchall()
    # أسماء "الريبورت المجمّع" المستخدمة أصلاً بأي تحليل — تُعرض كاقتراحات
    # بحقل الإدخال (datalist) حتى يعيد الأدمن استخدام نفس الاسم بالضبط بدل
    # ما يكتبه بصيغة مختلفة شوي فينفصل عن مجموعته بالخطأ (مثلاً "Viral
    # study" مرة و"Viral Study" مرة ثانية تصيران مجموعتين منفصلتين).
    report_groups = sorted({
        row["report_group"].strip() for row in tests if row["report_group"] and row["report_group"].strip()
    })
    return render_template("master/test_catalog.html", tests=tests, report_groups=report_groups)


# ------------------------------------------------------------ parameter order
# ترتيب باراميترات أي تحليل (بشاشة إدخال النتائج وبالتقرير المطبوع) —
# صفحة مستقلة عامة لأي تحليل (مو خاصة بس بتحاليل الـdifferential)، لأن
# ترتيب الباراميترات بقاعدة البيانات بيّن (test_parameters.sort_order)
# وما كان فيه شاشة تتحكم فيه سابقًا — كانت تعتمد على ترتيب id (الإدخال).
@app.route("/master/parameter-order", methods=["GET", "POST"])
@roles_required("supervisor")
def parameter_order():
    db = get_db()
    if request.method == "POST":
        test_id = request.form.get("test_id")
        param_ids = request.form.getlist("param_id")
        orders = request.form.getlist("order")
        for pid, order in zip(param_ids, orders):
            try:
                order_val = int(order)
            except (TypeError, ValueError):
                order_val = 0
            db.execute("UPDATE test_parameters SET sort_order=? WHERE id=?", (order_val, pid))
        db.commit()
        flash("تم حفظ ترتيب الباراميترات.")
        return redirect(url_for("parameter_order", test_id=test_id))

    tests = db.execute(
        "SELECT DISTINCT td.id, td.name, td.code FROM test_definitions td "
        "JOIN test_parameters tp ON tp.test_definition_id = td.id "
        "ORDER BY td.name"
    ).fetchall()

    selected_id = request.args.get("test_id", type=int)
    if not selected_id and tests:
        selected_id = tests[0]["id"]

    parameters = []
    if selected_id:
        parameters = db.execute(
            "SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id",
            (selected_id,),
        ).fetchall()

    return render_template("master/parameter_order.html", tests=tests, parameters=parameters, selected_id=selected_id)


def unique_test_code(db, name):
    """يولّد رمزًا فريدًا للتحليل من اسمه الإنكليزي عند الإضافة السريعة من
    قسم الطلبات (بدون الحاجة لإدخال رمز يدويًا)."""
    base = re.sub(r"[^A-Za-z0-9]+", "", name).upper()[:12] or "TEST"
    code = base
    n = 1
    while db.execute("SELECT 1 FROM test_definitions WHERE code=?", (code,)).fetchone():
        n += 1
        code = f"{base}{n}"
    return code


# إضافة/تعديل/إخفاء تحليل بسرعة من داخل نفس نافذة "قسم الطلبات" بشاشتي
# "زيارة جديدة" و"تعديل الزيارة" — دون مغادرة الصفحة أو فقدان بيانات الزيارة
# التي بدأ المستخدم بتعبئتها (لهذا الاستجابة JSON وليس إعادة توجيه صفحة).
@app.route("/api/tests/quick-add", methods=["POST"])
@roles_required("supervisor")
def api_quick_add_test():
    db = get_db()
    name = (request.form.get("name") or "").strip()
    department = (request.form.get("department") or "").strip()
    if not name:
        return jsonify({"error": "الاسم مطلوب"}), 400
    code = unique_test_code(db, name)
    cur = db.execute(
        "INSERT INTO test_definitions (code, name, department, sample_type, price, is_active, is_examining_test) "
        "VALUES (?, ?, ?, '', 0, 1, 0)",
        (code, name, department),
    )
    db.commit()
    log_action("QuickAddTest", "test_definition", cur.lastrowid, name)
    return jsonify({"id": cur.lastrowid, "name": name, "department": department, "code": code})

@app.route("/api/tests/quick-add-with-price", methods=["POST"])
@login_required
def api_quick_add_test_with_price():
    """يسمح لأي مستخدم مسجّل دخول (مو بس admin/supervisor) بإضافة تحليل غير
    موجود بالقائمة مباشرة أثناء إنشاء زيارة جديدة، مع تحديد سعره فورًا —
    يبقى هذا التحليل محفوظًا بالكتالوج بشكل دائم لأي زيارة قادمة. لو التحليل
    موجود مسبقًا بنفس الاسم، نرجع بياناته الحالية بدل ما ننشئ نسخة مكررة."""
    db = get_db()
    name = (request.form.get("name") or "").strip()
    try:
        price = float(request.form.get("price") or 0)
    except ValueError:
        price = 0
    if not name:
        return jsonify({"error": "اسم التحليل مطلوب"}), 400
    if price < 0:
        return jsonify({"error": "السعر يجب أن يكون رقمًا موجبًا"}), 400
    existing = db.execute(
        "SELECT id, price FROM test_definitions WHERE LOWER(TRIM(name))=LOWER(TRIM(?))", (name,)
    ).fetchone()
    if existing:
        return jsonify({"id": existing["id"], "name": name, "price": existing["price"], "existed": True})
    code = unique_test_code(db, name)
    cur = db.execute(
        "INSERT INTO test_definitions (code, name, department, sample_type, price, is_active, is_examining_test) "
        "VALUES (?, ?, '', '', ?, 1, 0)",
        (code, name, price),
    )
    db.commit()
    log_action("QuickAddTestWithPrice", "test_definition", cur.lastrowid, f"{name} ({price})")
    return jsonify({"id": cur.lastrowid, "name": name, "price": price, "existed": False})

# إضافة تحليل غير موجود بالقائمة مباشرة من قسم "الطلبات" بشاشة "زيارة
# جديدة" — متاحة لأي مستخدم مسجّل دخول (وليس فقط admin/supervisor مثل
# نافذة "إدارة التحاليل" وراوت /api/tests/quick-add أعلاه)، لأن موظف
# الاستقبال هو من يواجه هذا الموقف يوميًا (مريض طلب تحليل غير مُدرج بعد).
# يُطلب السعر إجباريًا هنا (بعكس /api/tests/quick-add اللي يحفظه 0 مؤقتًا
# بانتظار أن يعدّله المدير لاحقًا من كتالوج التحاليل) حتى يدخل التحليل
# فورًا بجدول الفاتورة/المجموع بسعره الصحيح دون انتظار أحد.
@app.route("/api/tests/quick-add-priced", methods=["POST"])
@login_required
def api_quick_add_test_priced():
    db = get_db()
    name = (request.form.get("name") or "").strip()
    price_raw = (request.form.get("price") or "").strip()
    if not name:
        return jsonify({"error": "الاسم مطلوب"}), 400
    try:
        price = float(price_raw)
        if price < 0:
            raise ValueError
    except ValueError:
        return jsonify({"error": "السعر يجب أن يكون رقمًا (0 أو أكثر)"}), 400
    code = unique_test_code(db, name)
    cur = db.execute(
        "INSERT INTO test_definitions (code, name, department, sample_type, price, is_active, is_examining_test) "
        "VALUES (?, ?, '', '', ?, 1, 0)",
        (code, name, price),
    )
    db.commit()
    log_action("QuickAddTestPriced", "test_definition", cur.lastrowid, f"{name} ({price})")
    return jsonify({"id": cur.lastrowid, "name": name, "price": price, "code": code})


# إضافة اسم طبيب فاحص جديد بسرعة من نفس شاشة "زيارة جديدة" (القائمتين
# "دكتور المختبر الفاحص" و"الدكتور الفاحص") — يُحفظ بنفس قائمة الإعدادات
# التي يديرها المدير من Management → Settings، فيظهر لاحقًا هناك أيضًا.
@app.route("/api/examining-doctors/quick-add", methods=["POST"])
@login_required
def api_quick_add_examining_doctor():
    db = get_db()
    name = (request.form.get("name") or "").strip()
    if not name:
        return jsonify({"error": "الاسم مطلوب"}), 400
    add_examining_doctor(db, name)
    log_action("QuickAddExaminingDoctor", "settings", 0, name)
    return jsonify({"name": name})


@app.route("/api/tests/<int:test_id>/quick-edit", methods=["POST"])
@roles_required("supervisor")
def api_quick_edit_test(test_id):
    db = get_db()
    name = (request.form.get("name") or "").strip()
    department = (request.form.get("department") or "").strip()
    if not name:
        return jsonify({"error": "الاسم مطلوب"}), 400
    test = db.execute("SELECT id FROM test_definitions WHERE id=?", (test_id,)).fetchone()
    if not test:
        return jsonify({"error": "not found"}), 404
    db.execute("UPDATE test_definitions SET name=?, department=? WHERE id=?", (name, department, test_id))
    db.commit()
    log_action("QuickEditTest", "test_definition", test_id, name)
    return jsonify({"id": test_id, "name": name, "department": department})


@app.route("/api/tests/<int:test_id>/quick-hide", methods=["POST"])
@roles_required("supervisor")
def api_quick_hide_test(test_id):
    db = get_db()
    test = db.execute("SELECT id, is_active FROM test_definitions WHERE id=?", (test_id,)).fetchone()
    if not test:
        return jsonify({"error": "not found"}), 404
    new_val = 0 if test["is_active"] else 1
    db.execute("UPDATE test_definitions SET is_active=? WHERE id=?", (new_val, test_id))
    db.commit()
    log_action("QuickToggleTestActive", "test_definition", test_id, str(new_val))
    return jsonify({"id": test_id, "is_active": new_val})


@app.route("/api/tests/active-list")
@roles_required("supervisor")
def api_tests_active_list():
    db = get_db()
    tests = db.execute(
        "SELECT id, name, department, is_active FROM test_definitions WHERE is_active=1 ORDER BY department, name"
    ).fetchall()
    return jsonify([dict(t) for t in tests])



# --------------------------------------------------------------- management
@app.route("/management/settings", methods=["GET", "POST"])
@roles_required("admin")
def app_settings():
    db = get_db()
    if request.method == "POST":
        name_en = request.form.get("app_name", "").strip()
        name_ar = request.form.get("app_name_ar", "").strip()
        if name_en:
            set_setting(db, "app_name", name_en)
        if name_ar:
            set_setting(db, "app_name_ar", name_ar)

        lab_address_raw = request.form.get("lab_address", "").strip()
        if lab_address_raw:
            set_setting(db, "lab_address", lab_address_raw)
        lab_phone_raw = request.form.get("lab_phone", "").strip()
        if lab_phone_raw:
            set_setting(db, "lab_phone", lab_phone_raw)

        file = request.files.get("logo")
        if file and file.filename:
            ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
            if ext in ALLOWED_LOGO_EXT:
                filename = secure_filename(f"logo.{ext}")
                # remove any previously uploaded logo with a different extension
                for old_ext in ALLOWED_LOGO_EXT:
                    old_path = os.path.join(UPLOAD_DIR, f"logo.{old_ext}")
                    if os.path.exists(old_path):
                        os.remove(old_path)
                file.save(os.path.join(UPLOAD_DIR, filename))
                set_setting(db, "logo_path", f"uploads/{filename}")
            else:
                flash("Unsupported logo file type. Use PNG, JPG, GIF, SVG or WEBP.")

        # خلفية شاشة الترحيب (dashboard) — نفس منطق رفع الشعار أعلاه بالضبط:
        # نحذف أي نسخة قديمة (بأي امتداد) ثم نحفظ الجديدة باسم ثابت
        # dashboard-bg.<ext> حتى يبقى مسار واحد معروف. زر "إزالة الصورة"
        # منفصل (checkbox/hidden باسم remove_dashboard_bg) يمسحها بدون رفع بديل.
        if request.form.get("remove_dashboard_bg") == "1":
            old_bg_path = get_setting(db, "dashboard_bg_path", "")
            if old_bg_path:
                old_bg_full = os.path.join(UPLOAD_DIR, os.path.basename(old_bg_path))
                if os.path.exists(old_bg_full):
                    os.remove(old_bg_full)
            set_setting(db, "dashboard_bg_path", "")
        else:
            bg_file = request.files.get("dashboard_bg")
            if bg_file and bg_file.filename:
                bg_ext = bg_file.filename.rsplit(".", 1)[-1].lower() if "." in bg_file.filename else ""
                if bg_ext in ALLOWED_DASHBOARD_BG_EXT:
                    bg_filename = secure_filename(f"dashboard-bg.{bg_ext}")
                    for old_ext in ALLOWED_DASHBOARD_BG_EXT:
                        old_bg = os.path.join(UPLOAD_DIR, f"dashboard-bg.{old_ext}")
                        if os.path.exists(old_bg):
                            os.remove(old_bg)
                    bg_file.save(os.path.join(UPLOAD_DIR, bg_filename))
                    set_setting(db, "dashboard_bg_path", f"uploads/{bg_filename}")
                else:
                    flash("صيغة صورة غير مدعومة لخلفية شاشة الترحيب. استخدم PNG أو JPG أو WEBP.")

        # ملاحظة: إدارة أسماء وشهادات دكاترة الفحص انتقلت لشاشة مستقلة
        # (management/examining_doctors.html عبر الرابط بهذي الصفحة) — ما عاد
        # فيها فورم هنا.

        # ارتفاع/عرض خلايا جدول نتائج التقرير — رقم صحيح بين 2 و30 بكسل فقط،
        # أي قيمة غير صالحة تُتجاهل ويبقى المحفوظ سابقًا كما هو.
        row_pad_raw = request.form.get("report_row_pad", "").strip()
        if row_pad_raw:
            try:
                row_pad = max(2, min(30, int(row_pad_raw)))
                set_setting(db, "report_row_pad", str(row_pad))
            except ValueError:
                pass
        col_pad_raw = request.form.get("report_col_pad", "").strip()
        if col_pad_raw:
            try:
                col_pad = max(2, min(40, int(col_pad_raw)))
                set_setting(db, "report_col_pad", str(col_pad))
            except ValueError:
                pass

        wa_code = request.form.get("whatsapp_country_code", "").strip()
        if wa_code:
            set_setting(db, "whatsapp_country_code", "".join(ch for ch in wa_code if ch.isdigit()))

        # حجم ونوع خط أسماء/شهادات الدكاترة بترويسة كل تقرير مطبوع (base_report.html)
        # — إعداد عام واحد ينطبق على الجميع دفعة وحدة، مو لكل شخص لحاله.
        lh_size_raw = request.form.get("letterhead_font_size", "").strip()
        if lh_size_raw:
            try:
                lh_size = max(8, min(30, int(lh_size_raw)))
                set_setting(db, "letterhead_font_size", str(lh_size))
            except ValueError:
                pass
        lh_family_raw = request.form.get("letterhead_font_family", "").strip()
        if lh_family_raw:
            set_setting(db, "letterhead_font_family", lh_family_raw)

        # موضع الشعار (يمين الافتراضي/وسط/يسار) وحجمه (عرضه بالبكسل) —
        # إعدادان عامان جديدان، نفس أسلوب باقي إعدادات التقرير.
        logo_pos_raw = request.form.get("logo_position", "").strip()
        if logo_pos_raw in ("right", "center", "left"):
            set_setting(db, "logo_position", logo_pos_raw)
        logo_width_raw = request.form.get("logo_width", "").strip()
        if logo_width_raw:
            try:
                logo_width = max(30, min(300, int(logo_width_raw)))
                set_setting(db, "logo_width", str(logo_width))
            except ValueError:
                pass

        # ترتيب أقسام اللوحة المجمّعة — سطر واحد لكل اسم قسم/ريبورت مجمّع
        # (report_group أو department)، بنفس الترتيب اللي يريده الأدمن
        # بالطباعة. حقل نصي فاضي بالكامل = رجوع للترتيب الأبجدي الافتراضي.
        panel_group_order_raw = request.form.get("combined_panel_group_order")
        if panel_group_order_raw is not None:
            set_setting(db, "combined_panel_group_order", panel_group_order_raw.strip())

        # ترتيب ثابت للتحاليل بشاشة "إدخال النتائج" (Results Entry) — سطر
        # لكل اسم تحليل بالضبط كما يظهر بالشاشة (مثلاً "HbA1c (BioChemstry)")،
        # بنفس الترتيب اللي يريده الأدمن، ثابت دائماً بغض النظر عن ترتيب
        # طلب التحاليل الفعلي بكل زيارة. فاضي بالكامل = رجوع لترتيب الطلب
        # الأصلي (ot.id) كالسابق.
        results_order_raw = request.form.get("results_entry_test_order")
        if results_order_raw is not None:
            set_setting(db, "results_entry_test_order", results_order_raw.strip())

        # حجم خط اسم التحليل وحجم خط النتيجة بجدول/بطاقات النتائج بكل
        # التقارير المطبوعة — إعدادان عامان منفصلان عن بعض (وعن حجم خط
        # ترويسة الدكاترة letterhead_font_size أعلاه)، دفعة وحدة لكل
        # التقارير. أي قيمة غير صالحة تُتجاهل ويبقى المحفوظ سابقًا كما هو.
        tn_size_raw = request.form.get("test_name_font_size", "").strip()
        if tn_size_raw:
            try:
                tn_size = max(8, min(30, int(tn_size_raw)))
                set_setting(db, "test_name_font_size", str(tn_size))
            except ValueError:
                pass
        rv_size_raw = request.form.get("result_value_font_size", "").strip()
        if rv_size_raw:
            try:
                rv_size = max(8, min(30, int(rv_size_raw)))
                set_setting(db, "result_value_font_size", str(rv_size))
            except ValueError:
                pass

        # مجلد أرشفة الـPDF الدائم (نقطة #8) — يُضبط مرة وحدة هنا وقت
        # التنصيب/الإعداد الأولي، ويُعاد استخدامه تلقائيًا بعدها لكل زيارة
        # تكتمل نتائجها. لا نتحقق من وجوده هنا (os.makedirs لاحقًا وقت
        # الأرشفة الفعلية يكفي) حتى يقدر المدير يكتب مسار جهاز آخر بالشبكة.
        archive_dir_raw = request.form.get("pdf_archive_dir", "").strip()
        if archive_dir_raw:
            set_setting(db, "pdf_archive_dir", archive_dir_raw)

        # ألوان رموز الـ Conclusion (نجمة/أسهم/استفهام/تعجب) — كل رمز إله
        # حقل <input type=color> باسم marker_color_<key> بشاشة الإعدادات.
        # نحفظ فقط الألوان اللي وصلت وبصيغة hex صحيحة؛ أي حقل فاضي أو غير
        # صالح يبقى على قيمته المحفوظة سابقًا (أو الافتراضي إذا ما انحفظ شي).
        marker_colors = {}
        raw_existing = get_setting(db, "conclusion_marker_colors", "")
        if raw_existing:
            try:
                marker_colors = json.loads(raw_existing)
            except (ValueError, TypeError):
                marker_colors = {}
        for m in CONCLUSION_MARKERS:
            val = request.form.get("marker_color_" + m["key"], "").strip()
            if val and _HEX_COLOR_RE.match(val):
                marker_colors[m["key"]] = val
        if marker_colors:
            set_setting(db, "conclusion_marker_colors", json.dumps(marker_colors))

        db.commit()
        log_action("UpdateSettings", "settings", 0)
        flash("Settings saved.")
        return redirect(url_for("app_settings"))

    current = {
        "app_name": get_setting(db, "app_name", ""),
        "app_name_ar": get_setting(db, "app_name_ar", ""),
        "lab_address": get_setting(db, "lab_address", ""),
        "lab_phone": get_setting(db, "lab_phone", ""),
        "logo_path": get_setting(db, "logo_path", ""),
        "dashboard_bg_path": get_setting(db, "dashboard_bg_path", ""),
        "report_row_pad": get_setting(db, "report_row_pad", "5"),
        "report_col_pad": get_setting(db, "report_col_pad", "12"),
        "whatsapp_country_code": get_setting(db, "whatsapp_country_code", "964"),
        "pdf_archive_dir": get_setting(db, "pdf_archive_dir", ""),
        "letterhead_font_size": get_setting(db, "letterhead_font_size", "14"),
        "letterhead_font_family": get_setting(db, "letterhead_font_family", "Segoe UI, Tahoma, Arial, sans-serif"),
        "test_name_font_size": get_setting(db, "test_name_font_size", "16"),
        "result_value_font_size": get_setting(db, "result_value_font_size", "16"),
        "logo_position": get_setting(db, "logo_position", "right"),
        "logo_width": get_setting(db, "logo_width", "100"),
        "combined_panel_group_order": get_setting(db, "combined_panel_group_order", ""),
        "results_entry_test_order": get_setting(db, "results_entry_test_order", ""),
    }
    marker_colors_by_char = get_conclusion_marker_colors(db)
    conclusion_markers = [
        {**m, "color": marker_colors_by_char[m["char"]]} for m in CONCLUSION_MARKERS
    ]
    return render_template(
        "management/settings.html", current=current, conclusion_markers=conclusion_markers
    )


# ------------------------------------------------------------------------
# تبديل "التحديث التلقائي" بضغطة وحدة من شريط الأعلى (بجانب اسم المستخدم) —
# متاح لدخول الأدمن العادي، بعكس /designer/update/toggle اللي يحتاج جلسة
# "مصمم" منفصلة. الاثنين يكتبان لنفس المفتاح (auto_update_enabled) بجدول
# الإعدادات، فتفعيل/تعطيل أي وحدة منهم ينعكس على الثاني وعلى auto_updater
# نفسه فورًا. يرجّع المستخدم لنفس الصفحة اللي كان فيها (request.referrer).
# ------------------------------------------------------------------------
@app.route("/settings/auto-update/toggle", methods=["POST"])
@roles_required("admin")
def toggle_auto_update():
    db = get_db()
    currently_on = get_setting(db, "auto_update_enabled", "1") == "1"
    set_setting(db, "auto_update_enabled", "0" if currently_on else "1")
    db.commit()
    log_action("ToggleAutoUpdate", "settings", 0, "off" if currently_on else "on")
    db.close()
    flash("تم إيقاف التحديث التلقائي." if currently_on else "تم تفعيل التحديث التلقائي.")
    return redirect(request.referrer or url_for("dashboard"))


# ------------------------------------------------------------------------
# إدارة قائمة (دكتور المختبر الفاحص) — اسم/لقب/شهادة عربي/شهادة انكليزي
# وترتيب يدوي (فوق/تحت)، بالإضافة لمفتاح "إظهار بترويسة التقرير" لكل واحد
# منهم. مفتوحة من بطاقة "دكتور المختبر الفاحص" بصفحة الإعدادات. الأسماء
# نفسها تبقى تُستخدم كنص بجداول أخرى (أجور الفحص، الزيارات) فتغيير الاسم
# هنا لا يحدّث تلقائيًا سجلات قديمة محفوظة بالاسم السابق.
# ------------------------------------------------------------------------
@app.route("/management/examining-doctors", methods=["GET"])
@roles_required("admin")
def examining_doctors_manage():
    db = get_db()
    doctors = get_examining_doctors_full(db)
    return render_template("management/examining_doctors.html", doctors=doctors)


@app.route("/management/examining-doctors/add", methods=["POST"])
@roles_required("admin")
def examining_doctor_add():
    db = get_db()
    name = request.form.get("name", "").strip()
    if not name:
        flash("اسم الدكتور مطلوب.")
        return redirect(url_for("examining_doctors_manage"))
    title = request.form.get("title", "الدكتور").strip()
    degree_ar = request.form.get("degree_ar", "").strip() or None
    degree_en = request.form.get("degree_en", "").strip() or None
    show_on_letterhead = bool(request.form.get("show_on_letterhead"))
    # حجم خط مخصص لهذا الدكتور تحديداً بالترويسة (اختياري) — يتجاوز الحجم
    # العام letterhead_font_size من الإعدادات لو تُرك فاضي يرث العام كالمعتاد.
    font_size_raw = request.form.get("font_size", "").strip()
    font_size = None
    if font_size_raw:
        try:
            font_size = max(8, min(30, int(font_size_raw)))
        except ValueError:
            font_size = None
    add_examining_doctor_full(db, name, title, degree_ar, degree_en, show_on_letterhead, font_size)
    log_action("AddExaminingDoctor", "examining_doctor", 0, name)
    flash(f"تمت إضافة \"{name}\".")
    return redirect(url_for("examining_doctors_manage"))


@app.route("/management/examining-doctors/<int:doctor_id>/update", methods=["POST"])
@roles_required("admin")
def examining_doctor_update(doctor_id):
    db = get_db()
    name = request.form.get("name", "").strip()
    if not name:
        flash("اسم الدكتور مطلوب.")
        return redirect(url_for("examining_doctors_manage"))
    title = request.form.get("title", "الدكتور").strip()
    degree_ar = request.form.get("degree_ar", "").strip() or None
    degree_en = request.form.get("degree_en", "").strip() or None
    show_on_letterhead = bool(request.form.get("show_on_letterhead"))
    font_size_raw = request.form.get("font_size", "").strip()
    font_size = None
    if font_size_raw:
        try:
            font_size = max(8, min(30, int(font_size_raw)))
        except ValueError:
            font_size = None
    update_examining_doctor(db, doctor_id, name, title, degree_ar, degree_en, show_on_letterhead, font_size)
    log_action("UpdateExaminingDoctor", "examining_doctor", doctor_id, name)
    flash(f"تم حفظ تعديلات \"{name}\".")
    return redirect(url_for("examining_doctors_manage"))


@app.route("/management/examining-doctors/<int:doctor_id>/delete", methods=["POST"])
@roles_required("admin")
def examining_doctor_delete(doctor_id):
    db = get_db()
    delete_examining_doctor(db, doctor_id)
    log_action("DeleteExaminingDoctor", "examining_doctor", doctor_id)
    flash("تم الحذف.")
    return redirect(url_for("examining_doctors_manage"))


@app.route("/management/examining-doctors/<int:doctor_id>/move", methods=["POST"])
@roles_required("admin")
def examining_doctor_move(doctor_id):
    db = get_db()
    direction = request.form.get("direction", "")
    if direction in ("up", "down"):
        move_examining_doctor(db, doctor_id, direction)
    return redirect(url_for("examining_doctors_manage"))


def _parse_docx_rows(file_storage, parameters):
    """Read an admin-uploaded Word (.docx) reference report and pull out an
    ordered list of {"param_name", "label"} rows by matching each line/cell
    of text in the document against this test's existing parameter names.
    Returns (rows, matched_count, total_lines_found)."""
    try:
        import docx  # python-docx
    except ImportError:
        return None, 0, 0

    document = docx.Document(file_storage)
    param_by_lower = {p["name"].strip().lower(): p["name"] for p in parameters}

    def match_param(text):
        low = text.strip().lower()
        low = low.rstrip(":.").strip()
        if not low:
            return None
        if low in param_by_lower:
            return param_by_lower[low]
        for pname_low, pname in param_by_lower.items():
            if pname_low and (pname_low in low or low in pname_low):
                return pname
        return None

    raw_lines = []
    for table in document.tables:
        for trow in table.rows:
            cells = [c.text.strip() for c in trow.cells if c.text.strip()]
            if cells:
                raw_lines.append(cells[0])
    if not raw_lines:
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                raw_lines.append(text)

    rows = []
    matched = 0
    seen = set()
    for line in raw_lines:
        pname = match_param(line)
        if pname and pname in seen:
            continue
        if pname:
            seen.add(pname)
            matched += 1
        rows.append({"param_name": pname or "", "label": line})
    return rows, matched, len(raw_lines)


@app.route("/management/report-designer", methods=["GET", "POST"])
@roles_required("admin")
def report_designer():
    db = get_db()
    hardcoded_codes = set(REPORT_TEMPLATE_MAP.keys())

    if request.method == "POST":
        test_id = request.form.get("test_definition_id")
        test = db.execute("SELECT * FROM test_definitions WHERE id=?", (test_id,)).fetchone()
        if not test:
            flash("التحليل غير موجود.")
            return redirect(url_for("report_designer"))

        parameters = db.execute(
            "SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id", (test_id,)
        ).fetchall()
        mode = request.form.get("mode", "manual")
        source_docx_name = None
        existing_tpl = get_report_template(db, test_id)
        heading_align = request.form.get("heading_align") or (existing_tpl["heading_align"] if existing_tpl else None) or "center"
        rows_align = request.form.get("rows_align") or (existing_tpl["rows_align"] if existing_tpl else None) or "right"

        if mode == "docx":
            file = request.files.get("docx_file")
            if not file or not file.filename:
                flash("لم تختر ملف Word.")
                return redirect(url_for("report_designer", test_definition_id=test_id))
            ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
            if ext not in ALLOWED_DOCX_EXT:
                flash("صيغة الملف يجب أن تكون .docx")
                return redirect(url_for("report_designer", test_definition_id=test_id))
            rows, matched, total = _parse_docx_rows(file, parameters)
            if rows is None:
                flash("مكتبة قراءة ملفات Word (python-docx) غير مثبتة بعد. شغّل تشغيل البرنامج.vbs مرة واحدة وأنت متصل بالإنترنت ليتم تثبيتها تلقائيًا، ثم أعد المحاولة.")
                return redirect(url_for("report_designer", test_definition_id=test_id))
            if not rows:
                flash("لم أجد أي صفوف أو نصوص داخل ملف Word هذا.")
                return redirect(url_for("report_designer", test_definition_id=test_id))
            rows_json = json.dumps(rows, ensure_ascii=False)
            source_docx_name = secure_filename(file.filename)
            save_report_template(db, test_id, test["name"], rows_json, source_docx_name, session["user_id"],
                                  heading_align=heading_align, rows_align=rows_align)
            log_action("SaveReportTemplate", "test_definition", int(test_id), f"docx:{source_docx_name}")
            flash(f"تم استيراد التصميم من Word: {matched} من {total} سطر تم ربطها تلقائيًا بنتائج التحليل. "
                  f"الشعار وترويسة الأطباء وذيل الصفحة ستُضاف تلقائيًا عند الطباعة — راجع الصفوف غير المرتبطة (إن وجدت) وعدّلها يدويًا أدناه.")
            return redirect(url_for("report_designer", test_definition_id=test_id))

        if mode == "panel":
            # إعدادات "اللوحة المجمّعة" — مستويان: (أ) التحليل كامل، تنطبق
            # على كل باراميتراته كسلوك احتياطي؛ (ب) كل باراميتر لحاله
            # (panel_color_<pid>/panel_page_break_<pid>) تكتسح مستوى التحليل
            # لنفس الباراميتر لو ضُبطت — تسمح بتلوين/فصل باراميتر وحدة بس
            # (زي NRBC) داخل تحليل متعدد الباراميترات دون التأثير على الباقي.
            panel_color = request.form.get("panel_color", "").strip()
            enable_panel_color = bool(request.form.get("enable_panel_color"))
            panel_color_val = panel_color if (enable_panel_color and panel_color and _HEX_COLOR_RE.match(panel_color)) else None
            panel_page_break = 1 if request.form.get("panel_page_break") else 0
            db.execute(
                "UPDATE test_definitions SET panel_color=?, panel_page_break=? WHERE id=?",
                (panel_color_val, panel_page_break, test_id),
            )
            for p in parameters:
                pid = p["id"]
                p_color = request.form.get(f"param_panel_color_{pid}", "").strip()
                p_enable = bool(request.form.get(f"enable_param_panel_color_{pid}"))
                p_color_val = p_color if (p_enable and p_color and _HEX_COLOR_RE.match(p_color)) else None
                p_page_break = 1 if request.form.get(f"param_panel_page_break_{pid}") else 0
                db.execute(
                    "UPDATE test_parameters SET panel_color=?, panel_page_break=? WHERE id=?",
                    (p_color_val, p_page_break, pid),
                )
            db.commit()
            log_action("UpdatePanelSettings", "test_definition", int(test_id), "panel")
            flash("تم حفظ إعدادات اللوحة المجمّعة.")
            return redirect(url_for("report_designer", test_definition_id=test_id))

        # mode == manual: rebuild the row order from the submitted list of
        # parameter ids (checked + reordered by the admin in the form).
        heading = request.form.get("heading", "").strip() or test["name"]
        ordered_param_ids = request.form.getlist("param_order")
        rows = []
        params_by_id = {str(p["id"]): p for p in parameters}
        for pid in ordered_param_ids:
            p = params_by_id.get(pid)
            if p:
                row = {"param_name": p["name"], "label": p["name"]}
                # حرية تحريك الحقول لكل صف — تُقرأ فقط لو مصمم التقرير
                # (report_designer.html) يرسلها فعليًا كحقول name_side_<id>/
                # range_position_<id> (زر أو قائمة اختيار جنب كل صف)؛ إذا ما
                # أُرسلت (كل الحالات الحالية) تبقى القيم الافتراضية كما هي
                # ولا يتغيّر أي تصميم محفوظ سابقًا.
                name_side = request.form.get(f"name_side_{pid}")
                if name_side in ("left", "right"):
                    row["name_side"] = name_side
                range_position = request.form.get(f"range_position_{pid}")
                if range_position in ("inline", "below"):
                    row["range_position"] = range_position
                # لون مخصص لهذا التحليل (اسمه ونتيجته) — اختياري، فاضي يعني
                # يبقى باللون الافتراضي بكل التقرير. وفاصل صفحة: لو مفعّل،
                # هذا التحليل يبدأ دائمًا بأعلى صفحة جديدة عند الطباعة (يُقرأ
                # بـ custom.html كـ page-break-before قبل بطاقة هذا الصف).
                row_color = request.form.get(f"row_color_{pid}", "").strip()
                if request.form.get(f"enable_color_{pid}") and row_color and _HEX_COLOR_RE.match(row_color):
                    row["color"] = row_color
                if request.form.get(f"page_break_{pid}"):
                    row["page_break_before"] = True
                rows.append(row)
        if not rows:
            flash("اختر باراميتر واحد على الأقل لتصميم التقرير.")
            return redirect(url_for("report_designer", test_definition_id=test_id))
        # تخطيط أعمدة بديل لهذا التقرير كامل (وحدة بعمود مستقل بأقصى اليمين
        # + نتيجة موسّطة بعمودها) بدل الوضع الافتراضي (الوحدة ملتصقة بنهاية
        # النتيجة بنفس العمود) — يُضبط مرة وحدة لكل التقرير من نفس الفورم.
        unit_column = 1 if request.form.get("unit_column") else 0
        rows_json = json.dumps(rows, ensure_ascii=False)
        save_report_template(db, test_id, heading, rows_json, None, session["user_id"],
                              heading_align=heading_align, rows_align=rows_align, unit_column=unit_column)
        log_action("SaveReportTemplate", "test_definition", int(test_id), "manual")
        flash("تم حفظ تصميم التقرير. الشعار سيُضاف تلقائيًا عند الطباعة.")
        return redirect(url_for("report_designer", test_definition_id=test_id))

    all_tests = db.execute(
        "SELECT * FROM test_definitions WHERE is_active=1 ORDER BY department, name"
    ).fetchall()
    tests_status = []
    for test in all_tests:
        has_builtin = test["code"] in hardcoded_codes
        custom = get_report_template(db, test["id"])
        tests_status.append({
            "test": test,
            "has_builtin": has_builtin,
            "has_custom": bool(custom),
            "needs_design": not has_builtin and not custom,
        })

    selected_id = request.args.get("test_definition_id", type=int)
    selected_test = None
    selected_parameters = []
    selected_template = None
    selected_rows_by_param = {}
    if selected_id:
        selected_test = db.execute("SELECT * FROM test_definitions WHERE id=?", (selected_id,)).fetchone()
        if selected_test:
            selected_parameters = db.execute(
                "SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id", (selected_id,)
            ).fetchall()
            selected_template = get_report_template(db, selected_id)
            if selected_template:
                for r in json.loads(selected_template["rows_json"] or "[]"):
                    selected_rows_by_param[r.get("param_name", "")] = r

    # ترتيب عرض الباراميترات بشاشة التصميم يدويًا: لو فيه تصميم محفوظ سابقًا،
    # نعرضهم بنفس ترتيب rows_json المحفوظ (نفس ترتيب الطباعة فعليًا) أولاً،
    # وأي باراميتر جديد أُضيف للتحليل لاحقًا (بعد آخر حفظ) يُلحق بالنهاية غير
    # مؤشر عليه — هذا يخلي "إضافة تحليل/باراميتر جديد لاحقًا" يظهر تلقائيًا
    # هنا جاهز للتأشير عليه دون أي خطوة إضافية. بدون تصميم سابق، الترتيب
    # الافتراضي هو ترتيب الإدخال بقاعدة البيانات كالمعتاد.
    ordered_selected_params = list(selected_parameters)
    if selected_template:
        params_by_name = {p["name"]: p for p in selected_parameters}
        saved_order = [r.get("param_name", "") for r in json.loads(selected_template["rows_json"] or "[]")]
        seen = set()
        ordered = []
        for name in saved_order:
            p = params_by_name.get(name)
            if p and name not in seen:
                ordered.append(p)
                seen.add(name)
        for p in selected_parameters:
            if p["name"] not in seen:
                ordered.append(p)
                seen.add(p["name"])
        ordered_selected_params = ordered

    return render_template(
        "management/report_designer.html",
        tests_status=tests_status, selected_test=selected_test,
        selected_parameters=ordered_selected_params, selected_template=selected_template,
        selected_rows_by_param=selected_rows_by_param,
    )


@app.route("/management/report-designer/preview/<int:test_definition_id>")
@roles_required("admin")
def preview_report_design(test_definition_id):
    # Renders the exact same report template print_report() uses, but with
    # clearly-labeled sample data instead of a real order/patient — so the
    # admin can actually SEE a built-in or custom design (logo, headings,
    # row layout) from the report-designer page without needing to create a
    # real visit/result first. Reads test_definitions/test_parameters only;
    # never touches patient, visit, or result data.
    db = get_db()
    test = db.execute("SELECT * FROM test_definitions WHERE id=?", (test_definition_id,)).fetchone()
    if not test:
        return "Not found", 404

    template_name = REPORT_TEMPLATE_MAP.get(test["code"])
    custom_template = None
    if not template_name:
        custom_template = get_report_template(db, test_definition_id)
        if not custom_template:
            flash("لا يوجد تصميم لهذا التحليل بعد لتتم معاينته. صممه أولاً بالأسفل.")
            return redirect(url_for("report_designer", test_definition_id=test_definition_id))
        template_name = "reports/custom.html"

    parameters = db.execute(
        "SELECT * FROM test_parameters WHERE test_definition_id=? ORDER BY sort_order, id", (test_definition_id,)
    ).fetchall()
    units_by_name = {p["name"]: p["unit"] for p in parameters}
    highlight_by_name = {p["name"]: bool(p["highlight"]) for p in parameters}
    show_prev_values = department_shows_previous_values(test["department"])
    logo_path = get_setting(db, "logo_path", "")
    logo_url = url_for("static", filename=logo_path) if logo_path else None

    cbc_groups = None
    if test["code"] == "CBC":
        cbc_groups = []
        for group in CBC_ROW_GROUPS:
            rows = [{"name": name, "result": "—", "unit": units_by_name.get(name, ""),
                     "low": "", "high": "", "highlight": highlight_by_name.get(name, False), "previous": "—" if show_prev_values else ""}
                    for name in group]
            cbc_groups.append(rows)

    custom_rows = None
    custom_heading = None
    custom_heading_align = "center"
    custom_rows_align = "right"
    custom_unit_column = False
    if custom_template:
        custom_heading = custom_template["heading"] or test["name"]
        custom_heading_align = custom_template["heading_align"] or "center"
        custom_rows_align = custom_template["rows_align"] or "right"
        custom_unit_column = bool(custom_template["unit_column"])
        row_defs = json.loads(custom_template["rows_json"] or "[]")
        custom_rows = [{
            "label": rd.get("label") or rd.get("param_name", ""),
            "result": "—", "unit": units_by_name.get(rd.get("param_name", ""), ""),
            "normal_range": "—", "range_tiers": [{"label": None, "value": "—"}],
            "normal_range2": None, "result2": None, "unit2": "",
            "previous": "—" if show_prev_values else "",
            "name_side": rd.get("name_side") or "left",
            "range_position": rd.get("range_position") or "inline",
            "name_align": rd.get("name_align"),
            "color": rd.get("color"),
            "page_break_before": rd.get("page_break_before", False),
        } for rd in row_defs]

    params = {p["name"]: "—" for p in parameters}
    units = units_by_name

    return render_template(
        template_name,
        ot={"test_name": test["name"]}, params=params, ranges={}, units=units, cbc_groups=cbc_groups,
        custom_rows=custom_rows, custom_heading=custom_heading,
        custom_heading_align=custom_heading_align, custom_rows_align=custom_rows_align,
        custom_unit_column=custom_unit_column,
        show_prev_values=show_prev_values, previous_visit_date=None, previous_values={},
        repeat_header_on_print=department_shows_previous_values(test["department"]),
        logo_url=logo_url, from_other_lab=False, font_size=14 if test["code"] == "CBC" else 16,
        show_exam_signature=False,
        visit_date=f"{datetime.now().day}/{datetime.now().month}/{datetime.now().year}", sex="—", age="—",
        patient_name="اسم المريض — معاينة تصميم فقط", patient_id="0000",
        referring_doctor_name="—", is_design_preview=True, preview_test_id=test_definition_id,
    )


if __name__ == "__main__":
    init_db()
    astm_host.start_listener_if_enabled()

    import threading
    threading.Thread(target=whatsapp_background_worker, daemon=True).start()
    # يحمّل توكن قراءة GitHub من الإعدادات المحلية (settings.github_read_token)
    # إلى الذاكرة قبل تشغيل خيط التحديث — التوكن نفسه ما يُكتب أبداً بكود
    # auto_updater.py حتى ما يترفع مع git push ويُلغى تلقائياً من GitHub
    # (راجع الشرح المفصّل بأعلى auto_updater.py).
    _startup_db = get_db()
    auto_updater.load_token_from_db(_startup_db)
    _startup_db.close()
    # خيط فحص التحديث التلقائي — يشتغل دائماً بالخلفية، لكن ما يتحقق فعلياً
    # من GitHub إلا إذا كان هذا الجهاز "مربوط بالإنترنت" من لوحة المصمم
    # (settings.auto_update_enabled = 1). راجع auto_updater.py.
    threading.Thread(target=auto_updater.background_loop, args=(get_db,), daemon=True).start()

    app.run(host="0.0.0.0", port=9090, debug=False)
